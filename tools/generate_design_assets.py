from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import tempfile
import xml.etree.ElementTree as ET

import svgwrite
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageColor, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.graphics import renderPM
from reportlab.graphics.shapes import Drawing, Rect
from svglib.svglib import svg2rlg


ROOT = Path(__file__).resolve().parents[1]
DOCUMENTS_DIR = ROOT / 'documents'
IMAGES_DIR = DOCUMENTS_DIR / 'images'
DOCX_PATH = DOCUMENTS_DIR / '系统详细设计文档.docx'
PNG_PADDING = 72
SVG_FONT = 'Microsoft YaHei'
SVG_NS = 'http://www.w3.org/2000/svg'
FONT_SEARCH_DIRS = [Path('C:/Windows/Fonts')]
FONT_CANDIDATES = {
    'normal': ['msyh.ttc', 'msyh.ttf', 'Deng.ttf', 'NotoSansSC-VF.ttf', 'simsun.ttc'],
    'bold': ['msyhbd.ttc', 'msyhbd.ttf', 'Dengb.ttf', 'NotoSansSC-VF.ttf', 'simhei.ttf'],
}


@dataclass(frozen=True)
class FigureAsset:
    key: str
    title: str
    svg_name: str
    png_name: str


FIGURES = [
    FigureAsset('module', '功能模块架构图', 'functional-module-architecture.svg', 'functional-module-architecture.png'),
    FigureAsset('business', '业务主流程图', 'business-process-flow.svg', 'business-process-flow.png'),
    FigureAsset('approval', '审批流程图', 'approval-flow.svg', 'approval-flow.png'),
    FigureAsset('er', '数据库 ER 图', 'database-er-diagram.svg', 'database-er-diagram.png'),
    FigureAsset('network', '网络拓扑图', 'network-topology.svg', 'network-topology.png'),
]


METHODS = [
    'Ontology: 以“博士生”为核心本体，统一学生、导师、团队、计划、培养、学位、成果、字典与审计实体。',
    '斯坦福七步法: 识别领域范围、核心概念、属性、对象关系、约束与实例化策略。',
    'Palantir 风格数据建模: 强调统一语义、单一事实来源、过程可追踪和跨系统联动。',
    '数据库治理原则: PostgreSQL 作为唯一运行数据源，页面选项、状态色和对象关系必须与数据库同步演进。',
]


ROLES = [
    ('学生', '查看个人档案、提交科研报告、研修申请、论文和评价问卷。'),
    ('导师', '确认导师关系、制定培养方案、审阅科研报告、组织预答辩。'),
    ('学合管理员', '维护主数据、配置流程、监督招生与学位流程、执行最终治理。'),
    ('评分人', '承担材料评审、评分推荐与异常复核。'),
    ('面试官', '执行面试分组、面试评分与成绩校算。'),
    ('HRBP', '确认实习与外出研修状态，联动中心管理。'),
    ('公寓保障', '配合在离校与住宿状态联动。'),
    ('党群负责人', '处理思政考核与资助资格审查。'),
]


MODULES = [
    ('学生主数据管理', '建立学生唯一主档案，并通过导师、团队、状态等受控主数据沉淀全周期信息。'),
    ('招生管理', '覆盖报名列表、资格审核、评分推荐、面试安排、预录取与录取闭环。'),
    ('导师关系管理', '维护新生关系确认、变更审批和权限联动。'),
    ('团队主数据管理', '围绕团队编码、负责人导师、团队导师集合、研究方向和学生归属统计进行治理。'),
    ('培养过程管理', '实现培养方案、科研报告、外出研修的查询筛选、字典选择、批量处理与成果沉淀。'),
    ('外出研修管理', '覆盖申请、审批、状态同步、归来评估与月度提醒。'),
    ('学位与毕业管理', '处理论文查重、盲审、答辩、授位与毕业离校。'),
    ('数据分析驾驶舱', '汇总招生效率、培养质量、学位进展和风险预警。'),
    ('权限与安全管理', '实现 RBAC、JWT、Redis 会话、字段级控制、审计与通知治理。'),
    ('系统集成管理', '联通招生系统、OA、飞书、科研、HR 与财务。'),
    ('通知与审计中心', '统一消息投递、回执追踪、日志归档和同步补偿。'),
    ('字典与主数据治理', '统一管理状态字典、颜色字典、枚举项和页面选择器来源，支撑全站配置收口。'),
]


BUSINESS_FLOW_DESCRIPTION = [
    '招生录取完成后，将录取结果与导师初始关系同步至学生主数据。',
    '学生入学后 15 日内导师制定培养方案，学生需在 7 日内确认。',
    '科研报告按照周期生成待办，逾期 7 日提醒导师，14 日升级至管理员。',
    '论文通过查重与盲审后，进入预答辩、正式答辩和授位审批。',
    '毕业离校完成后保留电子档案，并继续纳入就业跟踪分析。',
]


APPROVAL_FLOW_DESCRIPTION = [
    '导师变更审批: 支持学生申请、导师发起和管理员统筹三类入口，要求全过程留痕。',
    '外出研修审批: 学生提交后先由导师审核，再由学合管理员备案并同步飞书/OA。',
    '学位申请审批: 论文查重、盲审、预答辩、正式答辩和授位审议构成严格串行链路。',
]


DATABASE_GROUPS = {
    '系统治理类表': [
        ('dtlms_users', '系统账号主表', 'username, password_hash, is_active'),
        ('dtlms_roles', '角色定义表', 'role_code, role_name'),
        ('dtlms_permissions', '权限点定义表', 'permission_code, module_name'),
        ('dtlms_user_roles', '用户角色映射表', 'user_id, role_id'),
        ('dtlms_role_permissions', '角色权限映射表', 'role_id, permission_id'),
        ('dtlms_login_logs', '登录日志表', 'username, login_status, login_ip'),
        ('dtlms_operation_logs', '操作日志表', 'module_name, entity_name, old_value, new_value'),
        ('dtlms_data_sync_logs', '数据同步日志表', 'source_system, target_system, sync_status'),
        ('dtlms_notification_templates', '通知模板表', 'template_code, channel, content_template'),
        ('dtlms_system_configs', '系统参数表', 'config_key, config_value'),
        ('dtlms_dict_types', '字典类型表', 'dict_name, dict_type, status'),
        ('dtlms_dict_data', '字典数据表', 'dict_type, label, value, color_type'),
    ],
    '学生与培养类表': [
        ('dtlms_advisors', '导师主数据表', 'advisor_no, full_name, annual_quota'),
        ('dtlms_teams', '团队主数据表', 'team_code, team_name, department_name, lead_advisor_id, team_status'),
        ('dtlms_team_advisors', '团队导师映射表', 'team_id, advisor_id, advisor_role, joined_on, left_on'),
        ('dtlms_students', '学生主数据表', 'student_no, full_name, current_status, team_id, primary_advisor_id'),
        ('dtlms_student_team_history', '学生团队历史表', 'student_id, team_id, start_date, end_date'),
        ('dtlms_student_advisor_history', '导师关系历史表', 'student_id, advisor_id, start_date, end_date'),
        ('dtlms_research_projects', '科研项目表', 'project_code, project_name, principal_advisor_id'),
        ('dtlms_training_plans', '培养方案表', 'student_id, advisor_id, version_no, plan_status'),
        ('dtlms_training_plan_versions', '培养方案版本表', 'training_plan_id, version_no, plan_snapshot'),
        ('dtlms_scientific_reports', '科研报告表', 'student_id, period_label, report_status, review_score'),
        ('dtlms_outbound_studies', '外出研修表', 'student_id, study_type, destination, approval_status'),
        ('dtlms_achievements', '成果记录表', 'student_id, achievement_type, title'),
        ('dtlms_theses', '学位论文表', 'student_id, plagiarism_rate, thesis_status, degree_granted'),
        ('dtlms_thesis_reviews', '盲审评审表', 'thesis_id, expert_name, review_score, review_status'),
    ],
    '招生类表': [
        ('dtlms_recruitment_plans', '招生计划表', 'plan_code, academic_year, start_date, plan_status'),
        ('dtlms_research_fields', '研究领域表', 'field_code, field_name'),
        ('dtlms_recruitment_applications', '报名申请表', 'plan_id, candidate_no, application_status'),
        ('dtlms_application_materials', '申请材料表', 'application_id, material_type, material_status'),
        ('dtlms_qualification_reviews', '资格审核表', 'application_id, reviewer_username, review_status'),
        ('dtlms_reviewer_assignments', '评分人分配表', 'application_id, reviewer_username, reviewer_role'),
        ('dtlms_material_scores', '材料评分表', 'application_id, reviewer_assignment_id, material_score'),
        ('dtlms_interview_groups', '面试组表', 'plan_id, group_code, interview_mode'),
        ('dtlms_interview_schedules', '面试日程表', 'application_id, interview_group_id, admission_ticket_no'),
        ('dtlms_interview_scores', '面试评分表', 'schedule_id, evaluator_username, interview_score'),
        ('dtlms_written_exam_scores', '机试成绩表', 'application_id, exam_score, import_batch_no'),
        ('dtlms_admission_decisions', '预录取与录取表', 'application_id, decision_status, final_score'),
    ],
    '运行态镜像与同步类表': [
        ('dtlms_runtime_counters', '运行态计数器表', 'counter_name, counter_value'),
        ('dtlms_runtime_profiles', '运行态个人资料表', 'username, payload'),
        ('dtlms_runtime_students', '运行态学生镜像表', 'id, payload'),
        ('dtlms_runtime_recruitment_plans', '运行态招生计划镜像表', 'id, payload'),
        ('dtlms_runtime_recruitment_applications', '运行态申请镜像表', 'id, payload'),
        ('dtlms_runtime_training_plans', '运行态培养方案镜像表', 'id, payload'),
        ('dtlms_runtime_scientific_reports', '运行态科研报告镜像表', 'id, payload'),
        ('dtlms_runtime_outbound_studies', '运行态外出研修镜像表', 'id, payload'),
        ('dtlms_runtime_theses', '运行态论文镜像表', 'id, payload'),
        ('dtlms_runtime_thesis_reviews', '运行态盲审镜像表', 'id, payload'),
        ('dtlms_runtime_roles', '运行态角色镜像表', 'id, payload'),
        ('dtlms_runtime_system_users', '运行态系统用户镜像表', 'id, payload'),
        ('dtlms_runtime_audit_policies', '运行态审计策略镜像表', 'id, payload'),
        ('dtlms_runtime_integrations', '运行态集成链路镜像表', 'id, payload'),
        ('dtlms_runtime_operation_logs', '运行态操作日志镜像表', 'id, payload'),
        ('dtlms_runtime_sync_logs', '运行态同步日志镜像表', 'id, payload'),
        ('dtlms_runtime_workflow_tasks', '运行态流程任务镜像表', 'id, payload'),
    ],
}


RELATIONSHIPS = [
    ('字典类型 -> 字典数据', '一对多关系，所有稳定枚举值、状态值和颜色值统一通过字典表治理。'),
    ('团队 -> 导师', '多对多治理关系，lead_advisor_id 标识负责人导师，团队成员通过映射表维护。'),
    ('团队 -> 学生', '一对多归属关系，当前归属通过 students.team_id 表达，变更过程进入历史表。'),
    ('学生 -> 导师', '一对多历史关系，当前主导师通过 primary_advisor_id 指向。'),
    ('学生 -> 培养方案', '一对多版本关系，最近生效版本参与流程计算。'),
    ('学生 -> 科研报告', '一对多周期关系，支持周期催办和评分。'),
    ('学生 -> 外出研修', '一对多审批记录，记录审批状态与归来评估。'),
    ('学生 -> 学位论文', '一对多版本语义，当前论文进入盲审和答辩流程。'),
    ('招生计划 -> 申请', '一对多，申请贯穿资格审核、评分、面试与录取。'),
    ('申请 -> 评分人/面试日程/录取结果', '一条申请关联多个评分记录、一个日程和一个最终录取结果。'),
    ('用户 -> 角色 -> 权限', '通过映射表实现 RBAC 可配置授权。'),
]


VIEWS = [
    ('dtlms_v_student_lifecycle_snapshot', '学生生命周期快照视图', '供驾驶舱和学生全景卡片直接读取。'),
    ('dtlms_v_recruitment_dashboard', '招生驾驶舱聚合视图', '统计申请量、通过量、预录取量和平均材料分。'),
    ('dtlms_v_training_compliance', '培养合规视图', '监控培养方案、科研报告与外出研修状态。'),
    ('dtlms_v_degree_pipeline', '学位流程视图', '聚合论文、盲审、答辩与授位节点。'),
]


DATABASE_DESIGN_NOTES = [
    '关系主表与运行态镜像表共存，但业务系统运行期间只允许 PostgreSQL 作为唯一事实来源，不再依赖 JSON 演示文件。',
    '字典设计采用“字典类型 + 字典数据”两层结构，既管理 label/value，也管理 color_type、css_class 等展示语义。',
    '学生、导师、团队、招生、培养、学位均通过外键或历史表保持可追溯关系，避免页面层自由拼装对象关系。',
    '运行态镜像表承载 RuntimeManagementStore 的统一持久化结果，确保内存状态、接口返回和数据库对象持续同步。',
    '所有表采用 created_at、updated_at 与 is_deleted 的统一治理字段，支撑审计、回放与逻辑删除策略。',
]


INTEGRATIONS = [
    ('招生系统', '录取主数据导入、录取结果回传', '实时 + 每日对账'),
    ('实验室 OA', '考勤、门禁、请假、周报同步', '事件驱动 + 定时补偿'),
    ('飞书', '待办提醒、状态推送、回执跟踪', '实时'),
    ('科研系统', '成果和项目共享', '月度'),
    ('HR / 财务', '导师在岗、奖助与欠费状态', '月度'),
]


SECURITY_ITEMS = [
    ('认证授权', 'JWT + RBAC + Redis Session，控制接口与页面访问，并支持登出后立即失效。'),
    ('日志审计', '登录日志、操作日志、同步日志三层留痕，关键表保留 5 年。'),
    ('数据安全', '敏感字段脱敏导出，逻辑删除替代物理删除。'),
    ('缓存治理', 'Redis Sentinel + CTDTLMS_ 统一前缀，支撑登录会话、提醒状态与驾驶舱热点缓存。'),
    ('运维监控', 'Prometheus/Grafana 指标监控与日志归档联动。'),
]


AUTH_RUNTIME_NOTES = [
    '后端运行时读取 backend/.env 和 backend/.env.local，backend/.env.example 仅作为模板文件。',
    '登录接口在用户名密码校验通过后创建 Redis 会话，并将 sid 写入 access token 与 refresh token。',
    '前端路由守卫会记录受保护页面目标地址，登录成功后优先恢复 redirect 指向的页面。',
    '前端在收到 401 时自动清理本地 token，跳回登录页，并保留原始页面地址供重新登录后恢复。',
    '注销接口会撤销 Redis 会话，旧 token 会立即失效，避免仅删除前端 localStorage 造成的伪注销。',
]


TRAINING_RUNTIME_NOTES = [
    '培养方案、科研报告、外出研修三类页面均提供关键字、状态、导师或审阅人等查询条件。',
    '状态类字段统一改为字典选择，不再依赖自由文本录入。',
    '列表页统一支持单条删除与批量删除，按钮文案按业务语义显示。',
    '培养管理接口新增 /training/options，用于前端加载报告周期、方案状态、研修类型等字典项。',
]


STUDENT_TEAM_RUNTIME_NOTES = [
    '学生主档新增与编辑时，所属团队和导师均改为主数据受控选择，不再允许自由文本录入。',
    '团队实体采用本体化建模，核心属性包括团队编码、负责人导师、团队导师集合、研究方向、状态与成立日期。',
    '团队与学生、导师形成稳定关联，保存学生主档时会校验导师是否属于所选团队。',
    '团队管理页支持查询、状态筛选、负责人过滤、单条删除与批量删除，并实时统计团队学生数。',
]


def ensure_dirs() -> None:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def add_box(dwg: svgwrite.Drawing, x: int, y: int, width: int, height: int, title: str, subtitle: str = '', fill: str = '#eff6ff') -> None:
    dwg.add(
        dwg.rect(
            insert=(x, y),
            size=(width, height),
            rx=18,
            ry=18,
            fill=fill,
            stroke='#183a75',
            stroke_width=2,
        )
    )
    dwg.add(dwg.text(title, insert=(x + 16, y + 26), font_size=18, font_family=SVG_FONT, fill='#12284d', font_weight='bold'))
    if subtitle:
        for index, line in enumerate(subtitle.split('\n')):
            dwg.add(dwg.text(line, insert=(x + 16, y + 52 + index * 18), font_size=13, font_family=SVG_FONT, fill='#4d6287'))


def add_arrow(dwg: svgwrite.Drawing, start: tuple[int, int], end: tuple[int, int], label: str = '') -> None:
    x1, y1 = start
    x2, y2 = end
    dwg.add(dwg.line(start=start, end=end, stroke='#40649a', stroke_width=3))

    if x1 == x2:
        direction = 1 if y2 > y1 else -1
        arrow = [(x2, y2), (x2 - 8, y2 - 14 * direction), (x2 + 8, y2 - 14 * direction)]
    else:
        direction = 1 if x2 > x1 else -1
        arrow = [(x2, y2), (x2 - 14 * direction, y2 - 8), (x2 - 14 * direction, y2 + 8)]

    dwg.add(dwg.polygon(points=arrow, fill='#40649a'))
    if label:
        label_x = (x1 + x2) / 2
        label_y = (y1 + y2) / 2 - 10
        dwg.add(dwg.text(label, insert=(label_x, label_y), font_size=12, font_family=SVG_FONT, fill='#40649a', text_anchor='middle'))


def create_drawing(path: Path, width: int, height: int, background: str = '#fbfdff') -> svgwrite.Drawing:
    dwg = svgwrite.Drawing(str(path), size=(f'{width}px', f'{height}px'))
    dwg.viewbox(0, 0, width, height)
    dwg.add(dwg.rect(insert=(0, 0), size=(width, height), fill=background))
    return dwg


def render_module_architecture(path: Path) -> None:
    dwg = create_drawing(path, 1520, 940, '#f7fbff')
    dwg.add(dwg.text('博士生生命周期管理系统功能架构图', insert=(760, 50), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#10284d'))

    add_box(dwg, 550, 90, 300, 84, '门户展现层', 'Vue3 + Element Plus\n驾驶舱 / 工作台 / 检索', '#fff3df')
    add_box(dwg, 550, 220, 300, 84, '网关接入层', '认证 / 路由 / 审计 / API 文档', '#e7f0ff')

    service_specs = [
        (110, 370, '主数据域', '学生 / 导师 / 团队 / 字典'),
        (410, 370, '招生业务域', '报名 / 审核 / 面试 / 录取'),
        (710, 370, '培养业务域', '方案 / 报告 / 研修 / 成果'),
        (1010, 370, '学位业务域', '论文 / 盲审 / 答辩 / 授位'),
    ]
    for x, y, title, subtitle in service_specs:
        add_box(dwg, x, y, 270, 110, title, subtitle, '#eef7ff')

    add_box(dwg, 180, 570, 360, 110, '治理中心', 'RBAC / JWT / 登录审计\n操作日志 / 同步日志 / 字典管理', '#ebfff5')
    add_box(dwg, 610, 570, 360, 110, '集成中心', '招生系统 / OA / 飞书\n科研 / HR / 财务', '#fff7eb')
    add_box(dwg, 1040, 570, 220, 110, '分析驾驶舱', 'KPI / 预警 / 下钻', '#f1edff')
    add_box(dwg, 410, 730, 580, 84, '数据层', 'PostgreSQL + Redis Sentinel + MinIO', '#edf7ff')

    add_arrow(dwg, (700, 174), (700, 220))
    for x in (245, 545, 845, 1145):
        add_arrow(dwg, (700, 304), (x, 370))
    add_arrow(dwg, (320, 480), (360, 570))
    add_arrow(dwg, (845, 480), (790, 570))
    add_arrow(dwg, (1145, 480), (1150, 570))
    add_arrow(dwg, (360, 680), (540, 730))
    add_arrow(dwg, (790, 680), (700, 730))
    add_arrow(dwg, (1150, 680), (860, 730))
    dwg.save()


def render_business_flow(path: Path) -> None:
    dwg = create_drawing(path, 1740, 620, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期端到端业务流程图', insert=(870, 46), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#10284d'))

    stages = [
        ('招生管理', '报名 / 审核 / 评分'),
        ('录取入学', '面试 / 预录取 / 入学'),
        ('导师关系', '确认 / 变更 / 授权'),
        ('培养方案', '制定 / 确认 / 版本'),
        ('科研报告', '提交 / 审阅 / 预警'),
        ('外出研修', '申请 / 审批 / 评估'),
        ('学位流程', '查重 / 盲审 / 答辩'),
        ('毕业跟踪', '归档 / 离校 / 就业'),
    ]
    x = 70
    for title, subtitle in stages:
        add_box(dwg, x, 180, 170, 100, title, subtitle, '#eef5ff')
        x += 185

    x = 155
    for label in ['主数据同步', '角色联动', '规则校验', '自动提醒', '审批留痕', '授位决议', '档案归档']:
        add_arrow(dwg, (x, 230), (x + 105, 230), label)
        x += 185

    add_box(dwg, 520, 340, 220, 96, '审计留痕', '登录 / 操作 / 同步', '#ebfff5')
    add_box(dwg, 840, 340, 220, 96, '驾驶舱', 'KPI / 预警 / 趋势', '#fff8e8')
    add_arrow(dwg, (630, 280), (630, 340))
    add_arrow(dwg, (950, 280), (950, 340))
    dwg.save()


def render_approval_flow(path: Path) -> None:
    dwg = create_drawing(path, 1700, 1080, '#fbfdff')
    dwg.add(dwg.text('审批流程矩阵图', insert=(850, 44), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#10284d'))

    columns = [
        (90, '导师变更审批', ['学生/导师发起', '管理员审核', '原导师意见', '新导师确认', '权限切换']),
        (580, '外出研修审批', ['学生申请', '导师审核', '管理员备案', '飞书/OA 同步', '归来评估']),
        (1070, '学位申请审批', ['提交论文', '查重核验', '3 位盲审', '预答辩', '答辩与授位']),
    ]

    for x, title, steps in columns:
        add_box(dwg, x, 90, 390, 74, title, '', '#eaf4ff')
        y = 210
        for index, step in enumerate(steps):
            fill = '#eef8ff' if index % 2 == 0 else '#fff7e8'
            add_box(dwg, x + 20, y, 350, 82, step, '时限控制 / 审计留痕', fill)
            if index < len(steps) - 1:
                add_arrow(dwg, (x + 195, y + 82), (x + 195, y + 120), '')
            y += 128

    dwg.save()


def render_er_diagram(path: Path) -> None:
    dwg = create_drawing(path, 1800, 1100, '#fbfdff')
    dwg.add(dwg.text('核心数据库 ER 图', insert=(900, 44), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#10284d'))

    add_box(dwg, 680, 180, 300, 96, '学生主表', 'student_no\ncurrent_status\nprimary_advisor_id', '#eaf5ff')
    add_box(dwg, 330, 180, 240, 96, '导师主表', 'advisor_no\nannual_quota', '#fff7e8')
    add_box(dwg, 1040, 120, 280, 96, '培养方案表', 'student_id\nversion_no\nplan_status', '#eefcf2')
    add_box(dwg, 1040, 280, 280, 96, '科研报告表', 'student_id\nperiod_label\nreport_status', '#eefcf2')
    add_box(dwg, 1040, 440, 280, 96, '论文主表', 'student_id\nthesis_status\ndegree_status', '#eefcf2')
    add_box(dwg, 1040, 600, 280, 96, '盲审意见表', 'thesis_id\nexpert_name\nreview_score', '#eefcf2')
    add_box(dwg, 330, 380, 240, 96, '外出研修表', 'student_id\napproval_status', '#fff7e8')
    add_box(dwg, 330, 540, 240, 96, '成果记录表', 'student_id\nachievement_type\ntitle', '#fff7e8')
    add_box(dwg, 80, 640, 320, 96, '操作日志表', 'module_name\nentity_name\naction', '#f4f0ff')
    add_box(dwg, 80, 140, 220, 96, '系统用户表', 'username\npassword_hash', '#f4f0ff')
    add_box(dwg, 80, 300, 220, 96, '角色表', 'role_code\nrole_name', '#f4f0ff')
    add_box(dwg, 80, 460, 220, 96, '权限表', 'permission_code\nmodule_name', '#f4f0ff')
    add_box(dwg, 510, 720, 320, 96, '招生计划表', 'plan_code\nacademic_term\nplan_status', '#fff4ea')
    add_box(dwg, 900, 720, 320, 96, '招生申请表', 'candidate_no\napplication_status', '#fff4ea')
    add_box(dwg, 1290, 720, 250, 96, '录取决策表', 'decision_status\nfinal_score', '#fff4ea')

    add_arrow(dwg, (570, 228), (680, 228), '指导')
    add_arrow(dwg, (980, 228), (1040, 168), '拥有')
    add_arrow(dwg, (980, 228), (1040, 328), '提交')
    add_arrow(dwg, (980, 228), (1040, 488), '撰写')
    add_arrow(dwg, (1160, 536), (1160, 600), '评审')
    add_arrow(dwg, (570, 428), (680, 228), '关联')
    add_arrow(dwg, (570, 588), (680, 228), '沉淀')
    add_arrow(dwg, (400, 688), (680, 260), '审计')
    add_arrow(dwg, (300, 188), (300, 300), '关联')
    add_arrow(dwg, (300, 348), (300, 460), '授权')
    add_arrow(dwg, (830, 768), (900, 768), '接收')
    add_arrow(dwg, (1220, 768), (1290, 768), '决策')
    add_arrow(dwg, (900, 768), (830, 276), '转化学生')
    dwg.save()


def render_network_topology(path: Path) -> None:
    dwg = create_drawing(path, 1780, 980, '#fbfdff')
    dwg.add(dwg.text('网络拓扑与部署关系图', insert=(900, 44), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#10284d'))

    add_box(dwg, 80, 200, 240, 110, '用户终端', '学生 / 导师\n管理员 / 评审人', '#eef5ff')
    add_box(dwg, 380, 200, 220, 110, '接入网关', 'WAF + Nginx\nHTTPS / SSL / 路由', '#fff7e8')
    add_box(dwg, 680, 140, 260, 110, '前端门户', 'Vue3 / Element Plus\nVite 部署', '#eefcf2')
    add_box(dwg, 680, 320, 260, 110, '后端服务', 'FastAPI / JWT / RBAC\n审计中间件', '#eefcf2')
    add_box(dwg, 1020, 120, 240, 96, '异步任务节点', 'Celery Worker\n提醒 / 异步作业', '#f4f0ff')
    add_box(dwg, 1020, 270, 240, 96, '缓存哨兵集群', 'Redis Sentinel\n缓存 / 任务状态 / CTDTLMS_*', '#fff4ea')
    add_box(dwg, 1020, 430, 240, 96, '业务数据库', 'PostgreSQL\n主数据 / 视图 / 分析来源', '#fff4ea')
    add_box(dwg, 1320, 120, 220, 96, '对象存储', 'MinIO\n论文 / 报告 / 材料', '#eef5ff')
    add_box(dwg, 1320, 300, 220, 220, '外部系统', '招生系统\nOA\n飞书\n科研系统\nHR\n财务', '#eef5ff')
    add_box(dwg, 640, 560, 700, 140, '可观测性平台', 'Prometheus / Grafana / 日志归档 / 同步补偿', '#f7f3ff')

    add_arrow(dwg, (320, 255), (380, 255), 'HTTPS')
    add_arrow(dwg, (600, 220), (680, 195), '静态资源')
    add_arrow(dwg, (600, 290), (680, 360), '接口调用')
    add_arrow(dwg, (940, 360), (1020, 318), '缓存 / Broker')
    add_arrow(dwg, (940, 360), (1020, 478), 'SQL')
    add_arrow(dwg, (940, 360), (1020, 168), '任务派发')
    add_arrow(dwg, (1260, 168), (1320, 168), '文件读写')
    add_arrow(dwg, (940, 360), (1320, 410), 'REST / Webhook')
    add_arrow(dwg, (1140, 526), (1140, 560), '指标 / 日志')
    dwg.save()


def resolve_font_path(weight: str = 'normal') -> Path:
    candidates = FONT_CANDIDATES['bold' if weight == 'bold' else 'normal']
    for base_dir in FONT_SEARCH_DIRS:
        for candidate in candidates:
            font_path = base_dir / candidate
            if font_path.exists():
                return font_path
    raise FileNotFoundError(f'No suitable CJK font found for weight={weight}.')


def parse_svg_length(raw_value: str | None, fallback: float = 0.0) -> float:
    if not raw_value:
        return fallback
    value = raw_value.strip().lower().replace('px', '')
    try:
        return float(value)
    except ValueError:
        return fallback


def strip_svg_text(svg_path: Path) -> tuple[bytes, list[dict[str, str]], tuple[float, float]]:
    tree = ET.parse(svg_path)
    root = tree.getroot()
    texts: list[dict[str, str]] = []
    svg_width = parse_svg_length(root.attrib.get('width'))
    svg_height = parse_svg_length(root.attrib.get('height'))

    view_box = root.attrib.get('viewBox', '')
    if view_box:
        parts = [item for item in view_box.replace(',', ' ').split() if item]
        if len(parts) == 4:
            svg_width = parse_svg_length(parts[2], svg_width)
            svg_height = parse_svg_length(parts[3], svg_height)

    for parent in root.iter():
        children = list(parent)
        for child in children:
            if child.tag == f'{{{SVG_NS}}}text':
                texts.append({
                    'text': ''.join(child.itertext()),
                    'x': child.attrib.get('x', '0'),
                    'y': child.attrib.get('y', '0'),
                    'font_size': child.attrib.get('font-size', '12'),
                    'fill': child.attrib.get('fill', '#000000'),
                    'text_anchor': child.attrib.get('text-anchor', 'start'),
                    'font_weight': child.attrib.get('font-weight', 'normal'),
                })
                parent.remove(child)

    return ET.tostring(root, encoding='utf-8', xml_declaration=True), texts, (svg_width, svg_height)


def draw_svg_texts(image: Image.Image, texts: list[dict[str, str]], scale_x: float, scale_y: float) -> None:
    draw = ImageDraw.Draw(image)
    normal_font_path = resolve_font_path('normal')
    bold_font_path = resolve_font_path('bold')

    for text_item in texts:
        content = text_item['text'].strip()
        if not content:
            continue

        font_size = max(10, int(float(text_item['font_size'])))
        weight = 'bold' if text_item['font_weight'] == 'bold' else 'normal'
        font_path = bold_font_path if weight == 'bold' else normal_font_path
        scaled_font_size = max(8, int(round(font_size * min(scale_x, scale_y))))
        font = ImageFont.truetype(str(font_path), scaled_font_size)
        x = float(text_item['x']) * scale_x + PNG_PADDING
        y = float(text_item['y']) * scale_y + PNG_PADDING
        fill = ImageColor.getrgb(text_item['fill'])
        anchor = 'ms' if text_item['text_anchor'] == 'middle' else 'ls'
        draw.text((x, y), content, font=font, fill=fill, anchor=anchor)


def convert_svg_to_png(svg_path: Path, png_path: Path) -> None:
    stripped_svg, text_items, (svg_width, svg_height) = strip_svg_text(svg_path)
    with tempfile.NamedTemporaryFile(delete=False, suffix='.svg') as temp_svg:
        temp_svg.write(stripped_svg)
        temp_svg_path = Path(temp_svg.name)

    drawing = svg2rlg(str(temp_svg_path))
    wrapped = Drawing(drawing.width + PNG_PADDING * 2, drawing.height + PNG_PADDING * 2)
    wrapped.add(
        Rect(
            0,
            0,
            drawing.width + PNG_PADDING * 2,
            drawing.height + PNG_PADDING * 2,
            fillColor=colors.HexColor('#fbfdff'),
            strokeColor=None,
        )
    )
    group = drawing.asGroup()
    group.translate(PNG_PADDING, PNG_PADDING)
    wrapped.add(group)
    renderPM.drawToFile(wrapped, str(png_path), fmt='PNG')
    temp_svg_path.unlink(missing_ok=True)

    image = Image.open(png_path).convert('RGBA')
    scale_x = drawing.width / svg_width if svg_width else 1.0
    scale_y = drawing.height / svg_height if svg_height else 1.0
    draw_svg_texts(image, text_items, scale_x, scale_y)
    image.save(png_path)


def configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = document.styles['Normal']
    normal.font.name = '微软雅黑'
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for style_name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        style = document.styles[style_name]
        style.font.name = '微软雅黑'
        style.font.bold = True
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def apply_paragraph_style(paragraph, font_size: int = 12, bold: bool = False, center: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.2
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.add_run(item)
        apply_paragraph_style(paragraph)


def add_table(document: Document, title: str, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    heading = document.add_paragraph(title, style='Heading 3')
    apply_paragraph_style(heading, font_size=12, bold=True)

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, cell_value in enumerate(row):
            cells[index].text = cell_value

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                apply_paragraph_style(paragraph, font_size=10)


def add_figure(document: Document, figure: FigureAsset, width_cm: float = 15.2) -> None:
    png_path = IMAGES_DIR / figure.png_name
    document.add_picture(str(png_path), width=Cm(width_cm))
    caption = document.add_paragraph(f'图：{figure.title}')
    apply_paragraph_style(caption, font_size=10, center=True)


def generate_document() -> None:
    document = Document()
    configure_document_styles(document)

    title = document.add_paragraph()
    title.add_run('系统详细设计文档')
    apply_paragraph_style(title, font_size=18, bold=True, center=True)

    subtitle = document.add_paragraph()
    subtitle.add_run('博士生生命周期管理系统（DTLMS）')
    apply_paragraph_style(subtitle, font_size=14, center=True)

    document.add_paragraph('')

    overview = document.add_paragraph(style='Heading 1')
    overview.add_run('1. 文档概述')
    apply_paragraph_style(overview, font_size=16, bold=True)
    paragraph = document.add_paragraph(
        '本文档基于需求分析报告、建设方案、招生模块截图与统一需求摘要，形成博士生生命周期管理系统的详细设计基线。设计覆盖总体架构、数据库、流程、审批、模块、网络拓扑、安全审计与实施约束。'
    )
    apply_paragraph_style(paragraph)

    methodology = document.add_paragraph(style='Heading 1')
    methodology.add_run('2. 设计方法与原则')
    apply_paragraph_style(methodology, font_size=16, bold=True)
    add_bullet_list(document, METHODS)

    role_section = document.add_paragraph(style='Heading 1')
    role_section.add_run('3. 角色与治理边界')
    apply_paragraph_style(role_section, font_size=16, bold=True)
    add_table(document, '3.1 角色清单', ['角色', '职责边界'], ROLES)

    architecture = document.add_paragraph(style='Heading 1')
    architecture.add_run('4. 总体架构设计')
    apply_paragraph_style(architecture, font_size=16, bold=True)
    text = document.add_paragraph(
        '系统采用前后端分离模式。前端使用 Vue3 + Element Plus 构建管理门户，后端使用 Python3 FastAPI 输出 API 与业务服务，数据库使用 PostgreSQL，缓存与异步任务依赖 Redis Sentinel，统一通过 RBAC 与 JWT 管理访问控制，并通过字典治理保证全站状态与选项来源一致。'
    )
    apply_paragraph_style(text)
    add_figure(document, FIGURES[0])

    module_section = document.add_paragraph(style='Heading 2')
    module_section.add_run('4.1 功能模块说明')
    apply_paragraph_style(module_section, font_size=14, bold=True)
    add_table(document, '功能模块描述', ['模块', '设计说明'], MODULES)

    auth_runtime_section = document.add_paragraph(style='Heading 2')
    auth_runtime_section.add_run('4.2 认证与会话运行设计')
    apply_paragraph_style(auth_runtime_section, font_size=14, bold=True)
    add_bullet_list(document, AUTH_RUNTIME_NOTES)

    process_section = document.add_paragraph(style='Heading 1')
    process_section.add_run('5. 业务流程与审批设计')
    apply_paragraph_style(process_section, font_size=16, bold=True)
    process_intro = document.add_paragraph('业务流程以“招生 -> 入学 -> 培养 -> 学位 -> 毕业 -> 就业跟踪”为主线，同时通过通知、审计和驾驶舱形成治理闭环。')
    apply_paragraph_style(process_intro)
    add_figure(document, FIGURES[1])
    add_bullet_list(document, BUSINESS_FLOW_DESCRIPTION)
    add_figure(document, FIGURES[2])
    add_bullet_list(document, APPROVAL_FLOW_DESCRIPTION)

    training_runtime_section = document.add_paragraph(style='Heading 2')
    training_runtime_section.add_run('5.1 培养管理治理增强')
    apply_paragraph_style(training_runtime_section, font_size=14, bold=True)
    add_bullet_list(document, TRAINING_RUNTIME_NOTES)

    team_runtime_section = document.add_paragraph(style='Heading 2')
    team_runtime_section.add_run('5.2 学生与团队主数据治理设计')
    apply_paragraph_style(team_runtime_section, font_size=14, bold=True)
    add_bullet_list(document, STUDENT_TEAM_RUNTIME_NOTES)

    database_section = document.add_paragraph(style='Heading 1')
    database_section.add_run('6. 数据库设计')
    apply_paragraph_style(database_section, font_size=16, bold=True)
    database_intro = document.add_paragraph('数据库采用 PostgreSQL 关系建模，核心策略包括：主数据唯一、关键对象版本化、字典统一治理、日志软删除、统计视图前置和关键流程关系清晰可追溯。')
    apply_paragraph_style(database_intro)
    add_figure(document, FIGURES[3])
    add_bullet_list(document, DATABASE_DESIGN_NOTES)

    for group_name, rows in DATABASE_GROUPS.items():
        add_table(document, group_name, ['表名', '用途', '关键字段'], rows)

    add_table(document, '6.4 关联关系说明', ['关系', '说明'], RELATIONSHIPS)
    add_table(document, '6.5 视图设计', ['视图', '用途', '说明'], VIEWS)

    integration_section = document.add_paragraph(style='Heading 1')
    integration_section.add_run('7. 集成与驾驶舱设计')
    apply_paragraph_style(integration_section, font_size=16, bold=True)
    integration_intro = document.add_paragraph('驾驶舱建设遵循“汇总指标 + 风险预警 + 责任对象 + 下钻入口”的设计思路，面向学合管理员、导师和学生提供差异化看板。')
    apply_paragraph_style(integration_intro)
    add_table(document, '7.1 外部系统集成', ['系统', '交互内容', '节奏'], INTEGRATIONS)

    security_section = document.add_paragraph(style='Heading 1')
    security_section.add_run('8. 安全、审计与运维')
    apply_paragraph_style(security_section, font_size=16, bold=True)
    add_table(document, '8.1 安全与审计控制项', ['控制项', '设计说明'], SECURITY_ITEMS)

    network_section = document.add_paragraph(style='Heading 1')
    network_section.add_run('9. 网络拓扑与部署建议')
    apply_paragraph_style(network_section, font_size=16, bold=True)
    network_intro = document.add_paragraph('生产环境建议部署 Nginx 统一入口、FastAPI 服务实例、Celery 异步任务、PostgreSQL 数据库、Redis Sentinel 与对象存储，并接入监控与日志归档体系。')
    apply_paragraph_style(network_intro)
    add_figure(document, FIGURES[4])

    conclusion = document.add_paragraph(style='Heading 1')
    conclusion.add_run('10. 设计结论')
    apply_paragraph_style(conclusion, font_size=16, bold=True)
    add_bullet_list(
        document,
        [
            '学生主数据、导师关系、培养方案和学位论文是系统最关键的四条主轴数据。',
            '招生、培养和学位三个业务域均要求流程节点可审计、可回放和可升级提醒。',
            '字典类型与字典数据表已经纳入正式数据库设计，页面状态、选项与颜色均应优先由数据库字典驱动。',
            '运行态镜像表用于承载统一持久化结果，但系统运行期间仍以 PostgreSQL 为唯一事实来源。',
            'Redis Sentinel 统一承担缓存热点、提醒状态与异步任务状态缓存，所有 Key 使用 CTDTLMS_ 前缀。',
            '驾驶舱不是独立系统，而是建立在视图、日志和业务状态表之上的治理平面。',
        ],
    )

    document.save(DOCX_PATH)


def generate_figures() -> None:
    renderers = {
        'module': render_module_architecture,
        'business': render_business_flow,
        'approval': render_approval_flow,
        'er': render_er_diagram,
        'network': render_network_topology,
    }
    for figure in FIGURES:
        svg_path = IMAGES_DIR / figure.svg_name
        png_path = IMAGES_DIR / figure.png_name
        renderers[figure.key](svg_path)
        convert_svg_to_png(svg_path, png_path)


def main() -> None:
    ensure_dirs()
    generate_figures()
    generate_document()
    print(f'Generated design doc: {DOCX_PATH}')


if __name__ == '__main__':
    main()
