from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

try:
    from generate_design_assets import SVG_FONT, add_arrow, add_box, convert_svg_to_png, create_drawing
except ModuleNotFoundError:
    from tools.generate_design_assets import SVG_FONT, add_arrow, add_box, convert_svg_to_png, create_drawing


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "documents" / "产品需求规格说明书(EARS版本).docx"
FALLBACK_OUTPUT_PATH = ROOT / "documents" / "产品需求规格说明书(EARS版本)-更新版.docx"
EARS_IMAGES_DIR = ROOT / "documents" / "EARS_IMAGES"
BODY_FONT = "Microsoft YaHei"
BODY_FONT_FALLBACK = "微软雅黑"


FLOW_FIGURES = [
    {
        "key": "business-flow",
        "title": "博士生生命周期业务主流程图",
        "svg": "business-flow.svg",
        "png": "business-flow.png",
    },
    {
        "key": "approval-flow",
        "title": "博士生生命周期审批流程图",
        "svg": "approval-flow.svg",
        "png": "approval-flow.png",
    },
    {
        "key": "entity-relationship",
        "title": "博士生生命周期核心实体关系图",
        "svg": "entity-relationship.svg",
        "png": "entity-relationship.png",
    },
]


PROTOTYPE_SPECS = [
    {
        "key": "portal-auth",
        "title": "门户与认证原型图",
        "prototype_style": "auth",
        "section": "门户与认证",
        "route": "/portal/auth",
        "operator": "门户考生",
        "page_title": "门户认证中心",
        "summary": ["注册引导", "登录校验", "找回密码"],
        "filters": ["切换：登录/注册/找回", "验证码状态", "邮箱校验"],
        "table_headers": ["步骤", "输入项", "系统反馈", "操作"],
        "table_rows": [
            ["注册", "手机号、邮箱、姓名", "校验必填和格式", "下一步"],
            ["验证码", "图形验证码、邮箱验证码", "提示冷却和有效期", "发送验证码"],
            ["完成", "密码、确认密码", "注册成功提示", "提交注册"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["先切换认证模式", "输入账号资料", "完成双验证码校验", "提交后进入门户首页"],
        "description": "该原型图用于表达门户考生进入系统时的统一认证入口。页面应采用左右分区布局，左侧承载品牌引导与招生信息提示，右侧承载登录、注册与找回密码的切换面板。操作方式上，考生先选择认证模式，再输入手机号、邮箱、姓名等信息；注册时需同时完成图形验证码和邮箱验证码校验，系统实时反馈字段错误、验证码冷却时间及提交结果。",
    },
    {
        "key": "portal-home",
        "title": "学生门户首页原型图",
        "prototype_style": "portal",
        "section": "学生门户",
        "route": "/portal",
        "operator": "门户考生",
        "page_title": "博士生报名与服务门户",
        "summary": ["报名入口", "招生信息", "账号管理"],
        "filters": ["首页导航", "公告入口", "个人中心", "外链跳转"],
        "table_headers": ["服务卡片", "当前状态", "说明", "操作"],
        "table_rows": [
            ["在线申请", "未提交", "进入学生申报页", "立即填写"],
            ["账号管理", "正常", "修改密码与安全设置", "立即管理"],
            ["招生信息", "已开放", "查看实验室招生资讯", "立即查看"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["从顶部导航进入服务", "查看招生公告与日程", "进入在线申请或账号管理", "返回首页继续浏览信息"],
        "description": "该原型图对应学生门户首页，界面风格应明显区别于管理后台。页面采用门户式信息架构：顶部是品牌导航与账户入口，中部是招生宣传横幅和服务卡片，下部为报名入口、公告资讯和账号管理区。操作方式上，考生从首页顶部导航或中部卡片进入在线申请、查看招生信息或打开账号管理，不需要看到后台管理菜单、列表工作台或治理类抽屉。",
    },
    {
        "key": "portal-application",
        "title": "学生申报原型图",
        "prototype_style": "student-form",
        "section": "学生申报",
        "route": "/portal/application-v2",
        "operator": "门户考生",
        "page_title": "博士生在线申请",
        "summary": ["填写进度", "草稿状态", "附件完整度"],
        "filters": ["步骤导航", "自动保存", "附件上传", "提交确认"],
        "table_headers": ["申报分区", "当前状态", "关键内容", "操作"],
        "table_rows": [
            ["个人信息", "已完成", "身份信息、联系方式", "编辑"],
            ["教育经历", "填写中", "学历与培养背景", "继续填写"],
            ["成果与陈述", "未开始", "成果记录、个人陈述", "开始填写"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["沿步骤条逐段填写", "随时保存草稿", "上传附件并校验完整性", "最终提交前确认申报声明"],
        "description": "该原型图用于表达学生自助申报页面，不应采用后台列表管理模式，而应采用面向单个申请人的引导式长表单布局。页面顶部显示步骤进度和草稿状态，中部按个人信息、教育经历、实践经历、成果与陈述等分区组织内容，右侧显示填写进度、附件完整性和提交前检查。操作方式上，学生按步骤条依次填写各分区内容，可随时保存草稿、上传附件、返回修改，确认所有内容完整后再执行最终提交。",
    },
    {
        "key": "recruitment",
        "title": "招生管理原型图",
        "prototype_style": "admin",
        "section": "招生管理",
        "route": "/recruitment",
        "operator": "平台管理员、评分人、面试官",
        "page_title": "招生工作台",
        "summary": ["当前计划", "待审核申请", "待面试安排"],
        "filters": ["计划切换", "申请状态", "研究方向", "导入/导出"],
        "table_headers": ["考生", "资格审核", "材料评分", "面试安排", "录取状态"],
        "table_rows": [
            ["张三", "已通过", "89.5", "第1组 09:00", "待决策"],
            ["李四", "待审核", "--", "未安排", "未开始"],
            ["王五", "已通过", "92.0", "第2组 14:00", "预录取"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["先选择招生计划", "筛选待处理申请", "执行审核与评分", "安排面试并形成决策"],
        "description": "该原型图对应招生工作台，强调高密度作业界面和阶段化推进能力。页面顶部展示当前招生计划、待审核数量和待安排面试数量；中部提供计划切换、状态筛选、研究方向过滤以及导入导出操作；下方主表格承载考生审核、评分、面试安排和录取状态。操作方式上，管理员先切换计划，再按状态或方向筛选申请，随后逐条或批量执行资格审核、材料评分和面试安排，最后汇总形成预录取或录取决策。",
    },
    {
        "key": "student-master",
        "title": "学生主数据原型图",
        "prototype_style": "admin",
        "section": "学生主数据",
        "route": "/students",
        "operator": "平台管理员、导师、HRBP、党群负责人",
        "page_title": "学生与团队管理",
        "summary": ["在籍学生", "门户注册账号", "待维护团队"],
        "filters": ["学生状态", "所属团队", "导师", "门户账号状态"],
        "table_headers": ["学生", "团队", "导师", "当前状态", "门户账号"],
        "table_rows": [
            ["张三", "智能感知团队", "刘亚", "在培", "启用"],
            ["李四", "认知模型团队", "周青", "待入学", "未开通"],
            ["王五", "具身智能团队", "刘亚", "毕业", "停用"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["先筛选学生范围", "打开抽屉编辑主档", "选择团队与导师", "同步调整门户账号状态"],
        "description": "该原型图聚焦学生主档与团队治理的一体化维护。页面上方展示学生总量、门户注册账号数量和待维护团队数量；中部过滤区支持按学生状态、所属团队、导师和门户账号状态联查；主表格中展示学生与团队、导师、状态、门户账号等关键信息。操作方式上，管理员先筛选目标学生，再通过抽屉或弹窗进入主档编辑，选择团队与导师等受控主数据，并可同步执行门户账号启用、停用、密码重置等操作。",
    },
    {
        "key": "team-management",
        "title": "团队管理原型图",
        "prototype_style": "admin",
        "section": "团队管理",
        "route": "/students?tab=teams",
        "operator": "平台管理员",
        "page_title": "团队主数据管理",
        "summary": ["团队总数", "负责人导师", "团队学生数"],
        "filters": ["团队状态", "负责人导师", "研究方向"],
        "table_headers": ["团队", "负责人导师", "导师人数", "学生人数", "状态"],
        "table_rows": [
            ["智能感知团队", "刘亚", "8", "32", "启用"],
            ["认知模型团队", "周青", "6", "21", "启用"],
            ["具身智能团队", "袁野", "5", "17", "筹建"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["筛选团队范围", "打开团队详情", "维护导师集合和研究方向", "核对学生归属统计"],
        "description": "该原型图用于说明团队主数据治理页面的结构。页面应以“统计卡片 + 筛选区 + 团队列表 + 详情抽屉”的形式组织信息，突出团队编码、负责人导师、团队导师集合和学生归属统计。操作方式上，管理员可先按团队状态、负责人导师和研究方向筛选目标团队，再进入详情抽屉维护团队成员、研究方向和团队说明，并结合学生人数统计核验团队归属是否合理。",
    },
    {
        "key": "training-management",
        "title": "培养管理原型图",
        "prototype_style": "admin",
        "section": "培养管理",
        "route": "/training",
        "operator": "导师、平台管理员",
        "page_title": "培养过程管理",
        "summary": ["培养方案", "科研报告", "外出研修"],
        "filters": ["学生", "状态", "导师", "周期"],
        "table_headers": ["学生", "培养方案", "科研报告", "外出研修", "最近动作"],
        "table_rows": [
            ["张三", "已确认", "待审阅", "无", "提交科研报告"],
            ["李四", "待确认", "未提交", "已备案", "维护培养方案"],
            ["王五", "已完成", "已通过", "已归来", "归档成果"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["切换培养页签", "筛选学生与周期", "打开详情维护记录", "提交或审阅流程动作"],
        "description": "该原型图表现培养过程管理页面的分区方式。页面上半部分使用页签切换培养方案、科研报告、外出研修和成果记录；中部过滤区提供按学生、状态、导师和周期的组合检索；下半部分使用列表或时间线展示各类培养对象的最近状态。操作方式上，导师先切换目标页签，再通过筛选锁定学生和周期，随后进入详情页或抽屉维护培养方案、提交科研报告、审批外出研修，并查看最近一次业务动作。",
    },
    {
        "key": "degree-management",
        "title": "学位管理原型图",
        "prototype_style": "admin",
        "section": "学位管理",
        "route": "/degree",
        "operator": "导师、学位秘书、平台管理员",
        "page_title": "学位流程管理",
        "summary": ["论文总数", "盲审处理中", "待答辩安排"],
        "filters": ["论文状态", "学生", "导师", "学位阶段"],
        "table_headers": ["学生", "论文状态", "盲审结果", "答辩安排", "授位状态"],
        "table_rows": [
            ["张三", "待盲审", "未开始", "未安排", "未开始"],
            ["李四", "盲审中", "2/3 返回", "未安排", "未开始"],
            ["王五", "答辩完成", "通过", "已完成", "待授位"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["按阶段筛选论文", "进入里程碑详情", "录入盲审和答辩信息", "推动授位结论"],
        "description": "该原型图体现学位管理的强流程特征。页面顶部通过统计卡片展示论文总量、盲审处理中数量和待答辩安排数量；主列表以论文状态、盲审结果、答辩安排和授位状态为主字段；右侧详情面板使用时间轴或里程碑显示当前所处节点。操作方式上，学位秘书或导师先按论文状态、学生、导师筛选目标论文，再进入详情录入盲审结果、安排答辩、确认授位状态，并根据节点前置条件推进下一个环节。",
    },
    {
        "key": "workflow-center",
        "title": "流程中心原型图",
        "prototype_style": "admin",
        "section": "流程中心",
        "route": "/workflow",
        "operator": "全审批角色",
        "page_title": "流程待办中心",
        "summary": ["待办任务", "已办任务", "即将超时"],
        "filters": ["任务状态", "业务类型", "办理人", "创建时间"],
        "table_headers": ["任务", "业务类型", "当前节点", "发起人", "截止时间"],
        "table_rows": [
            ["资格审核-张三", "招生流程", "资格审核", "招生管理员", "今天 18:00"],
            ["科研报告-李四", "科研报告流程", "导师审阅", "李四", "明天 12:00"],
            ["论文-王五", "学位流程", "盲审确认", "王五", "后天 09:00"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["筛选待办任务", "打开详情查看业务对象", "执行通过或退回", "查看历史轨迹"],
        "description": "该原型图用于说明统一流程待办中心。页面需要突出“任务列表 + 详情面板 + 动作区 + 历史轨迹”的闭环结构，使审批角色在同一界面内完成筛选、查看、办理和追溯。操作方式上，用户先按任务状态、业务类型、办理人和时间筛选任务，再打开详情查看关联业务对象，随后通过通过、驳回、转办等按钮执行流程动作，并在底部或侧栏查看历史轨迹。",
    },
    {
        "key": "system-governance",
        "title": "系统治理原型图",
        "prototype_style": "admin",
        "section": "系统治理",
        "route": "/system",
        "operator": "平台管理员",
        "page_title": "系统治理中心",
        "summary": ["系统用户", "角色权限", "异常日志"],
        "filters": ["模块切换", "角色", "状态", "关键词"],
        "table_headers": ["对象", "类型", "当前状态", "最近变更", "操作"],
        "table_rows": [
            ["admin", "系统用户", "启用", "今天 09:10", "编辑/停用"],
            ["platform_admin", "角色", "生效", "昨天 14:00", "授权"],
            ["招生字典", "字典", "生效", "今天 08:30", "编辑"],
        ],
        "panel_title": "操作说明",
        "panel_lines": ["先切换治理模块", "筛选角色或状态", "编辑对象并确认风险操作", "查看日志追踪结果"],
        "description": "该原型图表达系统治理中心的统一治理体验。页面顶部通过卡片展示系统用户数、角色权限数和异常日志数量；中部使用模块切换控制用户、角色、字典、日志、集成链路等治理对象；底部主表格展示当前对象状态和最近变更。操作方式上，管理员先选择治理模块，再按角色、状态或关键词筛选对象，随后进入编辑弹窗完成配置变更；对于停用、删除、授权等高风险操作，系统应给出二次确认和日志留痕反馈。",
    },
]


ROLE_ROWS = [
    ("平台管理员", "系统级配置与全链路治理", "管理端全模块治理、角色权限配置、审计与集成管理"),
    ("博士生", "查看个人信息与业务概览", "管理端仅具备基础查看权限；门户侧作为独立考生身份使用"),
    ("导师", "培养方案制定、科研报告审阅、流程处理", "培养、学位、学生信息与流程任务"),
    ("学位秘书", "学位流程复核、送审与归档管理", "学位管理、流程中心、部分系统治理查看"),
    ("评分人", "招生材料评审与推荐", "招生查看、流程查看"),
    ("面试官", "面试分组、评分与校算", "招生读写、流程查看"),
    ("中心HRBP", "实习状态确认与过程监督", "学生查看、培养查看"),
    ("党群负责人", "思政考核与资助资格审查", "学生查看、审计查看"),
    ("门户考生", "注册、登录、找回密码、填写并提交申请", "学生门户独立认证域"),
]


MODULE_ROWS = [
    ("门户与认证", "门户考生、平台管理员", "注册、登录、找回密码、门户首页、在线申请入口控制", "认证隔离、短信/邮件/验证码反馈、登录后跳转恢复"),
    ("招生管理", "平台管理员、评分人、面试官", "招生计划、报名申请、资格审核、评分推荐、面试安排、录取决策", "工作台信息密度、列表与批量操作、阶段状态可视化"),
    ("学生主数据", "平台管理员、导师、HRBP、党群负责人", "学生主档、导师关系、门户注册学生、学生状态", "主数据受控选择、状态切换、历史可追溯"),
    ("团队管理", "平台管理员", "团队编码、负责人导师、团队导师集合、学生归属统计", "组织结构表达、成员关系维护、统计看板"),
    ("培养管理", "导师、平台管理员", "培养方案、科研报告、外出研修、成果信息", "长表单拆分、字典化选项、审批前后状态联动"),
    ("学位管理", "导师、学位秘书、平台管理员", "论文主档、盲审、答辩、授位", "流程里程碑、材料状态、节点阻断规则"),
    ("流程中心", "全审批角色", "待办任务、任务详情、动作执行、流程轨迹", "任务聚合展示、动作按钮控制、历史轨迹查看"),
    ("系统治理", "平台管理员", "系统用户、角色、权限、字典、审计策略、集成链路、日志", "高风险操作确认、角色权限矩阵、配置变更影响提示"),
]


BUSINESS_FLOW_ROWS = [
    ("门户注册", "考生在门户发起注册", "邮箱验证码校验、图形验证码校验、账号唯一性校验", "生成门户账号", "需设计注册页、验证码交互、错误提示与成功反馈"),
    ("在线申请", "门户考生进入在线申请", "按结构化分段填写资料、保存草稿、上传附件、最终提交", "生成报名申请记录并进入招生域", "需设计多分段表单、进度提示、草稿保存与提交确认"),
    ("招生评审", "管理员完成报名整理后", "资格审核、材料评分、面试安排、面试评分、录取决策", "形成录取结果", "需设计工作台列表、状态切换、批量导入导出与评分录入"),
    ("学生建档", "录取结果确认后", "学生主档、导师关系、团队归属初始化", "形成学生全周期主数据", "需设计主档编辑页、团队与导师联动选择"),
    ("培养执行", "学生在培期间", "培养方案维护、科研报告提交、外出研修申请与审批", "沉淀培养过程记录", "需设计分模块页签、时间线与状态提示"),
    ("学位申请", "学生满足学位阶段条件", "论文维护、盲审、预答辩、答辩、授位", "形成毕业/授位结论", "需设计里程碑链路、节点门禁与文档状态呈现"),
    ("系统治理", "管理员开展平台运维", "用户、角色、字典、日志、集成链路管理", "平台配置持续生效", "需设计治理台账、权限矩阵与审计追踪页"),
]


APPROVAL_FLOW_ROWS = [
    ("招生流程", "资格审核、评分推荐、面试安排、录取决策", "管理员、评分人、面试官", "流程中心、招生工作台", "阶段式流转，需展示当前环节、上一环节结果和下一动作"),
    ("科研报告流程", "提交、审阅、通过/退回", "学生、导师、管理员", "培养管理、流程中心", "需支持报告详情查看、评语记录和退回重提"),
    ("外出研修流程", "申请、导师审核、管理员备案", "学生、导师、管理员", "培养管理、流程中心", "需支持审批意见、出访信息、归来评估联动"),
    ("学位流程", "论文提交、盲审、预答辩、正式答辩、授位", "学生、导师、学位秘书、管理员", "学位管理、流程中心", "需支持严格前置条件控制和关键节点留痕"),
    ("系统治理操作", "用户/角色/字典等高风险变更", "平台管理员", "系统治理模块", "虽非传统审批，但需保留确认、校验、日志记录与结果反馈"),
]


ENTITY_ROWS = [
    ("系统用户", "管理端登录主体", "角色、登录日志、操作日志", "登录页、用户管理页、权限注入逻辑"),
    ("角色", "权限集合载体", "系统用户、权限目录", "角色维护页、菜单可见性、接口鉴权"),
    ("门户账号", "门户登录主体", "门户资料、在线申请", "门户认证页、账号状态管理页"),
    ("招生计划", "招生批次与配置载体", "报名申请、面试分组、招生简章素材", "招生工作台、简章展示"),
    ("招生申请", "考生正式报名业务对象", "资格审核、评分、面试、录取决策", "门户申请页、招生评审页、流程详情页"),
    ("学生", "博士生主档核心对象", "导师、团队、培养、学位", "学生主档页、学生画像、生命周期轨迹"),
    ("导师", "指导关系核心对象", "学生、团队、培养、学位", "导师关系维护、培养审批、团队管理"),
    ("团队", "研究组织单元", "学生、导师", "团队管理页、学生编辑页、统计页"),
    ("培养方案", "培养阶段版本化对象", "学生、导师、流程任务", "培养方案编辑页、版本对比"),
    ("科研报告", "周期性培养成果对象", "学生、导师、流程任务", "报告详情页、审阅页"),
    ("外出研修", "出访审批与归来评估对象", "学生、导师、流程任务", "申请页、审批页、台账页"),
    ("论文", "学位主线对象", "学生、盲审意见、答辩/授位阶段", "学位详情页、里程碑页"),
    ("流程任务", "统一审批载体", "业务键、审批人、处理动作", "流程中心、详情抽屉/弹窗"),
    ("字典项", "受控选项来源", "学生、培养、系统治理等多模块字段", "下拉框、状态标签、颜色呈现"),
]


ENTITY_MODEL_SPECS = [
    {
        "entity": "学生",
        "description": "学生是博士生生命周期的核心主实体，用于承载身份、团队、导师、状态和培养学位主线。",
        "rows": [
            ("student_id", "文本", "学生主实体唯一标识。", "是"),
            ("student_no", "文本", "学号，用于校内唯一识别学生。", "是"),
            ("full_name", "文本", "学生姓名，用于页面显示和检索。", "否"),
            ("gender", "枚举字典", "性别，来源于统一字典项。", "否"),
            ("current_status", "枚举字典", "当前生命周期状态，如待入学、在培、毕业。", "否"),
            ("team_id", "文本", "所属团队标识，关联团队实体。", "否"),
            ("primary_advisor_id", "文本", "主导师标识，关联导师实体。", "否"),
            ("phone", "文本", "联系电话，用于日常联络。", "否"),
            ("email", "文本", "邮箱地址，用于通知与校验。", "否"),
        ],
    },
    {
        "entity": "导师",
        "description": "导师实体用于维护指导教师主数据，并与学生、团队、培养和学位业务关联。",
        "rows": [
            ("advisor_id", "文本", "导师主实体唯一标识。", "是"),
            ("advisor_no", "文本", "导师编号，用于唯一识别导师。", "是"),
            ("full_name", "文本", "导师姓名。", "否"),
            ("title", "枚举字典", "职称信息，如教授、副教授。", "否"),
            ("research_direction", "文本", "主要研究方向描述。", "否"),
            ("annual_quota", "数字", "年度招生名额。", "否"),
            ("status", "枚举字典", "在岗、停用等状态。", "否"),
        ],
    },
    {
        "entity": "团队",
        "description": "团队实体用于描述研究组织单元，并维护负责人导师、团队成员导师与学生归属。",
        "rows": [
            ("team_id", "文本", "团队主实体唯一标识。", "是"),
            ("team_code", "文本", "团队编码，用于唯一识别团队。", "是"),
            ("team_name", "文本", "团队名称。", "否"),
            ("lead_advisor_id", "文本", "负责人导师标识，关联导师实体。", "否"),
            ("research_area", "文本", "团队研究方向。", "否"),
            ("team_status", "枚举字典", "团队状态，如启用、筹建。", "否"),
            ("student_count", "数字", "团队当前学生数，用于统计展示。", "否"),
        ],
    },
    {
        "entity": "门户账号",
        "description": "门户账号是考生与学生门户的登录主体，用于注册、登录、找回密码和在线申请。",
        "rows": [
            ("portal_user_id", "文本", "门户账号唯一标识。", "是"),
            ("phone", "文本", "手机号，用于登录或联系。", "是"),
            ("email", "文本", "邮箱地址，用于验证码和通知。", "是"),
            ("full_name", "文本", "考生姓名。", "否"),
            ("id_number", "文本", "身份证号，用于实名识别。", "是"),
            ("account_status", "枚举字典", "账号状态，如启用、停用。", "否"),
            ("last_login_at", "日期时间", "最近一次登录时间。", "否"),
        ],
    },
    {
        "entity": "招生计划",
        "description": "招生计划实体描述某一批次招生工作的基础配置，是报名申请和面试安排的上游对象。",
        "rows": [
            ("plan_id", "文本", "招生计划唯一标识。", "是"),
            ("plan_code", "文本", "招生计划编码。", "是"),
            ("academic_year", "文本", "所属学年。", "否"),
            ("plan_name", "文本", "计划名称。", "否"),
            ("plan_status", "枚举字典", "计划状态，如草稿、开放、结束。", "否"),
            ("brochure_url", "文本", "招生简章素材地址。", "否"),
            ("start_date", "日期", "计划开始日期。", "否"),
            ("end_date", "日期", "计划结束日期。", "否"),
        ],
    },
    {
        "entity": "招生申请",
        "description": "招生申请实体用于承载考生正式提交的报名信息，并贯穿资格审核、评分、面试和录取。",
        "rows": [
            ("application_id", "文本", "招生申请唯一标识。", "是"),
            ("plan_id", "文本", "所属招生计划标识。", "否"),
            ("portal_user_id", "文本", "提交申请的门户账号标识。", "否"),
            ("candidate_no", "文本", "考生编号。", "是"),
            ("application_status", "枚举字典", "申请状态，如待审核、面试中、预录取。", "否"),
            ("final_score", "数字", "综合成绩。", "否"),
            ("business_key", "文本", "流程关联业务键。", "是"),
        ],
    },
    {
        "entity": "培养方案",
        "description": "培养方案实体记录学生在培养阶段的方案内容、版本与确认状态。",
        "rows": [
            ("training_plan_id", "文本", "培养方案唯一标识。", "是"),
            ("student_id", "文本", "关联学生标识。", "否"),
            ("advisor_id", "文本", "关联导师标识。", "否"),
            ("version_no", "数字", "当前版本号。", "否"),
            ("plan_status", "枚举字典", "方案状态，如待确认、已确认。", "否"),
            ("plan_snapshot", "文本", "方案快照或摘要信息。", "否"),
            ("effective_date", "日期", "生效日期。", "否"),
        ],
    },
    {
        "entity": "科研报告",
        "description": "科研报告实体记录学生周期性科研汇报，是培养管理和流程审批的重要对象。",
        "rows": [
            ("report_id", "文本", "科研报告唯一标识。", "是"),
            ("student_id", "文本", "关联学生标识。", "否"),
            ("period_label", "文本", "报告周期，如 2026 春季。", "否"),
            ("report_status", "枚举字典", "报告状态，如待提交、待审阅、已通过。", "否"),
            ("review_score", "数字", "审阅分数。", "否"),
            ("business_key", "文本", "流程关联业务键。", "是"),
        ],
    },
    {
        "entity": "外出研修",
        "description": "外出研修实体用于记录学生外出学习、审批备案和归来评估过程。",
        "rows": [
            ("outbound_id", "文本", "外出研修唯一标识。", "是"),
            ("student_id", "文本", "关联学生标识。", "否"),
            ("study_type", "枚举字典", "研修类型。", "否"),
            ("destination", "文本", "目的地或机构名称。", "否"),
            ("approval_status", "枚举字典", "审批状态，如待审核、已备案。", "否"),
            ("return_assessment", "文本", "归来评估说明。", "否"),
            ("business_key", "文本", "流程关联业务键。", "是"),
        ],
    },
    {
        "entity": "论文",
        "description": "论文实体是学位流程主对象，用于承载盲审、答辩和授位相关状态。",
        "rows": [
            ("thesis_id", "文本", "论文唯一标识。", "是"),
            ("student_id", "文本", "关联学生标识。", "否"),
            ("thesis_title", "文本", "论文题目。", "否"),
            ("thesis_status", "枚举字典", "论文流程状态。", "否"),
            ("plagiarism_rate", "数字", "查重比例。", "否"),
            ("degree_granted", "布尔", "是否已授位。", "否"),
            ("business_key", "文本", "流程关联业务键。", "是"),
        ],
    },
    {
        "entity": "流程任务",
        "description": "流程任务实体统一承载待办、已办、动作执行与流程轨迹信息。",
        "rows": [
            ("task_id", "文本", "流程任务唯一标识。", "是"),
            ("business_key", "文本", "关联业务对象键。", "否"),
            ("task_name", "文本", "任务名称。", "否"),
            ("task_status", "枚举字典", "任务状态，如待办、已办、退回。", "否"),
            ("assignee", "文本", "当前办理人。", "否"),
            ("due_at", "日期时间", "截止办理时间。", "否"),
        ],
    },
]


OPERATION_MODE_ROWS = [
    ("管理端标准操作模式", "已登录内部角色", "菜单导航 + 列表检索 + 表单编辑 + 操作弹窗", "适用于招生、学生、培养、学位、系统治理等管理页面"),
    ("门户自助操作模式", "门户考生", "引导式页面 + 长表单分段 + 对话框反馈", "适用于注册、登录、密码找回、在线申请"),
    ("流程处理模式", "审批角色", "待办列表 + 详情面板 + 动作按钮 + 历史轨迹", "适用于跨业务统一审批中心"),
    ("系统配置模式", "平台管理员", "表格治理 + 配置编辑 + 高风险确认", "适用于角色、字典、审计策略、集成链路"),
    ("统一托管部署模式", "运维人员", "FastAPI 托管 API 与前端资源", "适用于一体化部署环境"),
    ("前后端分离部署模式", "运维人员", "Nginx 托管前端，FastAPI 提供 API", "适用于静态资源与 API 分域部署环境"),
]


PAGE_ROWS = [
    ("/login", "管理端登录", "用户名/密码输入、登录失败对话框、登录成功跳转", "需提供清晰登录状态反馈与错误提示策略"),
    ("/portal", "门户首页", "招生信息入口、在线申请入口、账号管理、外链跳转", "需体现招生品牌感与重要入口分层"),
    ("/portal/auth", "门户认证", "注册、登录、找回密码、邮箱验证码、图形验证码", "需设计双验证码流程、切换式认证面板"),
    ("/portal/application-v2", "在线申请", "多分段资料填写、草稿保存、附件上传、最终提交", "需提供步骤导航、保存状态、提交确认与字段校验反馈"),
    ("/recruitment", "招生工作台", "计划切换、申请列表、导入导出、评分与面试安排", "需突出阶段切换和高密度工作台操作"),
    ("/students", "学生与团队管理", "学生主档维护、门户账号状态管理、团队维护", "需处理主从关系、批量操作与状态标签"),
    ("/training", "培养管理", "培养方案、科研报告、外出研修、成果维护", "需采用分区布局或页签降低信息拥挤"),
    ("/degree", "学位管理", "论文、盲审、答辩与授位阶段管理", "需采用里程碑或时间轴强化流程感"),
    ("/workflow", "流程中心", "待办、已办、详情、动作执行、轨迹查看", "需强调任务优先级、状态与操作闭环"),
    ("/system / /dict", "系统治理", "用户、角色、字典、日志、集成链路、审计策略", "需强调风险提示、矩阵关系与操作可追溯"),
]


DESIGN_INPUT_ROWS = [
    ("技术负责人", "模块边界、核心业务对象、审批流转、运行模式、集成点、非功能约束", "用于详细设计、接口拆分、数据库与服务边界设计"),
    ("UI/UE 设计师", "页面清单、角色差异、关键任务流、交互反馈、状态与异常场景", "用于信息架构、页面原型、交互规格与视觉层次设计"),
    ("前端工程师", "页面路由、组件分区、表单分段、鉴权与异常反馈模式", "用于页面拆分、状态管理与组件抽象"),
    ("后端工程师", "领域对象、状态机、鉴权规则、配置项、日志与集成约束", "用于 API 边界、服务实现、数据库与审计设计"),
    ("测试工程师", "业务流程、审批路径、异常分支、权限矩阵、部署模式", "用于测试场景、用例分层与回归范围定义"),
]


REQUIREMENT_SECTIONS = [
    (
        "3.1 平台级与权限需求",
        [
            ("普适型", "系统应提供管理端与学生门户两套访问入口，并对两套入口的认证状态进行隔离。"),
            ("普适型", "系统应基于角色向已登录用户分配权限集合，并以该权限集作为页面展示和接口访问的共同依据。"),
            ("事件驱动型", "当用户访问管理端受保护路由时，系统应先校验登录状态，再决定放行或重定向到登录页。"),
            ("状态驱动型", "在用户权限集合不包含某菜单要求权限期间，系统应隐藏对应的管理端导航入口。"),
            ("异常型", "如果用户请求的管理端接口缺少必需权限，系统应拒绝请求并返回明确的缺失权限提示。"),
            ("可选特性型", "在角色维护场景中，系统应只允许选择权限目录中已登记的权限编码。"),
        ],
    ),
    (
        "3.2 管理端认证与会话需求",
        [
            ("事件驱动型", "当系统用户提交用户名和密码时，系统应校验账号是否存在、是否启用以及密码是否正确。"),
            ("事件驱动型", "当系统用户登录成功时，系统应签发访问令牌和刷新令牌。"),
            ("普适型", "系统应将管理端登录会话持久化，以支持令牌失效控制和服务重启后的会话校验。"),
            ("异常型", "如果访问令牌无效、过期或对应会话已失效，系统应拒绝请求并返回未认证提示。"),
            ("事件驱动型", "当用户主动注销时，系统应撤销当前会话，并使旧令牌立即失效。"),
            ("事件驱动型", "当登录成功且存在先前记录的目标页面时，系统应优先恢复用户原本准备访问的页面。"),
        ],
    ),
    (
        "3.3 招生管理需求",
        [
            ("普适型", "系统应支持维护招生计划，并保存学期、摘要、招生简章图片等基础信息。"),
            ("普适型", "系统应支持对报名申请进行查询、筛选、新增、编辑和删除。"),
            ("事件驱动型", "当招生计划被选中时，系统应加载该计划下的申请、审核、面试与录取相关数据。"),
            ("事件驱动型", "当管理员导入报名模板数据时，系统应返回导入成功数量、跳过数量和问题摘要。"),
            ("事件驱动型", "当用户对招生申请执行流程动作时，系统应推进对应流程状态。"),
            ("异常型", "如果用户试图删除仍被业务数据引用的招生对象，系统应阻止删除并返回原因。"),
        ],
    ),
    (
        "3.4 学生主数据与团队需求",
        [
            ("普适型", "系统应维护学生主档，包括身份、状态、导师、团队和联系方式等核心信息。"),
            ("普适型", "系统应维护研究中心/团队主数据，并记录负责人导师、导师集合和学生归属统计。"),
            ("事件驱动型", "当用户新增或编辑学生主档时，系统应要求其从受控主数据中选择角色、团队和相关对象，而不是自由录入。"),
            ("事件驱动型", "当管理员查看门户注册学生列表时，系统应展示门户注册账号状态并允许执行启用、停用和密码重置操作。"),
            ("状态驱动型", "在门户注册账号处于停用状态期间，系统应阻止该账号继续登录或提交报名。"),
            ("异常型", "如果学生与所选团队、导师关系不满足约束，系统应拒绝保存并返回错误信息。"),
        ],
    ),
    (
        "3.5 培养过程需求",
        [
            ("普适型", "系统应支持按学生维度维护培养方案、科研报告和外出研修业务数据。"),
            ("普适型", "系统应提供培养过程查询与过滤能力，并支持按状态、导师、关键字等条件检索。"),
            ("事件驱动型", "当用户保存培养方案时，系统应持久化当前版本并返回保存结果。"),
            ("事件驱动型", "当用户提交科研报告时，系统应为该报告创建或同步对应的流程任务。"),
            ("状态驱动型", "在培养方案达到年度修改上限期间，系统应阻止继续修改。"),
            ("异常型", "如果培养数据缺少关键字段或状态不满足前置条件，系统应拒绝提交并返回原因。"),
        ],
    ),
    (
        "3.6 学位管理需求",
        [
            ("普适型", "系统应维护论文主档、盲审意见、答辩与授位相关状态。"),
            ("事件驱动型", "当用户新增或更新论文主档时，系统应同步维护论文对应的流程状态和运行态镜像。"),
            ("事件驱动型", "当用户录入盲审意见时，系统应保存专家、评分与评审状态。"),
            ("状态驱动型", "在论文未通过查重、盲审或流程前置节点期间，系统应阻止进入后续学位环节。"),
            ("异常型", "如果学位业务对象不存在或状态非法，系统应拒绝修改并返回明确错误。"),
        ],
    ),
    (
        "3.7 工作流与审批需求",
        [
            ("普适型", "系统应提供统一的流程待办中心，用于展示审批任务、业务对象和当前处理节点。"),
            ("事件驱动型", "当用户打开某审批任务时，系统应返回任务详情、可执行动作和历史轨迹。"),
            ("事件驱动型", "当用户执行通过、驳回等动作时，系统应更新流程任务、业务对象状态和历史记录。"),
            ("普适型", "系统应通过统一的业务键将招生、培养、学位等业务对象与流程实例关联。"),
            ("异常型", "如果流程动作与当前节点不匹配或业务对象不存在，系统应拒绝执行该动作。"),
        ],
    ),
    (
        "3.8 系统治理与审计需求",
        [
            ("普适型", "系统应提供系统用户、角色、权限、字典、审计策略、集成链路、操作日志和同步日志的治理能力。"),
            ("事件驱动型", "当管理员新建或修改角色时，系统应校验角色编码唯一性并去重权限编码。"),
            ("事件驱动型", "当角色编码或角色名称发生变化时，系统应同步更新已关联用户和个人资料中的角色显示信息。"),
            ("普适型", "系统应将所有关键写操作记录到操作日志，以支持事后追溯。"),
            ("异常型", "如果角色仍被系统用户引用，系统应阻止删除该角色。"),
        ],
    ),
    (
        "3.9 学生门户需求",
        [
            ("普适型", "系统应为考生提供独立的门户注册、登录、找回密码、门户首页和在线申请能力。"),
            ("事件驱动型", "当考生注册时，系统应校验手机号、邮箱、姓名、身份证号、密码和邮件验证码。"),
            ("事件驱动型", "当考生请求邮件验证码时，系统应发送验证码并返回有效期与冷却时间。"),
            ("事件驱动型", "当考生登录成功时，系统应签发独立的门户访问令牌，并允许其进入受保护的门户页面。"),
            ("普适型", "系统应允许考生以结构化方式填写个人信息、教育经历、实践经历、英语能力、家庭成员、成果记录和个人陈述。"),
            ("事件驱动型", "当考生保存申请草稿时，系统应持久化当前草稿，并允许考生再次打开后继续填写。"),
            ("事件驱动型", "当考生提交完整申请时，系统应生成报名业务记录并启动后续招生流程。"),
            ("状态驱动型", "在在线申请入口关闭期间，系统应阻止考生进入在线申请页面与调用提交接口，并向考生展示开启提示。"),
            ("可选特性型", "在配置了招生信息外链地址的地方，系统应允许门户首页导航和横幅跳转到该地址。"),
        ],
    ),
    (
        "3.10 通知、附件与集成需求",
        [
            ("普适型", "系统应支持通过邮件发送注册验证码、密码重置和相关通知。"),
            ("事件驱动型", "当考生上传附件时，系统应校验附件分类、扩展名、内容类型和大小限制。"),
            ("异常型", "如果附件格式、类型或大小不符合要求，系统应拒绝上传并返回明确错误信息。"),
            ("普适型", "系统应保留对外部系统的集成链路配置，以支持与 OA、飞书等外部平台的协同。"),
            ("事件驱动型", "当外部集成执行或失败时，系统应记录同步结果并允许查询。"),
        ],
    ),
    (
        "4.1 安全与访问控制非功能需求",
        [
            ("普适型", "系统应对管理端和学生门户分别实施认证与授权控制，避免令牌混用。"),
            ("普适型", "系统应对受保护接口统一执行权限检查，不得仅依赖前端隐藏按钮实现访问控制。"),
            ("事件驱动型", "当用户修改密码时，系统应校验当前密码，并仅在校验通过后更新新密码。"),
            ("异常型", "如果账号被停用或不存在，系统应拒绝认证请求。"),
            ("普适型", "系统应支持通过配置文件管理数据库、缓存、令牌、邮件和门户开关等敏感参数。"),
        ],
    ),
    (
        "4.2 性能与可用性非功能需求",
        [
            ("普适型", "系统应提供健康检查接口，以便部署环境进行存活检测。"),
            ("事件驱动型", "当门户考生提交申请或保存草稿时，系统应在一次明确的请求响应中返回成功或失败结果，不得出现无提示中断。"),
            ("普适型", "系统应支持单机与高可用两种缓存部署模式，以适配不同可用性级别的部署环境。"),
            ("状态驱动型", "在前端统一托管模式期间，系统应允许应用同时提供 API 与已构建前端页面。"),
            ("状态驱动型", "在前后端分离部署模式期间，系统应允许 Web 服务器直接托管前端静态文件，而应用服务仅提供 API 和健康检查。"),
        ],
    ),
    (
        "4.3 审计、运维与部署非功能需求",
        [
            ("普适型", "系统应为登录、关键写操作和外部同步提供可查询的审计日志。"),
            ("普适型", "系统应支持通过标准进程管理工具托管应用服务。"),
            ("普适型", "系统应支持通过 Web 服务器提供 HTTPS、反向代理和静态资源访问。"),
            ("事件驱动型", "当部署采用前后端分离模式时，系统应允许通过统一 API 路径将浏览器请求转发到后端服务。"),
            ("异常型", "如果静态资源请求被错误回退为 HTML，部署方案应允许通过精确路径配置来避免脚本资源类型异常。"),
        ],
    ),
]


def apply_run_style(run, size: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_FALLBACK)


def configure_paragraph(paragraph, alignment=None, first_line_indent_cm: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.2
    fmt.space_after = Pt(6)
    fmt.space_before = Pt(0)
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)
    if alignment is not None:
        paragraph.alignment = alignment


def add_text_paragraph(document: Document, text: str, *, bold: bool = False, size: float = 12, alignment=None, indent: float | None = 0.74):
    paragraph = document.add_paragraph()
    configure_paragraph(paragraph, alignment=alignment, first_line_indent_cm=indent)
    run = paragraph.add_run(text)
    apply_run_style(run, size=size, bold=bold)
    return paragraph


def normalize_heading_text(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", text).strip()


def build_xml_element(tag: str, **attrs: str) -> OxmlElement:
    element = OxmlElement(tag)
    for key, value in attrs.items():
        element.set(qn(key), value)
    return element


def insert_toc(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-5" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "右键更新域后生成目录"
    placeholder_run.append(placeholder_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_separate)
    paragraph._p.append(placeholder_run)
    paragraph._p.append(fld_end)


def ensure_heading_numbering(document: Document) -> int:
    existing_num_id = getattr(document, "_auto_heading_num_id", None)
    if existing_num_id is not None:
        return existing_num_id

    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]

    abstract_num_id = str(max(abstract_ids, default=0) + 1)
    num_id = str(max(num_ids, default=0) + 1)

    abstract_num = build_xml_element("w:abstractNum", **{"w:abstractNumId": abstract_num_id})
    abstract_num.append(build_xml_element("w:multiLevelType", **{"w:val": "multilevel"}))

    left_values = [0, 720, 1440, 2160, 2880]
    for level in range(5):
        lvl = build_xml_element("w:lvl", **{"w:ilvl": str(level)})
        lvl.append(build_xml_element("w:start", **{"w:val": "1"}))
        lvl.append(build_xml_element("w:numFmt", **{"w:val": "decimal"}))
        pattern = ".".join(f"%{index}" for index in range(1, level + 2)) + "."
        lvl.append(build_xml_element("w:lvlText", **{"w:val": pattern}))
        lvl.append(build_xml_element("w:lvlJc", **{"w:val": "left"}))

        p_pr = build_xml_element("w:pPr")
        p_pr.append(
            build_xml_element(
                "w:ind",
                **{
                    "w:left": str(left_values[level]),
                    "w:hanging": "360",
                },
            )
        )
        lvl.append(p_pr)

        r_pr = build_xml_element("w:rPr")
        r_pr.append(build_xml_element("w:b", **{"w:val": "1"}))
        lvl.append(r_pr)
        abstract_num.append(lvl)

    numbering.append(abstract_num)

    num = build_xml_element("w:num", **{"w:numId": num_id})
    num.append(build_xml_element("w:abstractNumId", **{"w:val": abstract_num_id}))
    numbering.append(num)

    document._auto_heading_num_id = int(num_id)
    return document._auto_heading_num_id


def apply_heading_numbering(paragraph, level: int, num_id: int) -> None:
    paragraph.style = f"Heading {max(1, min(level, 5))}"
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(max(0, min(level - 1, 4))))

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), str(num_id))


def configure_heading_styles(document: Document) -> None:
    size_map = {1: 16, 2: 14, 3: 12, 4: 12, 5: 12}
    left_indent_map = {1: 0, 2: 0.74, 3: 1.48, 4: 2.22, 5: 2.96}
    for level in range(1, 6):
        style = document.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_FALLBACK)
        style.font.size = Pt(size_map[level])
        style.font.bold = True
        fmt = style.paragraph_format
        fmt.line_spacing = 1.2
        fmt.space_before = Pt(6 if level == 1 else 3)
        fmt.space_after = Pt(6)
        fmt.first_line_indent = Cm(0)
        fmt.left_indent = Cm(left_indent_map[level])

    title_style = document.styles["Title"]
    title_style.font.name = BODY_FONT
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_FALLBACK)
    title_style.font.size = Pt(18)
    title_style.font.bold = True


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    configure_paragraph(paragraph, first_line_indent_cm=0)
    apply_heading_numbering(paragraph, level, ensure_heading_numbering(document))
    run = paragraph.add_run(normalize_heading_text(text))
    apply_run_style(run, size={1: 16, 2: 14, 3: 12, 4: 12, 5: 12}.get(level, 12), bold=True)


def set_table_cell_text(cell, text: str, *, bold: bool = False, size: float = 11) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    configure_paragraph(paragraph, first_line_indent_cm=0)
    run = paragraph.add_run(text)
    apply_run_style(run, size=size, bold=bold)


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_table_cell_text(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_table_cell_text(cells[index], value)


def add_entity_model_tables(document: Document) -> None:
    for spec in ENTITY_MODEL_SPECS:
        add_heading(document, f"{spec['entity']}属性模型", 3)
        add_text_paragraph(document, str(spec['description']))
        add_table(
            document,
            ["属性（字段）", "属性类型", "描述", "是否为唯一关键属性"],
            list(spec['rows']),
        )


def ensure_asset_dirs() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    EARS_IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def add_figure(document: Document, png_path: Path, caption: str, width_cm: float = 16.0) -> None:
    document.add_picture(str(png_path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph()
    configure_paragraph(caption_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = caption_paragraph.add_run(caption)
    apply_run_style(run, size=10)


def add_operation_list(document: Document, title: str, items: list[str]) -> None:
    add_text_paragraph(document, title, bold=True, indent=0)
    for index, item in enumerate(items, start=1):
        add_text_paragraph(document, f"{index}. {item}", indent=0.74)


def render_business_flow_figure(svg_path: Path) -> None:
    dwg = create_drawing(svg_path, 1800, 720, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期业务主流程图', insert=(900, 50), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#17304f'))
    stages = [
        ('门户注册', '注册账号\n完成双验证码校验'),
        ('在线申请', '填写资料\n保存草稿并提交'),
        ('招生评审', '资格审核\n评分与面试安排'),
        ('录取建档', '形成录取结果\n建立学生主档'),
        ('培养执行', '培养方案\n科研报告与研修'),
        ('学位申请', '论文、盲审\n答辩与授位'),
        ('毕业归档', '离校归档\n就业跟踪'),
    ]
    positions = [70, 310, 550, 790, 1030, 1270, 1510]
    for x, (title, subtitle) in zip(positions, stages):
        add_box(dwg, x, 220, 210, 110, title, subtitle, '#eef6ff')
    for index in range(len(positions) - 1):
        add_arrow(dwg, (positions[index] + 210, 275), (positions[index + 1], 275), '状态推进')
    add_box(dwg, 360, 420, 300, 100, '协同支撑', '通知提醒\n附件上传\n操作留痕', '#fff5e8')
    add_box(dwg, 760, 420, 300, 100, '统一流程中心', '任务详情\n动作执行\n历史轨迹', '#eefcf2')
    add_box(dwg, 1160, 420, 300, 100, '系统治理', '角色权限\n字典配置\n日志审计', '#f4f0ff')
    add_arrow(dwg, (520, 330), (510, 420), '支撑')
    add_arrow(dwg, (920, 330), (910, 420), '驱动')
    add_arrow(dwg, (1320, 330), (1310, 420), '治理')
    dwg.save()


def render_approval_flow_figure(svg_path: Path) -> None:
    dwg = create_drawing(svg_path, 1780, 980, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期审批流程图', insert=(890, 46), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#17304f'))
    columns = [
        (80, '招生流程', ['资格审核', '评分推荐', '面试安排', '录取决策']),
        (500, '科研报告流程', ['学生提交', '导师审阅', '通过或退回', '归档']),
        (920, '外出研修流程', ['学生申请', '导师审核', '管理员备案', '归来评估']),
        (1340, '学位流程', ['论文提交', '盲审', '答辩安排', '授位结论']),
    ]
    for x, title, steps in columns:
        add_box(dwg, x, 110, 320, 74, title, '', '#eaf3ff')
        current_y = 230
        for step in steps:
            add_box(dwg, x + 15, current_y, 290, 88, step, '查看详情 / 执行动作 / 记录轨迹', '#eefcf2' if current_y % 2 == 0 else '#fff7eb')
            if current_y < 230 + (len(steps) - 1) * 150:
                add_arrow(dwg, (x + 160, current_y + 88), (x + 160, current_y + 140), '')
            current_y += 150
    add_box(dwg, 630, 830, 520, 90, '流程中心操作方式', '审批人先筛选任务，再查看业务详情，随后执行通过、驳回或转办动作，系统同步记录历史轨迹。', '#f4f0ff')
    dwg.save()


def render_entity_relationship_figure(svg_path: Path) -> None:
    dwg = create_drawing(svg_path, 1800, 1100, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期核心实体ER关系图', insert=(900, 44), text_anchor='middle', font_size=30, font_family=SVG_FONT, font_weight='bold', fill='#17304f'))
    add_box(dwg, 760, 120, 320, 132, '学生', 'PK：student_id\nUK：student_no\nFK：team_id、primary_advisor_id', '#eef6ff')
    add_box(dwg, 320, 130, 260, 120, '导师', 'PK：advisor_id\nUK：advisor_no\n属性：annual_quota', '#fff7e8')
    add_box(dwg, 1210, 130, 260, 120, '团队', 'PK：team_id\nUK：team_code\nFK：lead_advisor_id', '#eefcf2')
    add_box(dwg, 120, 370, 300, 132, '门户账号', 'PK：portal_user_id\nUK：phone、email、id_number\n属性：account_status', '#fff4ea')
    add_box(dwg, 470, 355, 300, 148, '招生申请', 'PK：application_id\nUK：candidate_no、business_key\nFK：plan_id、portal_user_id', '#fff4ea')
    add_box(dwg, 835, 355, 300, 148, '培养方案', 'PK：training_plan_id\nFK：student_id、advisor_id\n属性：version_no、plan_status', '#eefcf2')
    add_box(dwg, 1200, 355, 300, 148, '科研报告', 'PK：report_id\nUK：business_key\nFK：student_id', '#eefcf2')
    add_box(dwg, 470, 610, 300, 148, '外出研修', 'PK：outbound_id\nUK：business_key\nFK：student_id', '#eefcf2')
    add_box(dwg, 835, 610, 300, 148, '论文', 'PK：thesis_id\nUK：business_key\nFK：student_id', '#eefcf2')
    add_box(dwg, 1200, 610, 300, 148, '流程任务', 'PK：task_id\nFK：business_key\n属性：task_status、assignee', '#f4f0ff')
    add_box(dwg, 760, 860, 320, 120, '系统治理', '系统用户、角色、字典、日志\n为主业务实体提供权限和规则约束', '#f4f0ff')
    add_arrow(dwg, (600, 198), (760, 198), '指导')
    add_arrow(dwg, (1080, 188), (1210, 188), '归属')
    add_arrow(dwg, (420, 430), (760, 220), '注册/转化')
    add_arrow(dwg, (770, 430), (810, 430), '1:1/1:N')
    add_arrow(dwg, (1080, 430), (1200, 430), '1:N')
    add_arrow(dwg, (620, 252), (620, 355), '申请/培养')
    add_arrow(dwg, (920, 252), (920, 355), '培养方案')
    add_arrow(dwg, (1040, 252), (1330, 355), '科研报告')
    add_arrow(dwg, (620, 503), (620, 610), '外出研修')
    add_arrow(dwg, (985, 503), (985, 610), '进入学位阶段')
    add_arrow(dwg, (1135, 684), (1200, 684), '生成待办')
    add_arrow(dwg, (960, 758), (920, 860), '受治理规则约束')
    add_arrow(dwg, (1350, 758), (1040, 860), '日志与权限约束')
    dwg.save()


def render_prototype_figure(svg_path: Path, spec: dict[str, object]) -> None:
    style = str(spec.get('prototype_style', 'admin'))
    dwg = create_drawing(svg_path, 1680, 1040, '#f7fbff')
    dwg.add(dwg.text(spec['title'], insert=(840, 42), text_anchor='middle', font_size=28, font_family=SVG_FONT, font_weight='bold', fill='#17304f'))

    if style == 'auth':
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 900), rx=24, ry=24, fill='#ffffff', stroke='#cfdae8', stroke_width=2))
        dwg.add(dwg.rect(insert=(40, 80), size=(760, 900), rx=24, ry=24, fill='#103250', stroke='none'))
        dwg.add(dwg.rect(insert=(700, 80), size=(940, 900), rx=24, ry=24, fill='#fdfefe', stroke='none'))
        dwg.add(dwg.text('上海人工智能实验室', insert=(120, 170), font_size=36, font_family=SVG_FONT, fill='#ffffff', font_weight='bold'))
        dwg.add(dwg.text('博士生报名与服务门户', insert=(120, 225), font_size=30, font_family=SVG_FONT, fill='#dfeeff'))
        add_box(dwg, 110, 300, 600, 130, '统一认证入口', '登录、注册、找回密码均在同一页面切换完成', '#174467')
        add_box(dwg, 110, 470, 600, 120, '招生提示', '首页可查看招生信息、报名日程与申请开放状态', '#1b4c74')
        add_box(dwg, 110, 630, 600, 140, '操作路径', '选择认证模式\n填写账号资料\n完成图形与邮箱验证码\n进入门户首页', '#205684')
        dwg.add(dwg.rect(insert=(900, 190), size=(540, 660), rx=24, ry=24, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text(str(spec['page_title']), insert=(960, 260), font_size=32, font_family=SVG_FONT, fill='#17304f', font_weight='bold'))
        tab_x = 960
        for label in ['登录', '注册', '找回密码']:
            dwg.add(dwg.rect(insert=(tab_x, 300), size=(120, 42), rx=14, ry=14, fill='#eef6ff' if label == '注册' else '#f6f8fb', stroke='#d6dfeb', stroke_width=1))
            dwg.add(dwg.text(label, insert=(tab_x + 32, 328), font_size=18, font_family=SVG_FONT, fill='#17304f'))
            tab_x += 140
        field_y = 390
        for field in ['手机号', '邮箱地址', '姓名', '图形验证码', '邮箱验证码', '密码设置']:
            dwg.add(dwg.rect(insert=(960, field_y), size=(420, 54), rx=12, ry=12, fill='#f8fafc', stroke='#d6dfeb', stroke_width=1))
            dwg.add(dwg.text(field, insert=(985, field_y + 34), font_size=18, font_family=SVG_FONT, fill='#7a8798'))
            field_y += 74
        dwg.add(dwg.rect(insert=(1290, 612), size=(90, 54), rx=12, ry=12, fill='#e8f1fb', stroke='#bfd4ea', stroke_width=1))
        dwg.add(dwg.text('发送', insert=(1318, 646), font_size=18, font_family=SVG_FONT, fill='#295d8b'))
        dwg.add(dwg.rect(insert=(960, 720), size=(420, 60), rx=16, ry=16, fill='#17304f', stroke='none'))
        dwg.add(dwg.text('提交并进入门户首页', insert=(1078, 758), font_size=22, font_family=SVG_FONT, fill='#ffffff', font_weight='bold'))
        add_box(dwg, 900, 820, 540, 110, str(spec['panel_title']), '重点关注：双验证码、冷却时间、错误反馈和成功跳转', '#f7f3ff')
    elif style == 'portal':
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 900), rx=24, ry=24, fill='#fffdf8', stroke='#e2d8c6', stroke_width=2))
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 86), rx=24, ry=24, fill='#fff4de', stroke='none'))
        dwg.add(dwg.text('上海人工智能实验室 博士生报名门户', insert=(90, 133), font_size=30, font_family=SVG_FONT, fill='#834c11', font_weight='bold'))
        nav_x = 960
        for label in ['首页', '招生信息', '在线申请', '账号管理']:
            dwg.add(dwg.text(label, insert=(nav_x, 133), font_size=20, font_family=SVG_FONT, fill='#7a5522'))
            nav_x += 130
        add_box(dwg, 90, 210, 1490, 170, str(spec['page_title']), '欢迎进入博士生报名与服务门户\n首页以信息浏览、服务入口和账号操作为核心，不出现后台管理导航。', '#fff2d8')
        card_x = 100
        for summary in spec['summary']:
            add_box(dwg, card_x, 430, 300, 110, summary, '门户式服务卡片，适合考生点击进入下一步', '#ffffff')
            card_x += 340
        dwg.add(dwg.rect(insert=(90, 590), size=(980, 280), rx=22, ry=22, fill='#ffffff', stroke='#eadfce', stroke_width=2))
        dwg.add(dwg.text('服务入口', insert=(120, 635), font_size=24, font_family=SVG_FONT, fill='#834c11', font_weight='bold'))
        headers = spec['table_headers']
        rows = spec['table_rows']
        col_width = 980 / len(headers)
        for index, header in enumerate(headers):
            x = 90 + index * col_width
            dwg.add(dwg.rect(insert=(x, 660), size=(col_width, 52), fill='#fff6e8', stroke='#eadfce', stroke_width=1))
            dwg.add(dwg.text(header, insert=(x + 18, 692), font_size=18, font_family=SVG_FONT, fill='#7a5522', font_weight='bold'))
        row_y = 712
        for row in rows:
            for index, value in enumerate(row):
                x = 90 + index * col_width
                dwg.add(dwg.rect(insert=(x, row_y), size=(col_width, 62), fill='#ffffff', stroke='#f1e7d7', stroke_width=1))
                dwg.add(dwg.text(str(value), insert=(x + 18, row_y + 38), font_size=16, font_family=SVG_FONT, fill='#71553a'))
            row_y += 62
        add_box(dwg, 1120, 590, 460, 280, str(spec['panel_title']), '学生常用操作在门户右侧区域突出展示，避免进入复杂后台结构。', '#fff9ef')
        panel_y = 670
        for item in spec['panel_lines']:
            dwg.add(dwg.rect(insert=(1150, panel_y), size=(390, 48), rx=12, ry=12, fill='#ffffff', stroke='#eadfce', stroke_width=1))
            dwg.add(dwg.text(item, insert=(1170, panel_y + 31), font_size=16, font_family=SVG_FONT, fill='#71553a'))
            panel_y += 58
    elif style == 'student-form':
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 900), rx=24, ry=24, fill='#fcfdff', stroke='#d7e3ef', stroke_width=2))
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 76), rx=24, ry=24, fill='#eaf2fb', stroke='none'))
        dwg.add(dwg.text('博士生在线申请', insert=(90, 128), font_size=30, font_family=SVG_FONT, fill='#1a3d60', font_weight='bold'))
        dwg.add(dwg.text('学生自助填写界面', insert=(1360, 128), font_size=18, font_family=SVG_FONT, fill='#557392'))
        step_x = 120
        for idx, step in enumerate(['个人信息', '教育经历', '实践经历', '成果与陈述', '确认提交'], start=1):
            fill = '#1c5c91' if idx <= 3 else '#dfe8f3'
            text_fill = '#ffffff' if idx <= 3 else '#5c7088'
            dwg.add(dwg.circle(center=(step_x, 230), r=24, fill=fill))
            dwg.add(dwg.text(str(idx), insert=(step_x - 6, 238), font_size=18, font_family=SVG_FONT, fill=text_fill, font_weight='bold'))
            dwg.add(dwg.text(step, insert=(step_x - 34, 278), font_size=15, font_family=SVG_FONT, fill='#35516d'))
            if idx < 5:
                add_arrow(dwg, (step_x + 30, 230), (step_x + 170, 230), '')
            step_x += 250
        add_box(dwg, 90, 330, 1080, 120, '填写进度', '显示当前完成比例、草稿状态和系统自动保存提示', '#eef6ff')
        section_y = 490
        for section_title in ['个人信息', '教育经历', '英语能力', '成果与个人陈述']:
            add_box(dwg, 90, section_y, 1080, 96, section_title, '按分区组织表单，支持进入、返回、保存草稿和继续填写', '#ffffff')
            section_y += 116
        add_box(dwg, 1210, 330, 360, 520, str(spec['panel_title']), '右侧固定区展示完整性校验、附件上传和提交前检查。', '#f7f3ff')
        panel_y = 410
        for item in spec['panel_lines']:
            dwg.add(dwg.rect(insert=(1240, panel_y), size=(300, 50), rx=12, ry=12, fill='#ffffff', stroke='#d9d1eb', stroke_width=1))
            dwg.add(dwg.text(item, insert=(1260, panel_y + 32), font_size=16, font_family=SVG_FONT, fill='#4f4a78'))
            panel_y += 68
        dwg.add(dwg.rect(insert=(1210, 880), size=(160, 58), rx=14, ry=14, fill='#eef2f7', stroke='none'))
        dwg.add(dwg.text('保存草稿', insert=(1253, 916), font_size=20, font_family=SVG_FONT, fill='#4f647c'))
        dwg.add(dwg.rect(insert=(1410, 880), size=(160, 58), rx=14, ry=14, fill='#1c5c91', stroke='none'))
        dwg.add(dwg.text('提交申请', insert=(1453, 916), font_size=20, font_family=SVG_FONT, fill='#ffffff', font_weight='bold'))
    else:
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 900), rx=24, ry=24, fill='#ffffff', stroke='#cfdae8', stroke_width=2))
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 54), rx=24, ry=24, fill='#eef4fb', stroke='none'))
        for offset in (80, 110, 140):
            dwg.add(dwg.circle(center=(offset, 107), r=8, fill='#d96c5f' if offset == 80 else '#e1b45f' if offset == 110 else '#66b36a'))
        dwg.add(dwg.text(f"路由：{spec['route']}", insert=(1180, 112), font_size=18, font_family=SVG_FONT, fill='#4f6583'))

        dwg.add(dwg.rect(insert=(70, 160), size=(250, 780), rx=18, ry=18, fill='#12304f', stroke='none'))
        dwg.add(dwg.text('管理导航', insert=(120, 205), font_size=24, font_family=SVG_FONT, fill='#ffffff', font_weight='bold'))
        menu_items = ['首页', '招生管理', '学生管理', '培养管理', '学位管理', '流程中心', '系统治理']
        active_text = str(spec['section'])
        current_y = 255
        for item in menu_items:
            fill = '#1f4d7a' if item.startswith(active_text[:2]) or item == active_text else '#12304f'
            dwg.add(dwg.rect(insert=(92, current_y - 24), size=(205, 44), rx=12, ry=12, fill=fill, stroke='#2d5d8d', stroke_width=1))
            dwg.add(dwg.text(item, insert=(118, current_y + 4), font_size=18, font_family=SVG_FONT, fill='#ffffff'))
            current_y += 62

        dwg.add(dwg.rect(insert=(350, 160), size=(1260, 92), rx=18, ry=18, fill='#eef6ff', stroke='none'))
        dwg.add(dwg.text(spec['page_title'], insert=(390, 214), font_size=30, font_family=SVG_FONT, fill='#17304f', font_weight='bold'))
        dwg.add(dwg.text(f"操作角色：{spec['operator']}", insert=(1260, 214), font_size=18, font_family=SVG_FONT, fill='#49627f'))

        card_x = 350
        for summary in spec['summary']:
            add_box(dwg, card_x, 280, 250, 90, summary, '支持快速查看当前重点工作', '#fff8eb')
            card_x += 280

        dwg.add(dwg.rect(insert=(350, 402), size=(860, 84), rx=18, ry=18, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text('筛选与操作区', insert=(380, 438), font_size=22, font_family=SVG_FONT, fill='#17304f', font_weight='bold'))
        filter_x = 380
        for item in spec['filters']:
            dwg.add(dwg.rect(insert=(filter_x, 448), size=(180, 24), rx=10, ry=10, fill='#eef6ff', stroke='none'))
            dwg.add(dwg.text(item, insert=(filter_x + 12, 466), font_size=14, font_family=SVG_FONT, fill='#49627f'))
            filter_x += 200

        dwg.add(dwg.rect(insert=(350, 520), size=(860, 356), rx=18, ry=18, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        headers = spec['table_headers']
        rows = spec['table_rows']
        col_count = len(headers)
        col_width = 860 / col_count
        for index, header in enumerate(headers):
            x = 350 + index * col_width
            dwg.add(dwg.rect(insert=(x, 520), size=(col_width, 56), fill='#eef6ff', stroke='#d6dfeb', stroke_width=1))
            dwg.add(dwg.text(header, insert=(x + 18, 556), font_size=18, font_family=SVG_FONT, fill='#17304f', font_weight='bold'))
        row_y = 576
        for row in rows:
            for index, value in enumerate(row):
                x = 350 + index * col_width
                dwg.add(dwg.rect(insert=(x, row_y), size=(col_width, 72), fill='#ffffff', stroke='#edf1f6', stroke_width=1))
                dwg.add(dwg.text(str(value), insert=(x + 18, row_y + 42), font_size=16, font_family=SVG_FONT, fill='#49627f'))
            row_y += 72

        add_box(dwg, 1240, 402, 340, 474, str(spec['panel_title']), '右侧面板用于说明详情、抽屉或弹窗中的关键动作。', '#f7f3ff')
        panel_y = 472
        for item in spec['panel_lines']:
            dwg.add(dwg.rect(insert=(1270, panel_y), size=(280, 54), rx=12, ry=12, fill='#ffffff', stroke='#d9d1eb', stroke_width=1))
            dwg.add(dwg.text(item, insert=(1292, panel_y + 33), font_size=16, font_family=SVG_FONT, fill='#4f4a78'))
            panel_y += 72

        dwg.add(dwg.rect(insert=(1240, 900), size=(340, 76), rx=18, ry=18, fill='#17304f', stroke='none'))
        dwg.add(dwg.text('主按钮：保存 / 提交 / 执行动作', insert=(1280, 946), font_size=18, font_family=SVG_FONT, fill='#ffffff', font_weight='bold'))
    dwg.save()


def generate_figure_assets() -> dict[str, Path]:
    generated: dict[str, Path] = {}
    for figure in FLOW_FIGURES:
        svg_path = EARS_IMAGES_DIR / figure['svg']
        png_path = EARS_IMAGES_DIR / figure['png']
        if figure['key'] == 'business-flow':
            render_business_flow_figure(svg_path)
        elif figure['key'] == 'approval-flow':
            render_approval_flow_figure(svg_path)
        else:
            render_entity_relationship_figure(svg_path)
        convert_svg_to_png(svg_path, png_path)
        generated[str(figure['key'])] = png_path

    for spec in PROTOTYPE_SPECS:
        svg_path = EARS_IMAGES_DIR / f"{spec['key']}.svg"
        png_path = EARS_IMAGES_DIR / f"{spec['key']}.png"
        render_prototype_figure(svg_path, spec)
        convert_svg_to_png(svg_path, png_path)
        generated[str(spec['key'])] = png_path
    return generated


def build_document(assets: dict[str, Path]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_FALLBACK)
    normal_style.font.size = Pt(12)
    configure_heading_styles(document)

    title = document.add_paragraph()
    title.style = "Title"
    configure_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = title.add_run("博士生生命周期管理系统\n产品需求规格说明书（EARS版本）")
    apply_run_style(run, size=18, bold=True)

    subtitle = document.add_paragraph()
    configure_paragraph(subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = subtitle.add_run(f"版本：V1.0（逆向编制）\n日期：{date.today().isoformat()}")
    apply_run_style(run, size=12)

    document.add_paragraph().add_run("")

    add_heading(document, "目录", 1)
    toc_paragraph = document.add_paragraph()
    configure_paragraph(toc_paragraph, first_line_indent_cm=0)
    insert_toc(toc_paragraph)
    document.add_section(WD_SECTION_START.NEW_PAGE)

    add_heading(document, "1. 文档说明", 1)
    add_text_paragraph(document, "本文档基于当前仓库中的前端、后端、配置、测试和部署文档进行逆向编制，用于形成符合 EARS（Easy Approach to Requirements Syntax）规范的产品需求规格说明。")
    add_text_paragraph(document, "本文档描述的是当前系统已实现或明确体现出的产品能力边界，不替代后续需由业务方正式签署的需求基线文档。")

    add_heading(document, "1.1 EARS 规范说明", 2)
    add_text_paragraph(document, "本文件的需求条目统一采用 EARS 语法模式，包括：普适型（系统应……）、事件驱动型（当……时，系统应……）、状态驱动型（在……期间，系统应……）、可选特性型（在……的地方，系统应……）、异常型（如果……，系统应……）。")

    add_heading(document, "2. 产品概述", 1)
    add_text_paragraph(document, "博士生生命周期管理系统面向博士研究生培养全过程提供管理端与学生门户两套产品能力，覆盖招生、学生主数据、培养、学位、审批、系统治理、审计和门户申请等核心业务。")
    add_text_paragraph(document, "产品默认支持“FastAPI 统一托管前后端”的部署方式，同时支持“前后端分离”部署模式；数据存储以 PostgreSQL 为主，缓存与会话以 Redis 为主。")

    add_heading(document, "2.1 主要用户角色", 2)
    add_table(document, ["角色", "职责定位", "主要使用范围"], ROLE_ROWS)

    add_heading(document, "2.2 产品边界", 2)
    add_text_paragraph(document, "本产品当前主要包含以下业务域：招生管理、学生主数据与团队管理、培养过程管理、学位管理、流程审批中心、系统治理、学生门户、通知与附件管理。")
    add_text_paragraph(document, "本产品当前通过集成链路管理保留与 OA、飞书等外部系统协同的能力，但不在本文件中展开描述外部系统的内部需求。")

    add_heading(document, "2.3 模块视图与详细设计输入", 2)
    add_text_paragraph(document, "EARS 规范约束的是需求条目的句法表达，不排斥也不替代业务流程、审批流程、实体关系、操作模式和页面原型输入等设计前置信息。为了支撑技术设计与 UI/UE 设计，本文件补充以下设计输入视图。")
    add_table(document, ["业务域", "主要使用者", "当前实现范围", "详细设计关注点"], MODULE_ROWS)

    add_heading(document, "2.4 业务主流程视图", 2)
    add_figure(document, assets['business-flow'], '图 2-1 博士生生命周期业务主流程图')
    add_text_paragraph(document, '该流程图将门户注册、在线申请、招生评审、录取建档、培养执行、学位申请和毕业归档串联为完整生命周期主线。操作方式上，业务人员应先识别当前对象处于哪一个阶段，再根据该阶段允许的动作推进到下一节点；流程中心、通知提醒和系统治理贯穿全链路，负责提供任务协同、状态提醒和审计约束。')
    add_table(document, ["流程阶段", "触发点", "核心处理", "阶段产出", "设计含义"], BUSINESS_FLOW_ROWS)

    add_heading(document, "2.5 审批流程视图", 2)
    add_figure(document, assets['approval-flow'], '图 2-2 博士生生命周期审批流程图')
    add_text_paragraph(document, '该审批流程图聚焦招生流程、科研报告流程、外出研修流程和学位流程四类关键审批链。操作方式上，审批角色统一在流程中心筛选待办任务，打开详情查看业务对象，随后执行通过、驳回或转办动作；系统同步刷新业务状态，并保留历史轨迹以便追溯。')
    add_table(document, ["审批类型", "关键节点", "主要角色", "入口模块", "设计含义"], APPROVAL_FLOW_ROWS)

    add_heading(document, "2.6 核心实体关系视图", 2)
    add_text_paragraph(document, "本视图用于说明详细设计时应重点建模的业务对象与关系，不等同于数据库 DDL 全量描述。", indent=0.74)
    add_figure(document, assets['entity-relationship'], '图 2-3 博士生生命周期核心实体关系图')
    add_text_paragraph(document, '该实体关系图突出学生、导师、团队、招生申请、培养方案、科研报告、外出研修、论文与学位、流程任务和系统治理之间的关系。操作方式上，详细设计应围绕这些核心对象建立稳定关联，确保每个对象既能在业务页面中被维护，也能在流程和日志中被追踪。')
    add_table(document, ["实体", "对象说明", "主要关联", "设计影响"], ENTITY_ROWS)
    add_heading(document, "2.6.1 核心实体属性模型", 3)
    add_text_paragraph(document, "以下表格给出关键实体的属性（字段）、属性类型、字段含义以及是否属于唯一关键属性，用于支撑后续数据库设计、接口设计和页面模型设计。")
    add_entity_model_tables(document)

    add_heading(document, "2.7 运行与操作模式", 2)
    add_table(document, ["模式", "使用者", "交互方式", "设计影响"], OPERATION_MODE_ROWS)

    add_heading(document, "2.8 页面与交互原型输入", 2)
    add_text_paragraph(document, "当前仓库未内置正式高保真原型图，因此本节提供页面级原型输入清单，作为 UI/UE 后续产出原型、交互说明和组件规范的基础。", indent=0.74)
    add_table(document, ["页面/路由", "功能定位", "关键操作", "原型设计提示"], PAGE_ROWS)
    for index, spec in enumerate(PROTOTYPE_SPECS, start=1):
        add_heading(document, f"2.8.{index} {spec['title']}", 3)
        add_figure(document, assets[str(spec['key'])], f"图 2-8-{index} {spec['title']}")
        add_text_paragraph(document, str(spec['description']))
        add_operation_list(document, '建议操作方式：', list(spec['panel_lines']))

    add_heading(document, "2.9 设计角色使用建议", 2)
    add_table(document, ["使用角色", "重点关注内容", "使用目的"], DESIGN_INPUT_ROWS)

    add_heading(document, "3. 功能需求（EARS）", 1)
    requirement_counter = 1
    for section_title, items in REQUIREMENT_SECTIONS[:10]:
        add_heading(document, section_title, 2)
        rows: list[tuple[str, ...]] = []
        for mode, statement in items:
            requirement_id = f"REQ-{requirement_counter:03d}"
            rows.append((requirement_id, mode, statement))
            requirement_counter += 1
        add_table(document, ["编号", "EARS模式", "需求条目"], rows)

    add_heading(document, "4. 非功能需求（EARS）", 1)
    for section_title, items in REQUIREMENT_SECTIONS[10:]:
        add_heading(document, section_title, 2)
        rows = []
        for mode, statement in items:
            requirement_id = f"REQ-{requirement_counter:03d}"
            rows.append((requirement_id, mode, statement))
            requirement_counter += 1
        add_table(document, ["编号", "EARS模式", "需求条目"], rows)

    add_heading(document, "5. 逆向编制说明", 1)
    add_text_paragraph(document, "本规格说明书由现有实现逆向提取而成，因此需求条目更偏重当前系统已经具备或已通过配置体现的产品能力。")
    add_text_paragraph(document, "若后续需要作为正式交付件签署，建议在本文件基础上增加业务优先级、边界排除项、验收准则、原型链接及版本变更记录。")

    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(document, "附录 A 需求组织建议", 1)
    add_text_paragraph(document, "若后续继续维护 EARS 版需求文档，建议按“平台与权限、认证与会话、业务模块、门户模块、横切能力、非功能要求”的结构持续扩展。")
    add_text_paragraph(document, "新增需求时，应优先复用本文件中的 EARS 模式分类，并同步补充业务规则、页面原型和验收口径。")

    document.core_properties.title = "博士生生命周期管理系统 产品需求规格说明书（EARS版本）"
    document.core_properties.subject = "EARS 需求规格说明"
    document.core_properties.author = "GitHub Copilot"
    return document


def main() -> None:
    ensure_asset_dirs()
    assets = generate_figure_assets()
    document = build_document(assets)
    try:
        document.save(OUTPUT_PATH)
        print(f"Generated: {OUTPUT_PATH}")
    except PermissionError:
        document.save(FALLBACK_OUTPUT_PATH)
        print(f"Generated: {FALLBACK_OUTPUT_PATH}")


if __name__ == "__main__":
    main()