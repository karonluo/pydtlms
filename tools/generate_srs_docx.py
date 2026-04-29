from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

try:
    import generate_ears_prd_docx as base
except ModuleNotFoundError:
    from tools import generate_ears_prd_docx as base


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / 'documents' / '软件需求规格说明书(SRS版).docx'
FALLBACK_OUTPUT_PATH = ROOT / 'documents' / '软件需求规格说明书(SRS版)-更新版.docx'
SRS_IMAGES_DIR = ROOT / 'documents' / 'SRS_IMAGES'

DOCUMENT_TITLE = '博士生生命周期管理系统\n软件需求规格说明书（SRS版）'
DOCUMENT_SUBTITLE = '版本：V1.0'


REFERENCE_ROWS = [
    ('博士生生命周期管理系统分析报告', '业务背景、对象范围、生命周期目标与角色背景'),
    ('前端页面与路由实现', '界面清单、操作方式、页面结构与视觉边界'),
    ('后端接口与领域服务实现', '业务边界、流程状态、实体关系和外部接口能力'),
    ('部署与运维脚本', '运行模式、反向代理、统一托管与前后端分离部署约束'),
    ('自动化测试与模拟脚本', '审批分支、关键路径、异常处理与回归边界'),
]


TERM_ROWS = [
    ('SRS', '软件需求规格说明书，用于定义系统范围、接口、功能、数据和非功能要求。'),
    ('管理端', '供平台管理员、导师、学位秘书、评分人等内部角色使用的后台工作台。'),
    ('学生门户', '供考生或学生注册、登录、找回密码、查看信息和提交申请的独立门户。'),
    ('业务键', '用于把招生、培养、学位对象与统一审批流程关联的唯一业务标识。'),
    ('在线申请 V2', '当前推荐的门户申请版本，采用章节导航、长表单与分段保存结构。'),
]


ROLE_ROWS = [
    ('平台管理员', '系统级配置、用户权限、字典、日志、集成和全局治理'),
    ('招生管理员', '招生计划、报名申请、资格审核、面试安排和录取决策'),
    ('评分人/面试官', '招生材料评分、面试安排和录取建议处理'),
    ('导师', '学生主档协同、培养方案、科研报告、外出研修与学位阶段办理'),
    ('学位秘书', '论文、盲审、答辩、授位进度维护与节点把控'),
    ('中心/团队负责人', '研究中心、团队组织与人员归属治理'),
    ('门户考生/学生', '门户注册、登录、找回密码、在线申请、查看申请进度'),
]


CONSTRAINT_ROWS = [
    ('双认证域隔离', '管理端与学生门户分别使用独立令牌、独立路由守卫和独立接口域。'),
    ('流程托管约束', '招生申请、科研报告、外出研修、论文主档等托管对象的状态应由流程动作推进。'),
    ('字典驱动 UI', '大部分枚举值来源于系统字典，前端应使用动态字典选项而非硬编码。'),
    ('主数据稳定性', '学生、导师、团队等主数据应作为相对稳定的数据源，不应被流程瞬时状态污染。'),
    ('部署双模式', '系统既支持 FastAPI 统一托管，也支持 Nginx + FastAPI 的前后端分离部署。'),
]


ENVIRONMENT_ROWS = [
    ('客户端环境', 'Chrome/Edge 现代浏览器', '支持管理端工作台和门户长表单交互'),
    ('前端运行环境', 'Vue 3 + TypeScript + Vite + Element Plus', '负责管理端与门户界面渲染'),
    ('后端运行环境', 'FastAPI + Python 3.12', '负责业务接口、认证、审批与配置能力'),
    ('关系数据库', 'PostgreSQL', '负责业务数据、流程兼容数据和治理数据存储'),
    ('缓存/会话', 'Redis 单机或 Sentinel', '负责缓存、会话和验证码相关场景'),
    ('邮件服务', 'SMTP/邮件网关', '负责门户注册验证码与通知发送'),
]


SOFTWARE_INTERFACE_ROWS = [
    ('浏览器 <-> 前端', 'HTTPS / HTML / JS / CSS', '页面访问、表单交互、静态资源加载'),
    ('前端 <-> 后端 API', 'HTTPS / JSON', '登录、业务查询、保存、提交流程、审批动作'),
    ('后端 <-> PostgreSQL', 'SQL / 连接池', '业务主数据、流程兼容层、审计与配置数据读写'),
    ('后端 <-> Redis', 'Redis 协议', '会话、缓存、验证码和部分运行态加速'),
    ('后端 <-> SMTP', 'SMTP', '邮件验证码和通知发送'),
]


COMMUNICATION_ROWS = [
    ('认证通信', 'Bearer Token', '管理端与门户分别签发并校验访问令牌'),
    ('接口格式', 'application/json', '前后端交互以 JSON 结构为主'),
    ('文件上传', 'multipart/form-data', '门户附件、招生简章图片等上传场景'),
    ('健康检查', 'HTTP GET', '供部署环境探测服务健康状态'),
]


HARDWARE_ROWS = [
    ('应用服务节点', '至少 2 核 CPU / 4GB 内存', '支撑 Web/API 服务与基础后台处理'),
    ('数据库节点', '建议 SSD 存储', '支撑业务对象、日志与流程兼容数据增长'),
    ('缓存节点', '低延迟网络连接', '支撑会话、验证码与热数据访问'),
    ('文件存储', '本地或对象存储', '支撑门户附件、招生简章与证据材料保存'),
]


MODULE_ROWS = [
    ('公共与认证', '管理端登录、门户认证、个人中心', '用户登录、登出、资料修改、密码修改、门户注册找回'),
    ('经营总览', '生命周期总览、KPI、预警', '指标查看、预警提示、业务阶段概览'),
    ('招生管理', '招生计划、报名申请、资格审核、评分、面试与录取', '计划维护、申请检索、导入导出、评分与决策'),
    ('学生与中心管理', '学生主档、注册学生、研究中心/团队', '主档维护、账号状态治理、组织关系维护'),
    ('培养管理', '培养方案、科研报告、外出研修', '培养资料维护、报告评阅、研修申请与审批'),
    ('学位管理', '论文主档、盲审意见、答辩与授位', '学位节点维护、盲审录入、授位结论形成'),
    ('流程中心', '待办任务、详情、动作执行、轨迹查看', '统一审批入口、历史追溯和动作闭环'),
    ('系统治理', '用户、角色、字典、审计、集成、日志', '基础配置、权限矩阵、日志查询和集成治理'),
    ('学生门户申请', '门户首页、在线申请 V2 与 V2 章节', '自助填报、附件上传、草稿保存与提交'),
]


PAGE_CATALOG_ROWS = [
    ('/login', '管理端登录页', '账号密码登录、失败提示、重定向恢复', '独立登录页'),
    ('/dashboard', '经营总览', '指标卡片、生命周期阶段、预警与图表', '管理端'),
    ('/recruitment', '招生工作台', '计划维护、报名检索、评分与面试安排', '管理端'),
    ('/students/records', '学生主档', '学生基础信息维护与状态管理', 'StudentsView section'),
    ('/students/portal-registrations', '注册学生', '门户注册账号治理、邮件与密码操作', 'StudentsView section'),
    ('/students/centers', '研究中心管理', '中心/团队组织维护', 'StudentsView section'),
    ('/training/plans', '培养方案管理', '方案列表、编辑和版本维护', 'TrainingView section'),
    ('/training/reports', '科研报告管理', '报告评阅、评分与状态跟踪', 'TrainingView section'),
    ('/training/outbound', '外出研修管理', '申请、审批与归来评估', 'TrainingView section'),
    ('/degree/theses', '论文主档管理', '论文状态、查重、答辩节点', 'DegreeView section'),
    ('/degree/reviews', '盲审意见管理', '盲审记录维护与评分', 'DegreeView section'),
    ('/workflow/tasks', '审批中心', '待办、详情、动作与历史轨迹', '管理端'),
    ('/system/users', '系统用户管理', '账号维护与启停', 'SystemView section'),
    ('/system/roles', '角色权限管理', '角色维护与权限配置', 'SystemView section'),
    ('/system/dict-types', '字典类型管理', '字典分类维护', 'DictView section'),
    ('/system/dict-data', '字典数据管理', '字典值维护与停启', 'DictView section'),
    ('/system/audit', '审计策略管理', '审计规则维护', 'SystemView section'),
    ('/system/integrations', '集成链路管理', '外部系统连接维护', 'SystemView section'),
    ('/system/operation-logs', '操作日志查询', '关键写操作追踪', 'SystemView section'),
    ('/system/sync-logs', '同步日志查询', '同步结果与失败追踪', 'SystemView section'),
    ('/profile', '个人空间', '个人信息维护与主题设置', '管理端'),
    ('/portal', '门户认证中心', '登录、注册、重置密码、协议查看', '门户'),
    ('/portal/home', '门户首页', '申请进度、计划列表、密码修改、信息浏览', '门户'),
    ('/portal/applicationv2', '申请表 V2', '章节导航、长表单、附件、草稿与提交', '门户'),
]


OPERATION_MODE_ROWS = [
    ('管理端表格治理模式', '顶部统计 + 筛选区 + 表格 + 对话框/抽屉', '适用于招生、学生、培养、学位、系统治理'),
    ('审批闭环模式', '待办列表 + 详情 + 动作按钮 + 历史时间线', '适用于统一流程中心和托管业务对象'),
    ('门户认证模式', '模式切换卡片 + 图形验证码 + 邮件验证码 + 结果提示', '适用于注册、登录、找回密码'),
    ('门户申请模式', '章节导航 + 长表单 + 附件上传 + 草稿保存 + 最终提交', '适用于在线申请 V2'),
    ('运维部署模式', '统一托管或前后端分离', '适用于部署手册、健康检查和反向代理配置'),
]


BUSINESS_FLOW_ROWS = [
    ('门户注册与认证', '考生进入门户并选择注册/登录/找回密码', '完成验证码校验、身份校验和账号初始化', '门户账号与认证令牌'),
    ('在线申请', '已登录门户用户进入申请页', '填写资料、上传附件、保存草稿、最终提交', '招生申请记录与业务键'),
    ('招生评审', '招生管理员在工作台处理申请', '资格审核、评分推荐、面试安排、录取决策', '录取结果或预录取结果'),
    ('学生建档', '录取结果确认', '建立学生主档、导师关系和团队归属', '可持续维护的学生主数据'),
    ('培养执行', '学生进入在培阶段', '维护培养方案、科研报告、外出研修记录', '培养轨迹与审批记录'),
    ('学位推进', '学生满足学位阶段条件', '维护论文、盲审、答辩和授位节点', '授位结论与归档状态'),
    ('治理运维', '平台管理员维护平台', '调整用户、角色、字典、审计与集成策略', '受控生效的治理配置与日志'),
]


APPROVAL_FLOW_ROWS = [
    ('招生流程', '资格审核 -> 评分推荐 -> 面试安排 -> 录取决策', '招生管理员、评分人、面试官', '招生工作台、流程中心'),
    ('科研报告流程', '学生提交 -> 导师审阅 -> 通过/退回 -> 归档', '学生、导师、管理员', '培养管理、流程中心'),
    ('外出研修流程', '学生申请 -> 导师审核 -> 管理员备案 -> 归来评估', '学生、导师、管理员', '培养管理、流程中心'),
    ('学位流程', '论文提交 -> 盲审 -> 预答辩/答辩 -> 授位结论', '学生、导师、学位秘书、管理员', '学位管理、流程中心'),
]


BUSINESS_RULE_ROWS = [
    ('BR-001', '门户注册必须同时通过图形验证码和邮件验证码校验。'),
    ('BR-002', '门户账号停用后不得继续登录门户，也不得继续提交在线申请。'),
    ('BR-003', '招生申请、科研报告、外出研修和论文主档等托管对象不得通过普通编辑接口直接改写托管状态。'),
    ('BR-004', '角色删除前必须确认无系统用户继续引用该角色。'),
    ('BR-005', '字典项变化应自动影响依赖该字典的前端下拉与状态标签渲染。'),
    ('BR-006', '门户附件必须按分类校验格式、大小与内容类型。'),
    ('BR-007', '论文未满足前置条件时不得进入下一学位节点。'),
    ('BR-008', '关键写操作必须记录操作日志，必要时还应记录同步日志或流程轨迹。'),
]


DEPLOYMENT_ROWS = [
    ('统一托管模式', 'FastAPI 同时提供 API 与构建后的前端静态资源', '部署结构简单，适合单域名场景'),
    ('前后端分离模式', 'Nginx 提供前端静态资源，FastAPI 提供 /api 与健康检查', '适合分域、静态资源缓存和独立扩缩容场景'),
    ('缓存高可用模式', 'Redis Sentinel 负责主从与主节点发现', '适合对可用性要求更高的验证码与会话场景'),
]


NONFUNCTIONAL_ROWS = {
    '性能要求': [
        ('NFR-PERF-001', '普通列表查询应提供分页与筛选能力，避免一次返回过大数据集。'),
        ('NFR-PERF-002', '门户草稿保存和申请提交应在单次请求中明确返回成功或失败，不得出现无提示中断。'),
        ('NFR-PERF-003', '系统应支持通过缓存降低会话、验证码和热点读取场景的响应开销。'),
    ],
    '安全要求': [
        ('NFR-SEC-001', '管理端与门户必须实施独立认证和授权控制，禁止令牌混用。'),
        ('NFR-SEC-002', '关键接口不得仅依赖前端隐藏按钮做权限控制，后端必须再次校验。'),
        ('NFR-SEC-003', '密码修改必须校验原密码，新密码应满足长度与一致性规则。'),
    ],
    '可靠性与可用性': [
        ('NFR-AVL-001', '系统应提供健康检查接口供部署环境进行存活探测。'),
        ('NFR-AVL-002', '系统应支持 Redis 单机与哨兵两种模式，以适配不同级别的高可用部署要求。'),
        ('NFR-AVL-003', '系统应在关键失败场景返回清晰错误信息，便于用户重试或管理员排障。'),
    ],
    '可维护性与扩展性': [
        ('NFR-MNT-001', '系统应通过统一配置管理数据库、缓存、邮件和门户开关等参数。'),
        ('NFR-MNT-002', '审批兼容层变更时应同步维护运行态、历史态和 Flowable 风格兼容数据。'),
        ('NFR-MNT-003', '系统应允许通过字典维护、角色维护和门户开关实现部分功能的配置化扩展。'),
    ],
    '审计与运维': [
        ('NFR-OPS-001', '关键写操作和同步动作应可在日志中追溯到操作人、时间和结果。'),
        ('NFR-OPS-002', '系统应支持通过标准 Web 服务器和进程托管工具部署。'),
        ('NFR-OPS-003', '前后端分离部署时应允许通过统一 API 路径进行反向代理。'),
    ],
}


ACCEPTANCE_ROWS = [
    ('认证与权限', '验证管理端登录、门户注册/登录、角色权限控制与停用账户拦截'),
    ('门户申请', '验证 V2 申请填写、草稿保存、附件上传、最终提交与关闭开关'),
    ('招生链路', '验证计划维护、申请导入、资格审核、评分、面试和录取决策'),
    ('学生与中心', '验证主档维护、注册学生治理、研究中心/团队维护'),
    ('培养与学位', '验证培养方案、科研报告、外出研修、论文与盲审链路'),
    ('流程中心', '验证待办筛选、详情查看、动作执行与轨迹留痕'),
    ('系统治理', '验证用户、角色、字典、日志和集成配置治理'),
    ('部署运维', '验证统一托管、前后端分离、健康检查与日志可追溯'),
]


INTERFACE_OPERATION_SCENARIOS = {
    '平台管理员典型操作方式': [
        '进入系统登录页，输入账号密码并完成登录。',
        '从左侧菜单选择目标模块，通过顶部统计与筛选区缩小处理范围。',
        '在表格中定位业务对象，进入对话框或抽屉编辑详细信息。',
        '对高风险操作执行二次确认，并在完成后核验日志或同步结果。',
    ],
    '导师/学位秘书典型操作方式': [
        '进入培养管理、学位管理或流程中心，按学生、状态、周期或节点进行过滤。',
        '打开详情查看业务对象当前阶段、相关附件和历史记录。',
        '执行审阅、通过、退回、录入评语或推进节点等动作。',
        '回到列表确认状态已同步更新并保留处理痕迹。',
    ],
    '门户考生典型操作方式': [
        '进入门户认证页，选择注册、登录或找回密码模式。',
        '完成图形验证码、邮件验证码与身份信息校验。',
        '进入门户首页后查看申请进度、招生信息并打开在线申请。',
        '按章节填写申请表、保存草稿、上传附件并最终提交。',
    ],
}


FUNCTIONAL_REQUIREMENTS = {
    '4.1 公共认证与个人中心': [
        ('FR-AUTH-001', '管理端登录', '管理端用户', '进入登录页并提交账号密码', '系统校验账号、密码与启用状态，成功后进入原目标页面。'),
        ('FR-AUTH-002', '门户注册', '门户考生', '在门户认证页提交注册信息', '系统校验邮箱、手机号、身份证号、邮件验证码并创建门户账号。'),
        ('FR-AUTH-003', '找回密码', '门户考生', '在认证页切换到找回密码模式', '系统校验身份后允许重置密码并提示重新登录。'),
        ('FR-AUTH-004', '个人资料维护', '管理端用户', '进入个人空间或门户密码修改入口', '系统允许修改个人资料或密码并反馈结果。'),
    ],
    '4.2 经营总览': [
        ('FR-DASH-001', '生命周期指标展示', '管理端用户', '进入经营总览页', '系统展示招生、学生、培养、学位与流程相关关键指标。'),
        ('FR-DASH-002', '预警信息提示', '管理端用户', '存在超期或风险事项', '系统在总览页展示预警列表和责任归属。'),
        ('FR-DASH-003', '生命周期阶段图', '管理端用户', '打开总览页', '系统展示从招生准备到毕业归档的阶段化总览。'),
    ],
    '4.3 招生管理': [
        ('FR-REC-001', '招生计划维护', '招生管理员', '进入招生工作台', '系统支持新增、编辑、查询和删除招生计划。'),
        ('FR-REC-002', '申请检索', '招生管理员', '选择计划或输入筛选条件', '系统按计划、状态、方向等维度检索报名申请。'),
        ('FR-REC-003', '材料导入导出', '招生管理员', '执行模板导入或导出动作', '系统返回导入结果摘要或导出文件。'),
        ('FR-REC-004', '资格审核与评分', '招生管理员/评分人', '打开申请详情或执行流程动作', '系统支持资格审核、评分录入与过程状态更新。'),
        ('FR-REC-005', '面试安排与录取决策', '招生管理员/面试官', '进入面试安排或录取阶段', '系统支持面试分组、时间安排和录取结论形成。'),
    ],
    '4.4 学生与研究中心管理': [
        ('FR-STU-001', '学生主档维护', '平台管理员/导师', '进入学生主档工作台', '系统支持维护学生基础资料、状态、导师与中心归属。'),
        ('FR-STU-002', '注册学生治理', '平台管理员', '进入注册学生工作台', '系统展示门户账号状态并支持启用、停用、发送邮件和重置密码。'),
        ('FR-STU-003', '研究中心维护', '平台管理员', '进入研究中心工作台', '系统支持新增、编辑、批量删除和状态维护。'),
        ('FR-STU-004', '主数据联动约束', '平台管理员', '保存学生与中心/导师关系', '系统校验归属关系有效性并保持统计一致。'),
    ],
    '4.5 培养管理': [
        ('FR-TRN-001', '培养方案维护', '导师/管理员', '进入培养方案工作台', '系统支持方案查询、编辑、版本记录和保存。'),
        ('FR-TRN-002', '科研报告评阅', '导师/管理员', '进入科研报告工作台', '系统支持报告录入、评分评语和状态跟踪。'),
        ('FR-TRN-003', '外出研修办理', '学生/导师/管理员', '进入外出研修工作台', '系统支持申请、审核、备案和归来评估。'),
        ('FR-TRN-004', '培养过滤与分页', '培养角色', '输入学生、导师、状态等过滤条件', '系统应以分页方式返回目标记录。'),
    ],
    '4.6 学位管理': [
        ('FR-DEG-001', '论文主档维护', '导师/学位秘书/管理员', '进入论文主档工作台', '系统支持维护论文题目、状态、查重率和阶段信息。'),
        ('FR-DEG-002', '盲审意见维护', '学位秘书/评阅专家', '进入盲审意见工作台', '系统支持维护评分、意见和评阅状态。'),
        ('FR-DEG-003', '节点前置校验', '学位角色', '推进论文到下一节点', '系统校验前置条件，不满足时阻止推进。'),
        ('FR-DEG-004', '授位状态维护', '学位秘书/管理员', '答辩完成并满足条件', '系统形成授位结论并保留过程信息。'),
    ],
    '4.7 流程中心': [
        ('FR-WF-001', '待办列表展示', '审批角色', '进入流程中心', '系统展示待办、已办、超时和业务类型相关任务。'),
        ('FR-WF-002', '任务详情查看', '审批角色', '打开任务详情', '系统展示业务对象信息、可执行动作和处理历史。'),
        ('FR-WF-003', '动作执行', '审批角色', '点击通过、驳回等动作', '系统同步更新流程状态、业务状态和历史记录。'),
        ('FR-WF-004', '历史轨迹追溯', '审批角色', '查看任务历史', '系统展示处理人、处理意见、节点变化和时间信息。'),
    ],
    '4.8 系统治理': [
        ('FR-SYS-001', '系统用户维护', '平台管理员', '进入用户管理工作台', '系统支持新增、编辑、启用、停用和查询系统用户。'),
        ('FR-SYS-002', '角色权限配置', '平台管理员', '进入角色管理工作台', '系统支持角色新增、权限分配和唯一性校验。'),
        ('FR-SYS-003', '字典维护', '平台管理员', '进入字典类型或字典数据工作台', '系统支持维护枚举分类和值，并供业务页面复用。'),
        ('FR-SYS-004', '审计与集成治理', '平台管理员', '进入审计或集成工作台', '系统支持维护审计策略、外部集成配置并查询日志。'),
        ('FR-SYS-005', '日志查询', '平台管理员/审计员', '进入操作日志或同步日志页面', '系统支持按关键字、状态或时间过滤日志。'),
    ],
    '4.9 学生门户': [
        ('FR-PORTAL-001', '门户首页', '门户考生', '登录后进入门户首页', '系统展示申请进度、计划列表、公告入口与密码修改入口。'),
        ('FR-PORTAL-002', '申请表 V2', '门户考生', '进入在线申请页', '系统支持章节导航、长表单、附件上传、保存草稿和提交。'),
        ('FR-PORTAL-003', '附件上传', '门户考生', '在申请表中上传附件', '系统校验附件分类、格式、大小并返回可用地址。'),
        ('FR-PORTAL-004', '申请入口开关', '平台管理员/门户考生', '申请入口被关闭', '系统阻止进入申请页并给出友好提示。'),
    ],
}


DATA_SUMMARY_ROWS = [
    ('主数据层', '学生、导师、团队/中心、系统用户、角色、字典', '提供稳定基础对象与受控选项来源'),
    ('招生业务层', '招生计划、招生申请、面试安排、录取结果', '支撑招生阶段业务推进'),
    ('培养业务层', '培养方案、科研报告、外出研修', '支撑培养过程维护与审批协同'),
    ('学位业务层', '论文、盲审意见、答辩/授位信息', '支撑学位阶段节点管理'),
    ('流程层', '流程任务、动作日志、流程兼容数据', '支撑统一审批和历史追溯'),
    ('审计治理层', '操作日志、同步日志、审计策略、集成链路', '支撑平台治理与追踪'),
]


ENTITY_ROWS = [
    ('系统用户', '管理端登录主体，承载角色与权限引用。', '角色、操作日志'),
    ('角色', '权限集合载体，用于菜单和接口鉴权。', '系统用户、权限目录'),
    ('门户账号', '门户登录主体，用于注册、登录、找回密码和申请。', '门户资料、招生申请'),
    ('招生计划', '招生批次和基础配置的载体。', '招生申请、招生简章'),
    ('招生申请', '考生正式报名业务对象。', '门户账号、流程任务、评分与录取'),
    ('学生', '博士生生命周期核心主实体。', '导师、团队、培养、学位'),
    ('导师', '指导教师主数据。', '学生、团队、培养和学位'),
    ('团队/研究中心', '研究组织单元。', '学生、导师'),
    ('培养方案', '培养阶段版本化对象。', '学生、导师'),
    ('科研报告', '周期性科研汇报对象。', '学生、流程任务'),
    ('外出研修', '出访审批与归来评估对象。', '学生、流程任务'),
    ('论文', '学位阶段主线对象。', '学生、盲审意见、流程任务'),
    ('流程任务', '统一审批待办与处理轨迹载体。', '业务键、业务对象、动作日志'),
    ('字典项', '动态枚举选项来源。', '多个业务字段和状态标签'),
]


SRS_ENTITY_MODEL_SPECS = list(base.ENTITY_MODEL_SPECS) + [
    {
        'entity': '系统用户',
        'description': '系统用户实体用于承载管理端登录账号、角色和启停状态。',
        'rows': [
            ('user_id', '文本', '系统用户唯一标识。', '是'),
            ('username', '文本', '登录用户名。', '是'),
            ('display_name', '文本', '展示名称。', '否'),
            ('role_codes', '文本集合', '关联角色编码集合。', '否'),
            ('status', '枚举字典', '启用、停用等状态。', '否'),
            ('last_login_at', '日期时间', '最近登录时间。', '否'),
        ],
    },
    {
        'entity': '角色',
        'description': '角色实体用于聚合权限目录并向系统用户分配权限边界。',
        'rows': [
            ('role_id', '文本', '角色唯一标识。', '是'),
            ('role_code', '文本', '角色编码。', '是'),
            ('role_name', '文本', '角色名称。', '否'),
            ('permission_codes', '文本集合', '权限编码集合。', '否'),
            ('status', '枚举字典', '生效、停用等状态。', '否'),
        ],
    },
    {
        'entity': '字典项',
        'description': '字典项实体用于维护业务页面中的枚举值、标签和颜色信息。',
        'rows': [
            ('dict_data_id', '文本', '字典数据唯一标识。', '是'),
            ('dict_type_code', '文本', '所属字典类型编码。', '否'),
            ('label', '文本', '显示标签。', '否'),
            ('value', '文本', '实际值。', '是'),
            ('status', '枚举字典', '启用、停用等状态。', '否'),
            ('color_token', '文本', '前端状态标签颜色提示。', '否'),
        ],
    },
]


PROTOTYPE_SPECS = [
    {
        'key': 'login-view',
        'title': '管理端登录页原型图',
        'prototype_style': 'login',
        'section': '公共认证',
        'route': '/login',
        'operator': '管理端用户',
        'page_title': '博士生生命周期管理系统',
        'summary': ['统一登录', '失败反馈', '重定向恢复'],
        'filters': ['账号', '密码', '登录提示'],
        'table_headers': ['区域', '主要元素', '操作方式'],
        'table_rows': [
            ['品牌区', '系统名称、说明文案', '展示系统定位'],
            ['登录区', '账号、密码、按钮', '输入并提交登录'],
            ['反馈区', '错误对话框、跳转提示', '失败或成功后反馈'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['输入账号和密码', '提交登录请求', '失败时弹窗提示', '成功后跳转目标页面'],
        'description': '管理端登录页应采用简洁的双分区布局，左侧展示系统定位与业务口号，右侧展示账号密码登录卡片。页面不承载复杂导航，仅突出认证入口与反馈逻辑。',
    },
    {
        'key': 'dashboard-view',
        'title': '经营总览原型图',
        'prototype_style': 'dashboard',
        'section': '经营总览',
        'route': '/dashboard',
        'operator': '管理端用户',
        'page_title': '经营总览',
        'summary': ['生命周期覆盖', '在途审批', '风险预警'],
        'filters': ['阶段图', '指标图表', '预警面板'],
        'table_headers': ['指标', '说明', '展示方式'],
        'table_rows': [
            ['招生计划', '当前开放计划规模', 'KPI 卡片'],
            ['学生规模', '当前学生主档数量', 'KPI 卡片'],
            ['盲审待办', '当前待处理学位事项', '柱线图/预警'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['查看 KPI 卡片', '横向浏览生命周期阶段', '关注风险预警', '进入对应模块处理'],
        'description': '经营总览页用于快速观察招生、培养、学位、审批和治理的整体运行态势，应突出生命周期阶段、KPI 卡片和预警事项。',
    },
    {
        'key': 'recruitment-view',
        'title': '招生工作台原型图',
        'prototype_style': 'admin',
        'section': '招生管理',
        'route': '/recruitment',
        'operator': '平台管理员、评分人、面试官',
        'page_title': '招生工作台',
        'summary': ['当前计划', '待审核申请', '待面试安排'],
        'filters': ['计划切换', '申请状态', '研究方向', '导入/导出'],
        'table_headers': ['考生', '资格审核', '材料评分', '面试安排', '录取状态'],
        'table_rows': [
            ['张三', '已通过', '89.5', '第1组 09:00', '待决策'],
            ['李四', '待审核', '--', '未安排', '未开始'],
            ['王五', '已通过', '92.0', '第2组 14:00', '预录取'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['先切换招生计划', '筛选目标申请', '执行审核和评分', '安排面试并形成决策'],
        'description': '招生工作台是高密度作业界面，应突出阶段化处理、申请检索、导入导出和评分/面试安排能力。',
    },
    {
        'key': 'students-records-view',
        'title': '学生主档原型图',
        'prototype_style': 'admin',
        'section': '学生主档',
        'route': '/students/records',
        'operator': '平台管理员、导师',
        'page_title': '学生主档管理',
        'summary': ['在籍学生', '导师归属', '状态卡片'],
        'filters': ['学生状态', '团队/中心', '导师', '关键词'],
        'table_headers': ['学生', '团队', '导师', '当前状态', '操作'],
        'table_rows': [
            ['张三', '智能感知团队', '刘亚', '在培', '编辑'],
            ['李四', '认知模型团队', '周青', '待入学', '编辑'],
            ['王五', '具身智能团队', '刘亚', '毕业', '查看'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['筛选学生对象', '打开编辑对话框', '选择团队与导师', '保存后刷新状态'],
        'description': '学生主档页应突出主数据稳定性，支持围绕学生、导师、中心/团队和状态进行治理。',
    },
    {
        'key': 'students-registrations-view',
        'title': '注册学生管理原型图',
        'prototype_style': 'admin',
        'section': '注册学生',
        'route': '/students/portal-registrations',
        'operator': '平台管理员',
        'page_title': '注册学生管理',
        'summary': ['门户账号', '停用账号', '待通知对象'],
        'filters': ['账号状态', '手机号/邮箱', '关键词', '批量操作'],
        'table_headers': ['姓名', '手机号', '邮箱', '账号状态', '操作'],
        'table_rows': [
            ['张三', '138****0001', 'a@example.com', '启用', '停用/重置密码'],
            ['李四', '138****0002', 'b@example.com', '未开通', '发送通知'],
            ['王五', '138****0003', 'c@example.com', '停用', '启用'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['按状态检索账号', '执行启停或密码重置', '必要时发送邮件通知', '回查操作结果'],
        'description': '注册学生管理页用于治理门户侧已注册账号，重点体现启用/停用、邮件通知和密码重置。',
    },
    {
        'key': 'students-centers-view',
        'title': '研究中心管理原型图',
        'prototype_style': 'admin',
        'section': '研究中心管理',
        'route': '/students/centers',
        'operator': '平台管理员',
        'page_title': '研究中心管理',
        'summary': ['中心总数', '负责人导师', '学生规模'],
        'filters': ['中心状态', '负责人', '研究方向'],
        'table_headers': ['中心/团队', '负责人导师', '导师数', '学生数', '状态'],
        'table_rows': [
            ['智能感知团队', '刘亚', '8', '32', '启用'],
            ['认知模型团队', '周青', '6', '21', '启用'],
            ['具身智能团队', '袁野', '5', '17', '筹建'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['筛选中心或团队', '打开详情维护成员', '维护研究方向', '核验学生归属统计'],
        'description': '研究中心管理页用于维护组织单元及其成员关系，应突出负责人、研究方向和学生归属统计。',
    },
    {
        'key': 'training-plans-view',
        'title': '培养方案管理原型图',
        'prototype_style': 'admin',
        'section': '培养管理',
        'route': '/training/plans',
        'operator': '导师、平台管理员',
        'page_title': '培养方案管理',
        'summary': ['待确认方案', '最近维护', '版本状态'],
        'filters': ['学生', '导师', '状态', '周期'],
        'table_headers': ['学生', '方案版本', '状态', '生效日期', '操作'],
        'table_rows': [
            ['张三', 'V3', '已确认', '2026-03-01', '编辑'],
            ['李四', 'V1', '待确认', '2026-04-15', '维护'],
            ['王五', 'V2', '已归档', '2025-09-01', '查看'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['切换方案页签', '筛选学生与状态', '进入表单编辑', '保存后刷新版本信息'],
        'description': '培养方案管理页用于维护培养阶段计划版本，应突出学生、导师、状态和版本信息。',
    },
    {
        'key': 'training-reports-view',
        'title': '科研报告管理原型图',
        'prototype_style': 'admin',
        'section': '培养管理',
        'route': '/training/reports',
        'operator': '导师、平台管理员',
        'page_title': '科研报告管理',
        'summary': ['待审报告', '已通过', '待退回'],
        'filters': ['学生', '周期', '状态', '导师'],
        'table_headers': ['学生', '报告周期', '状态', '审阅分数', '操作'],
        'table_rows': [
            ['张三', '2026 春季', '待审阅', '--', '审阅'],
            ['李四', '2026 春季', '已通过', '92', '查看'],
            ['王五', '2026 冬季', '退回修改', '80', '再次审阅'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['筛选报告周期', '进入评阅对话框', '填写评语和分数', '提交后同步流程状态'],
        'description': '科研报告管理页聚焦周期性报告提交与评阅，需体现评分、评语与流程状态联动。',
    },
    {
        'key': 'training-outbound-view',
        'title': '外出研修管理原型图',
        'prototype_style': 'admin',
        'section': '培养管理',
        'route': '/training/outbound',
        'operator': '导师、平台管理员',
        'page_title': '外出研修管理',
        'summary': ['待审核申请', '已备案', '归来评估'],
        'filters': ['学生', '状态', '类型', '时间'],
        'table_headers': ['学生', '研修类型', '目的地', '审批状态', '操作'],
        'table_rows': [
            ['张三', '联合培养', '新加坡', '待审核', '审核'],
            ['李四', '短期访问', '北京', '已备案', '查看'],
            ['王五', '国际会议', '巴黎', '已归来', '归来评估'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['检索外出申请', '打开详情核验信息', '执行审核或备案', '记录归来评估'],
        'description': '外出研修管理页应突出申请、审批、备案和归来评估的闭环。',
    },
    {
        'key': 'degree-theses-view',
        'title': '论文主档管理原型图',
        'prototype_style': 'admin',
        'section': '学位管理',
        'route': '/degree/theses',
        'operator': '导师、学位秘书、平台管理员',
        'page_title': '论文主档管理',
        'summary': ['论文总数', '盲审处理中', '待答辩安排'],
        'filters': ['论文状态', '学生', '导师', '学位阶段'],
        'table_headers': ['学生', '论文状态', '查重率', '答辩安排', '授位状态'],
        'table_rows': [
            ['张三', '待盲审', '8%', '未安排', '未开始'],
            ['李四', '盲审中', '11%', '未安排', '未开始'],
            ['王五', '答辩完成', '6%', '已完成', '待授位'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['按节点筛选论文', '打开里程碑详情', '维护阶段性信息', '推进授位结论'],
        'description': '论文主档管理页应体现学位阶段的里程碑感，支持盲审、答辩和授位状态展示。',
    },
    {
        'key': 'degree-reviews-view',
        'title': '盲审意见管理原型图',
        'prototype_style': 'admin',
        'section': '学位管理',
        'route': '/degree/reviews',
        'operator': '学位秘书、评阅专家',
        'page_title': '盲审意见管理',
        'summary': ['待回收意见', '已评分', '待复核'],
        'filters': ['论文', '评阅状态', '专家', '时间'],
        'table_headers': ['学生', '评阅专家', '评阅状态', '综合评分', '操作'],
        'table_rows': [
            ['张三', '专家A', '待回收', '--', '提醒'],
            ['李四', '专家B', '已提交', '88', '查看'],
            ['王五', '专家C', '待复核', '91', '复核'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['筛选目标评阅记录', '录入或查看意见', '复核评分', '同步论文状态'],
        'description': '盲审意见管理页侧重专家意见、评分和复核动作的组织。',
    },
    {
        'key': 'workflow-center-view',
        'title': '审批中心原型图',
        'prototype_style': 'admin',
        'section': '流程中心',
        'route': '/workflow/tasks',
        'operator': '全审批角色',
        'page_title': '流程待办中心',
        'summary': ['待办任务', '已办任务', '即将超时'],
        'filters': ['任务状态', '业务类型', '办理人', '创建时间'],
        'table_headers': ['任务', '业务类型', '当前节点', '发起人', '截止时间'],
        'table_rows': [
            ['资格审核-张三', '招生流程', '资格审核', '招生管理员', '今天 18:00'],
            ['科研报告-李四', '科研报告流程', '导师审阅', '李四', '明天 12:00'],
            ['论文-王五', '学位流程', '盲审确认', '王五', '后天 09:00'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['筛选待办任务', '查看业务详情', '执行通过或退回', '查看历史轨迹'],
        'description': '审批中心原型应突出任务列表、详情面板、动作区和历史轨迹的完整闭环。',
    },
    {
        'key': 'system-users-view',
        'title': '系统用户管理原型图',
        'prototype_style': 'admin',
        'section': '系统治理',
        'route': '/system/users',
        'operator': '平台管理员',
        'page_title': '系统用户管理',
        'summary': ['系统用户', '停用用户', '最近变更'],
        'filters': ['状态', '角色', '关键词'],
        'table_headers': ['用户名', '姓名', '角色', '状态', '操作'],
        'table_rows': [
            ['admin', '管理员', 'platform_admin', '启用', '编辑/停用'],
            ['liu.ya', '刘亚', 'advisor', '启用', '编辑'],
            ['zhou.qing', '周青', 'degree_secretary', '停用', '启用'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['检索目标用户', '编辑角色和状态', '执行启停', '核查日志留痕'],
        'description': '系统用户管理页用于维护后台账号，需体现角色、状态和高风险操作确认。',
    },
    {
        'key': 'system-roles-view',
        'title': '角色权限管理原型图',
        'prototype_style': 'admin',
        'section': '系统治理',
        'route': '/system/roles',
        'operator': '平台管理员',
        'page_title': '角色权限管理',
        'summary': ['角色数量', '权限矩阵', '待审配置'],
        'filters': ['角色状态', '关键词', '权限树'],
        'table_headers': ['角色编码', '角色名称', '权限数', '状态', '操作'],
        'table_rows': [
            ['platform_admin', '平台管理员', '36', '生效', '编辑'],
            ['advisor', '导师', '12', '生效', '编辑'],
            ['degree_secretary', '学位秘书', '15', '生效', '授权'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['新增或检索角色', '配置权限树', '保存角色配置', '核验用户生效范围'],
        'description': '角色权限管理页应突出权限矩阵配置和唯一性校验。',
    },
    {
        'key': 'dict-view',
        'title': '字典管理原型图',
        'prototype_style': 'admin',
        'section': '系统治理',
        'route': '/system/dict-types + /system/dict-data',
        'operator': '平台管理员',
        'page_title': '字典管理',
        'summary': ['字典类型', '字典数据', '业务复用'],
        'filters': ['类型切换', '状态', '关键词'],
        'table_headers': ['对象', '编码', '标签/名称', '状态', '操作'],
        'table_rows': [
            ['字典类型', 'degree_type', '学位类型', '启用', '编辑'],
            ['字典数据', 'doctor', '博士', '启用', '编辑'],
            ['字典数据', 'master', '硕士', '停用', '启用'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['切换类型或数据页签', '检索目标条目', '维护标签和值', '核验前端下拉展示'],
        'description': '字典管理页用于维护系统枚举来源，应体现类型与数据的级联关系。',
    },
    {
        'key': 'profile-view',
        'title': '个人空间原型图',
        'prototype_style': 'admin',
        'section': '个人空间',
        'route': '/profile',
        'operator': '管理端用户',
        'page_title': '个人空间',
        'summary': ['个人资料', '密码修改', '主题设置'],
        'filters': ['资料区', '密码区', '主题色'],
        'table_headers': ['区域', '内容', '操作'],
        'table_rows': [
            ['资料区', '姓名、邮箱、电话', '编辑保存'],
            ['密码区', '原密码、新密码', '修改密码'],
            ['主题区', '主题色预设', '切换预览'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['查看和编辑资料', '修改密码', '切换主题色', '保存个人设置'],
        'description': '个人空间页用于展示和维护当前登录用户的个人信息与个性化设置。',
    },
    {
        'key': 'portal-auth-view',
        'title': '门户认证原型图',
        'prototype_style': 'auth',
        'section': '门户与认证',
        'route': '/portal',
        'operator': '门户考生',
        'page_title': '门户认证中心',
        'summary': ['注册引导', '登录校验', '找回密码'],
        'filters': ['模式切换', '验证码状态', '邮箱校验'],
        'table_headers': ['步骤', '输入项', '系统反馈', '操作'],
        'table_rows': [
            ['注册', '手机号、邮箱、姓名', '校验必填和格式', '下一步'],
            ['验证码', '图形验证码、邮箱验证码', '提示冷却和有效期', '发送验证码'],
            ['完成', '密码、确认密码', '注册成功提示', '提交注册'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['选择认证模式', '输入账号资料', '完成双验证码校验', '提交后进入门户首页'],
        'description': '门户认证页用于承载登录、注册和找回密码三种模式，并通过图形验证码和邮件验证码保护注册与重置链路。',
    },
    {
        'key': 'portal-home-view',
        'title': '门户首页原型图',
        'prototype_style': 'portal',
        'section': '学生门户',
        'route': '/portal/home',
        'operator': '门户考生',
        'page_title': '博士生报名与服务门户',
        'summary': ['报名入口', '招生信息', '账号管理'],
        'filters': ['首页导航', '公告入口', '个人中心', '外链跳转'],
        'table_headers': ['服务卡片', '当前状态', '说明', '操作'],
        'table_rows': [
            ['在线申请', '未提交', '进入学生申报页', '立即填写'],
            ['账号管理', '正常', '修改密码与安全设置', '立即管理'],
            ['招生信息', '已开放', '查看实验室招生资讯', '立即查看'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['从顶部导航进入服务', '查看公告与进度', '打开在线申请或修改密码', '返回首页继续浏览信息'],
        'description': '门户首页应体现面向考生的信息架构，突出申请进度、服务入口和招生资讯，而不是后台工作台风格。',
    },
    {
        'key': 'portal-application-v2-view',
        'title': '在线申请 V2 主界面原型图',
        'prototype_style': 'portal-v2',
        'section': '学生申报',
        'route': '/portal/applicationv2',
        'operator': '门户考生',
        'page_title': '博士研究生申请表 V2',
        'summary': ['章节导航', '状态指示', '保存与提交'],
        'filters': ['左侧章节', '当前章节', '右侧表单'],
        'table_headers': ['模块', '说明', '操作方式'],
        'table_rows': [
            ['左侧章节栏', '显示 8 个章节及完成状态', '点击跳转'],
            ['主表单区', '显示当前章节字段', '填写并保存'],
            ['底部操作区', '上一章、下一章、保存草稿、提交', '推进流程'],
        ],
        'panel_title': '操作说明',
        'panel_lines': ['沿左侧章节逐步填写', '当前章节内完成字段与附件', '随时保存草稿', '最终提交并确认声明'],
        'description': '在线申请 V2 采用章节导航与长表单嵌入式布局，是当前推荐的门户申请版本。',
    },
    {
        'key': 'portal-v2-basic-section',
        'title': 'V2 基本信息章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalBasicSection',
        'operator': '门户考生',
        'page_title': '基本信息',
        'summary': ['身份信息', '联系方式', '头像上传'],
        'form_fields': ['姓名', '身份证号', '出生日期', '联系方式', '个人照片'],
        'panel_title': '操作说明',
        'panel_lines': ['核验身份资料', '上传个人照片', '补充联系方式', '保存章节'],
        'description': '基本信息章节用于填写身份、联系方式与个人照片，是申请表后续章节的前置基础。',
    },
    {
        'key': 'portal-v2-application-section',
        'title': 'V2 报名信息章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalApplicationSection',
        'operator': '门户考生',
        'page_title': '报名信息',
        'summary': ['志愿信息', '导师意向', '获知渠道'],
        'form_fields': ['招生计划', '第一志愿中心', '第二志愿中心', '意向导师', '获知渠道'],
        'panel_title': '操作说明',
        'panel_lines': ['选择招生计划', '设置志愿', '关联导师', '保存章节'],
        'description': '报名信息章节用于采集招生计划、志愿和导师意向，是招生匹配的关键信息入口。',
    },
    {
        'key': 'portal-v2-education-section',
        'title': 'V2 教育经历章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalEducationSection',
        'operator': '门户考生',
        'page_title': '教育经历',
        'summary': ['多条记录', '成绩单附件', '学位证附件'],
        'form_fields': ['学校名称', '专业', '起止时间', '平均成绩/GPA', '成绩单/学位证上传'],
        'panel_title': '操作说明',
        'panel_lines': ['新增多条教育经历', '上传成绩与证书附件', '校验起止时间', '保存章节'],
        'description': '教育经历章节应支持多条记录和多附件上传，是申请材料核验的核心章节。',
    },
    {
        'key': 'portal-v2-practice-section',
        'title': 'V2 实践经历章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalPracticeSection',
        'operator': '门户考生',
        'page_title': '实践经历',
        'summary': ['多条记录', '项目职责', '证明人'],
        'form_fields': ['机构名称', '职位/项目', '起止时间', '职责描述', '证明人'],
        'panel_title': '操作说明',
        'panel_lines': ['新增实践经历', '填写职责描述', '补充证明人信息', '保存章节'],
        'description': '实践经历章节用于补充项目、实习和工程实践等经历。',
    },
    {
        'key': 'portal-v2-english-section',
        'title': 'V2 英语能力章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalEnglishSection',
        'operator': '门户考生',
        'page_title': '英语语言能力',
        'summary': ['考试记录', '证书上传', '多条维护'],
        'form_fields': ['考试名称', '成绩', '考试时间', '证书附件'],
        'panel_title': '操作说明',
        'panel_lines': ['维护多条考试成绩', '上传证书附件', '检查成绩格式', '保存章节'],
        'description': '英语能力章节用于维护多条考试记录和证书材料。',
    },
    {
        'key': 'portal-v2-family-section',
        'title': 'V2 家庭情况章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalFamilySection',
        'operator': '门户考生',
        'page_title': '家庭情况',
        'summary': ['家庭成员', '父母必填', '联系方式'],
        'form_fields': ['成员姓名', '关系', '工作单位', '职务', '联系电话'],
        'panel_title': '操作说明',
        'panel_lines': ['至少填写父母信息', '补充其他成员', '核对联系方式', '保存章节'],
        'description': '家庭情况章节用于采集家庭成员基础信息，并强调父母信息必填。',
    },
    {
        'key': 'portal-v2-achievement-section',
        'title': 'V2 获奖经历章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalAchievementSection',
        'operator': '门户考生',
        'page_title': '获奖经历',
        'summary': ['论文成果', '获奖信息', '多条维护'],
        'form_fields': ['成果类型', '题目/奖项', '发表或获奖时间', '级别', '说明'],
        'panel_title': '操作说明',
        'panel_lines': ['新增成果记录', '补充奖项和论文信息', '维护时间与级别', '保存章节'],
        'description': '获奖经历章节用于维护论文、科研和获奖成果等记录。',
    },
    {
        'key': 'portal-v2-statement-section',
        'title': 'V2 个人陈述章节原型图',
        'prototype_style': 'portal-v2-section',
        'section': '学生申报',
        'route': 'PortalStatementSection',
        'operator': '门户考生',
        'page_title': '个人陈述',
        'summary': ['个人陈述', 'AI 思考', '简历上传与提交'],
        'form_fields': ['个人陈述', 'AI 问题理解', '行业观点', '简历附件', '提交声明'],
        'panel_title': '操作说明',
        'panel_lines': ['填写陈述和 AI 思考', '上传简历', '勾选声明', '最终提交申请'],
        'description': '个人陈述章节用于完成申请末端内容、简历上传和提交确认，是整个申请页的收口章节。',
    },
]


PROTOTYPE_GROUPS = [
    ('3.1 公共与认证界面', ['login-view', 'profile-view', 'portal-auth-view']),
    ('3.2 管理端业务界面', ['dashboard-view', 'recruitment-view', 'students-records-view', 'students-registrations-view', 'students-centers-view']),
    ('3.3 培养、学位与治理界面', ['training-plans-view', 'training-reports-view', 'training-outbound-view', 'degree-theses-view', 'degree-reviews-view', 'workflow-center-view', 'system-users-view', 'system-roles-view', 'dict-view']),
    ('3.4 学生门户界面', ['portal-home-view', 'portal-application-v2-view']),
    ('3.5 申请 V2 章节界面', ['portal-v2-basic-section', 'portal-v2-application-section', 'portal-v2-education-section', 'portal-v2-practice-section', 'portal-v2-english-section', 'portal-v2-family-section', 'portal-v2-achievement-section', 'portal-v2-statement-section']),
]


FIGURE_SPECS = [
    ('system-context', '系统上下文图'),
    ('module-map', '模块功能图'),
    ('business-flow', '业务主流程图'),
    ('portal-application-flow', '门户在线申请流程图'),
    ('approval-flow', '统一审批流程图'),
    ('deployment-topology', '部署拓扑图'),
    ('entity-relationship', '核心实体关系图'),
]


def ensure_asset_dirs() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    SRS_IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def add_cover(document: Document) -> None:
    title = document.add_paragraph()
    title.style = 'Title'
    base.configure_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = title.add_run(DOCUMENT_TITLE)
    base.apply_run_style(run, size=18, bold=True)

    subtitle = document.add_paragraph()
    base.configure_paragraph(subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = subtitle.add_run(f'{DOCUMENT_SUBTITLE}\n日期：{date.today().isoformat()}')
    base.apply_run_style(run, size=12)

    owner = document.add_paragraph()
    base.configure_paragraph(owner, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = owner.add_run('文档性质：正式 SRS 文档')
    base.apply_run_style(run, size=12)

    document.add_paragraph().add_run('')


def add_figure(document: Document, png_path: Path, caption: str, width_cm: float = 16.0) -> None:
    document.add_picture(str(png_path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph()
    base.configure_paragraph(caption_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=0)
    run = caption_paragraph.add_run(caption)
    base.apply_run_style(run, size=10)


def add_requirement_table(document: Document, rows: list[tuple[str, ...]]) -> None:
    base.add_table(document, ['编号', '功能项', '使用角色', '触发/前置', '系统响应'], rows)


def add_nonfunctional_table(document: Document, rows: list[tuple[str, str]]) -> None:
    base.add_table(document, ['编号', '需求描述'], rows)


def add_entity_model_tables(document: Document) -> None:
    for spec in SRS_ENTITY_MODEL_SPECS:
        base.add_heading(document, f"{spec['entity']}属性模型", 3)
        base.add_text_paragraph(document, str(spec['description']))
        base.add_table(document, ['属性（字段）', '属性类型', '描述', '是否为唯一关键属性'], list(spec['rows']))


def render_system_context_figure(svg_path: Path) -> None:
    dwg = base.create_drawing(svg_path, 1800, 980, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期管理系统上下文图', insert=(900, 50), text_anchor='middle', font_size=30, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    base.add_box(dwg, 120, 150, 260, 110, '管理端用户', '管理员、导师、学位秘书、评分人、面试官', '#eef6ff')
    base.add_box(dwg, 1420, 150, 260, 110, '门户考生/学生', '注册、登录、找回密码、在线申请', '#fff7e8')
    base.add_box(dwg, 580, 180, 640, 170, '博士生生命周期管理系统', '前端：Vue3 + Element Plus\n后端：FastAPI\n支持统一托管与前后端分离部署', '#eefcf2')
    base.add_arrow(dwg, (380, 205), (580, 205), '管理端访问')
    base.add_arrow(dwg, (1420, 205), (1220, 205), '门户访问')
    base.add_box(dwg, 220, 450, 260, 120, 'PostgreSQL', '业务主数据\n流程兼容数据\n治理数据', '#f4f0ff')
    base.add_box(dwg, 560, 450, 260, 120, 'Redis', '会话\n验证码\n缓存', '#fff5e8')
    base.add_box(dwg, 900, 450, 260, 120, '邮件服务', '注册验证码\n通知消息', '#fff4ea')
    base.add_box(dwg, 1240, 450, 260, 120, 'Web 服务器/Nginx', '静态资源\nHTTPS\n反向代理', '#eef6ff')
    base.add_arrow(dwg, (900, 350), (350, 450), '数据持久化')
    base.add_arrow(dwg, (900, 350), (690, 450), '缓存/会话')
    base.add_arrow(dwg, (900, 350), (1030, 450), '邮件通知')
    base.add_arrow(dwg, (900, 350), (1370, 450), '静态资源/代理')
    base.add_box(dwg, 520, 700, 760, 140, '边界说明', '系统对外主要暴露 Web 页面与 JSON API；对内依赖 PostgreSQL、Redis 和邮件服务。管理端与门户使用双认证域，统一审批与审计能力横向贯穿全部业务模块。', '#ffffff')
    dwg.save()


def render_module_map_figure(svg_path: Path) -> None:
    dwg = base.create_drawing(svg_path, 1800, 1020, '#fbfdff')
    dwg.add(dwg.text('博士生生命周期管理系统模块功能图', insert=(900, 50), text_anchor='middle', font_size=30, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    base.add_box(dwg, 700, 410, 400, 160, '系统核心', '统一认证\n统一流程\n统一日志\n统一配置', '#eef6ff')
    modules = [
        (180, 160, '经营总览', 'KPI、预警、阶段总览', '#fff7e8'),
        (700, 140, '招生管理', '计划、申请、评分、面试、录取', '#eefcf2'),
        (1220, 160, '学生与中心', '主档、注册学生、团队/中心', '#fff7e8'),
        (160, 660, '培养管理', '方案、科研报告、外出研修', '#eefcf2'),
        (700, 720, '学位管理', '论文、盲审、答辩、授位', '#fff7e8'),
        (1240, 660, '流程中心', '待办、动作、历史', '#eefcf2'),
        (180, 410, '系统治理', '用户、角色、字典、审计、集成', '#f4f0ff'),
        (1220, 410, '学生门户', '认证、首页、在线申请 V2', '#fff4ea'),
    ]
    for x, y, title, subtitle, fill in modules:
        base.add_box(dwg, x, y, 320, 120, title, subtitle, fill)
    centers = [
        ((500, 220), (700, 450)),
        ((1020, 220), (980, 450)),
        ((500, 720), (700, 530)),
        ((1020, 780), (980, 530)),
        ((500, 470), (700, 490)),
        ((1220, 470), (1100, 490)),
        ((540, 470), (700, 490)),
        ((1220, 470), (1100, 490)),
    ]
    for start, end in centers[:4]:
        base.add_arrow(dwg, start, end, '')
    base.add_arrow(dwg, (500, 470), (700, 490), '')
    base.add_arrow(dwg, (1220, 470), (1100, 490), '')
    base.add_arrow(dwg, (500, 470), (700, 490), '')
    base.add_arrow(dwg, (1220, 470), (1100, 490), '')
    base.add_box(dwg, 620, 880, 560, 90, '说明', '系统治理、统一认证、统一流程和统一日志为各业务模块提供横向支撑；学生门户与管理端在界面风格和认证域上保持隔离，但共享招生与生命周期主线。', '#ffffff')
    dwg.save()


def render_portal_application_flow_figure(svg_path: Path) -> None:
    dwg = base.create_drawing(svg_path, 1860, 860, '#fbfdff')
    dwg.add(dwg.text('门户在线申请流程图', insert=(930, 48), text_anchor='middle', font_size=30, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    steps = [
        ('认证登录', '注册/登录/找回密码'),
        ('进入首页', '查看进度与入口'),
        ('填写基本信息', '身份、联系方式、照片'),
        ('填写报名信息', '志愿、导师、渠道'),
        ('完善经历资料', '教育、实践、英语、家庭、成果'),
        ('个人陈述', 'AI 思考、简历上传'),
        ('保存草稿', '可反复返回修改'),
        ('最终提交', '勾选声明并提交招生申请'),
    ]
    x = 40
    for index, (title, subtitle) in enumerate(steps):
        base.add_box(dwg, x, 250, 200, 120, title, subtitle, '#eef6ff' if index % 2 == 0 else '#fff7e8')
        if index < len(steps) - 1:
            base.add_arrow(dwg, (x + 200, 310), (x + 230, 310), '')
        x += 230
    base.add_box(dwg, 220, 520, 520, 120, '附件管理', '在教育经历、英语能力和个人陈述等章节上传成绩单、证书和简历，系统按分类校验格式、大小与内容类型。', '#eefcf2')
    base.add_box(dwg, 820, 520, 420, 120, '状态指示', '左侧章节导航显示未开始、进行中和已完成状态，支持用户快速定位未完成项。', '#f4f0ff')
    base.add_box(dwg, 1320, 520, 360, 120, '结果输出', '提交后生成招生申请记录、业务键和后续招生评审入口。', '#fff4ea')
    dwg.save()


def render_deployment_topology_figure(svg_path: Path) -> None:
    dwg = base.create_drawing(svg_path, 1780, 920, '#fbfdff')
    dwg.add(dwg.text('部署拓扑图', insert=(890, 48), text_anchor='middle', font_size=30, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    base.add_box(dwg, 110, 160, 300, 120, '浏览器', '管理端和学生门户\n静态资源与 API 调用入口', '#eef6ff')
    base.add_box(dwg, 530, 160, 300, 120, 'Nginx / Web 服务器', 'HTTPS\n静态资源缓存\n反向代理', '#fff7e8')
    base.add_box(dwg, 950, 160, 300, 120, 'FastAPI 应用', '统一托管或 API 服务\n业务处理\n认证与流程', '#eefcf2')
    base.add_box(dwg, 1370, 160, 300, 120, '后台任务/通知', '邮件发送\n异步处理\n扩展接口', '#fff4ea')
    base.add_arrow(dwg, (410, 220), (530, 220), 'HTTPS')
    base.add_arrow(dwg, (830, 220), (950, 220), 'API/静态')
    base.add_arrow(dwg, (1250, 220), (1370, 220), '通知/异步')
    base.add_box(dwg, 370, 470, 280, 120, 'PostgreSQL', '主数据\n业务数据\n日志与流程兼容数据', '#f4f0ff')
    base.add_box(dwg, 760, 470, 280, 120, 'Redis', '会话\n验证码\n缓存', '#fff5e8')
    base.add_box(dwg, 1150, 470, 280, 120, 'SMTP/邮件网关', '验证码与通知消息', '#eef6ff')
    base.add_arrow(dwg, (1080, 280), (510, 470), 'SQL')
    base.add_arrow(dwg, (1080, 280), (900, 470), 'Redis 协议')
    base.add_arrow(dwg, (1370, 280), (1290, 470), 'SMTP')
    base.add_box(dwg, 250, 700, 1280, 110, '部署说明', '统一托管模式下 FastAPI 同时提供 API 与已构建前端资源；前后端分离模式下 Nginx 提供前端静态资源，FastAPI 仅提供 /api 和健康检查接口。', '#ffffff')
    dwg.save()


def render_portal_v2_main_figure(svg_path: Path, spec: dict[str, object]) -> None:
    dwg = base.create_drawing(svg_path, 1760, 1040, '#f7fbff')
    dwg.add(dwg.text(str(spec['title']), insert=(880, 42), text_anchor='middle', font_size=28, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    dwg.add(dwg.rect(insert=(40, 80), size=(1680, 900), rx=24, ry=24, fill='#ffffff', stroke='#d2deea', stroke_width=2))
    dwg.add(dwg.rect(insert=(40, 80), size=(1680, 78), rx=24, ry=24, fill='#edf4ff', stroke='none'))
    dwg.add(dwg.text('上海人工智能实验室 博士研究生在线申请', insert=(90, 128), font_size=28, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
    nav_x = 1080
    for label in ['首页', '招生信息', '在线申请', '申请进度']:
        dwg.add(dwg.text(label, insert=(nav_x, 128), font_size=18, font_family=base.SVG_FONT, fill='#4f6583'))
        nav_x += 120
    dwg.add(dwg.rect(insert=(90, 200), size=(280, 700), rx=20, ry=20, fill='#12304f', stroke='none'))
    section_y = 250
    for idx, label in enumerate(['基本信息', '报名信息', '教育经历', '实践经历', '英语能力', '家庭情况', '获奖经历', '个人陈述'], start=1):
        fill = '#1f4d7a' if idx == 1 else '#183956'
        dwg.add(dwg.rect(insert=(112, section_y), size=(236, 60), rx=14, ry=14, fill=fill, stroke='#2f5c86', stroke_width=1))
        dwg.add(dwg.text(str(idx), insert=(130, section_y + 38), font_size=18, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
        dwg.add(dwg.text(label, insert=(168, section_y + 38), font_size=18, font_family=base.SVG_FONT, fill='#ffffff'))
        section_y += 74
    dwg.add(dwg.rect(insert=(410, 200), size=(860, 700), rx=20, ry=20, fill='#fbfdff', stroke='#d6dfeb', stroke_width=2))
    dwg.add(dwg.text(str(spec['page_title']), insert=(450, 260), font_size=28, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
    field_y = 340
    for field in ['当前章节说明', '字段组 1', '字段组 2', '字段组 3', '附件或备注区域']:
        dwg.add(dwg.rect(insert=(450, field_y), size=(780, 88), rx=16, ry=16, fill='#ffffff', stroke='#d6dfeb', stroke_width=1))
        dwg.add(dwg.text(field, insert=(480, field_y + 50), font_size=18, font_family=base.SVG_FONT, fill='#5d6f85'))
        field_y += 108
    dwg.add(dwg.rect(insert=(1310, 200), size=(320, 700), rx=20, ry=20, fill='#f7f3ff', stroke='#ddd3ef', stroke_width=2))
    dwg.add(dwg.text('状态与操作', insert=(1360, 250), font_size=24, font_family=base.SVG_FONT, fill='#4f4a78', font_weight='bold'))
    panel_y = 320
    for item in spec['panel_lines']:
        dwg.add(dwg.rect(insert=(1340, panel_y), size=(260, 54), rx=12, ry=12, fill='#ffffff', stroke='#d9d1eb', stroke_width=1))
        dwg.add(dwg.text(str(item), insert=(1360, panel_y + 34), font_size=16, font_family=base.SVG_FONT, fill='#4f4a78'))
        panel_y += 70
    dwg.add(dwg.rect(insert=(1340, 810), size=(110, 52), rx=14, ry=14, fill='#eff3f8', stroke='none'))
    dwg.add(dwg.text('上一步', insert=(1365, 843), font_size=18, font_family=base.SVG_FONT, fill='#51687f'))
    dwg.add(dwg.rect(insert=(1470, 810), size=(110, 52), rx=14, ry=14, fill='#17304f', stroke='none'))
    dwg.add(dwg.text('下一步', insert=(1498, 843), font_size=18, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
    dwg.save()


def render_portal_v2_section_figure(svg_path: Path, spec: dict[str, object]) -> None:
    dwg = base.create_drawing(svg_path, 1680, 980, '#f7fbff')
    dwg.add(dwg.text(str(spec['title']), insert=(840, 42), text_anchor='middle', font_size=28, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
    dwg.add(dwg.rect(insert=(40, 80), size=(1600, 850), rx=24, ry=24, fill='#ffffff', stroke='#d2deea', stroke_width=2))
    dwg.add(dwg.rect(insert=(60, 120), size=(260, 740), rx=20, ry=20, fill='#12304f', stroke='none'))
    dwg.add(dwg.text('章节导航', insert=(120, 170), font_size=24, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
    nav_y = 230
    for label in ['基本信息', '报名信息', '教育经历', '实践经历', '英语能力', '家庭情况', '获奖经历', '个人陈述']:
        fill = '#1f4d7a' if label == spec['page_title'] else '#183956'
        dwg.add(dwg.rect(insert=(88, nav_y), size=(204, 46), rx=12, ry=12, fill=fill, stroke='#2f5c86', stroke_width=1))
        dwg.add(dwg.text(label, insert=(118, nav_y + 29), font_size=16, font_family=base.SVG_FONT, fill='#ffffff'))
        nav_y += 60
    dwg.add(dwg.rect(insert=(360, 120), size=(860, 740), rx=20, ry=20, fill='#fbfdff', stroke='#d6dfeb', stroke_width=2))
    dwg.add(dwg.text(str(spec['page_title']), insert=(400, 175), font_size=28, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
    dwg.add(dwg.text('根据当前章节显示差异化字段和附件区域', insert=(400, 212), font_size=16, font_family=base.SVG_FONT, fill='#5d6f85'))
    field_y = 270
    for field in spec['form_fields']:
        dwg.add(dwg.rect(insert=(400, field_y), size=(760, 72), rx=14, ry=14, fill='#ffffff', stroke='#d6dfeb', stroke_width=1))
        dwg.add(dwg.text(str(field), insert=(430, field_y + 44), font_size=18, font_family=base.SVG_FONT, fill='#5d6f85'))
        field_y += 92
    dwg.add(dwg.rect(insert=(1260, 120), size=(320, 740), rx=20, ry=20, fill='#f7f3ff', stroke='#ddd3ef', stroke_width=2))
    dwg.add(dwg.text(str(spec['panel_title']), insert=(1320, 174), font_size=24, font_family=base.SVG_FONT, fill='#4f4a78', font_weight='bold'))
    panel_y = 250
    for item in spec['panel_lines']:
        dwg.add(dwg.rect(insert=(1290, panel_y), size=(260, 52), rx=12, ry=12, fill='#ffffff', stroke='#d9d1eb', stroke_width=1))
        dwg.add(dwg.text(str(item), insert=(1310, panel_y + 33), font_size=16, font_family=base.SVG_FONT, fill='#4f4a78'))
        panel_y += 68
    dwg.add(dwg.rect(insert=(1290, 770), size=(110, 50), rx=14, ry=14, fill='#eff3f8', stroke='none'))
    dwg.add(dwg.text('保存', insert=(1325, 802), font_size=18, font_family=base.SVG_FONT, fill='#51687f'))
    dwg.add(dwg.rect(insert=(1420, 770), size=(130, 50), rx=14, ry=14, fill='#17304f', stroke='none'))
    dwg.add(dwg.text('下一章节', insert=(1450, 802), font_size=18, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
    dwg.save()


def render_prototype_figure(svg_path: Path, spec: dict[str, object]) -> None:
    style = str(spec.get('prototype_style', 'admin'))
    if style in {'auth', 'portal', 'student-form', 'admin'}:
        base.render_prototype_figure(svg_path, spec)
        return
    if style == 'login':
        dwg = base.create_drawing(svg_path, 1680, 980, '#f7fbff')
        dwg.add(dwg.text(str(spec['title']), insert=(840, 42), text_anchor='middle', font_size=28, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
        dwg.add(dwg.rect(insert=(40, 80), size=(1600, 860), rx=24, ry=24, fill='#ffffff', stroke='#d2deea', stroke_width=2))
        dwg.add(dwg.rect(insert=(40, 80), size=(780, 860), rx=24, ry=24, fill='#103250', stroke='none'))
        dwg.add(dwg.text('博士生生命周期管理系统', insert=(110, 180), font_size=36, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
        dwg.add(dwg.text('管理端统一登录入口', insert=(110, 235), font_size=28, font_family=base.SVG_FONT, fill='#dfeeff'))
        base.add_box(dwg, 110, 330, 620, 130, '登录用途', '统一承载后台角色登录与登录失败反馈', '#174467')
        base.add_box(dwg, 110, 500, 620, 150, '操作路径', '输入用户名和密码\n提交认证\n失败时使用对话框提示\n成功后跳回原目标页面', '#205684')
        dwg.add(dwg.rect(insert=(940, 220), size=(470, 520), rx=26, ry=26, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text('系统登录', insert=(1070, 290), font_size=30, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
        for index, field in enumerate(['用户名', '密码']):
            y = 370 + index * 96
            dwg.add(dwg.rect(insert=(1000, y), size=(350, 56), rx=14, ry=14, fill='#f8fafc', stroke='#d6dfeb', stroke_width=1))
            dwg.add(dwg.text(field, insert=(1025, y + 35), font_size=18, font_family=base.SVG_FONT, fill='#7a8798'))
        dwg.add(dwg.rect(insert=(1000, 580), size=(350, 60), rx=16, ry=16, fill='#17304f', stroke='none'))
        dwg.add(dwg.text('登录系统', insert=(1112, 618), font_size=22, font_family=base.SVG_FONT, fill='#ffffff', font_weight='bold'))
        dwg.add(dwg.rect(insert=(1000, 664), size=(350, 44), rx=12, ry=12, fill='#fff7e8', stroke='#eadcc3', stroke_width=1))
        dwg.add(dwg.text('登录失败通过对话框提示', insert=(1068, 692), font_size=16, font_family=base.SVG_FONT, fill='#9a6a28'))
        dwg.save()
        return
    if style == 'dashboard':
        dwg = base.create_drawing(svg_path, 1720, 1020, '#f7fbff')
        dwg.add(dwg.text(str(spec['title']), insert=(860, 42), text_anchor='middle', font_size=28, font_family=base.SVG_FONT, font_weight='bold', fill='#17304f'))
        dwg.add(dwg.rect(insert=(40, 80), size=(1640, 900), rx=24, ry=24, fill='#ffffff', stroke='#d2deea', stroke_width=2))
        dwg.add(dwg.rect(insert=(60, 120), size=(1600, 82), rx=18, ry=18, fill='#eef4fb', stroke='none'))
        dwg.add(dwg.text('经营总览', insert=(100, 172), font_size=30, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
        card_x = 100
        for title in ['开放招生计划', '学生总量', '科研报告待审', '盲审待办', '流程超期']:
            base.add_box(dwg, card_x, 240, 280, 96, title, 'KPI 卡片 / 趋势说明', '#fff7e8' if card_x % 2 == 0 else '#eef6ff')
            card_x += 300
        dwg.add(dwg.rect(insert=(100, 390), size=(980, 250), rx=18, ry=18, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text('生命周期阶段视图', insert=(130, 434), font_size=24, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
        step_x = 150
        for step in ['招生准备', '入学录取', '导师建立', '培养执行', '学位收口', '毕业归档']:
            base.add_box(dwg, step_x, 480, 130, 90, step, '阶段摘要', '#eefcf2')
            if step_x < 930:
                base.add_arrow(dwg, (step_x + 130, 525), (step_x + 160, 525), '')
            step_x += 160
        dwg.add(dwg.rect(insert=(1120, 390), size=(460, 250), rx=18, ry=18, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text('风险预警', insert=(1160, 434), font_size=24, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
        warn_y = 480
        for item in ['招生申请超时未审', '科研报告待审数量偏高', '盲审意见回收滞后']:
            dwg.add(dwg.rect(insert=(1150, warn_y), size=(400, 48), rx=12, ry=12, fill='#fff7e8', stroke='#eadcc3', stroke_width=1))
            dwg.add(dwg.text(item, insert=(1170, warn_y + 31), font_size=16, font_family=base.SVG_FONT, fill='#8f6530'))
            warn_y += 60
        dwg.add(dwg.rect(insert=(100, 690), size=(1480, 220), rx=18, ry=18, fill='#ffffff', stroke='#d6dfeb', stroke_width=2))
        dwg.add(dwg.text('图表区', insert=(130, 734), font_size=24, font_family=base.SVG_FONT, fill='#17304f', font_weight='bold'))
        dwg.add(dwg.rect(insert=(150, 770), size=(1350, 90), rx=14, ry=14, fill='#eef6ff', stroke='none'))
        dwg.add(dwg.text('柱线图：招生计划 / 学生规模 / 培养方案 / 报告待审 / 论文总量 / 流程待办', insert=(220, 826), font_size=20, font_family=base.SVG_FONT, fill='#4f6583'))
        dwg.save()
        return
    if style == 'portal-v2':
        render_portal_v2_main_figure(svg_path, spec)
        return
    if style == 'portal-v2-section':
        render_portal_v2_section_figure(svg_path, spec)
        return
    base.render_prototype_figure(svg_path, spec)


def generate_figure_assets() -> dict[str, Path]:
    generated: dict[str, Path] = {}
    for key, _ in FIGURE_SPECS:
        svg_path = SRS_IMAGES_DIR / f'{key}.svg'
        png_path = SRS_IMAGES_DIR / f'{key}.png'
        if key == 'system-context':
            render_system_context_figure(svg_path)
        elif key == 'module-map':
            render_module_map_figure(svg_path)
        elif key == 'business-flow':
            base.render_business_flow_figure(svg_path)
        elif key == 'portal-application-flow':
            render_portal_application_flow_figure(svg_path)
        elif key == 'approval-flow':
            base.render_approval_flow_figure(svg_path)
        elif key == 'deployment-topology':
            render_deployment_topology_figure(svg_path)
        else:
            base.render_entity_relationship_figure(svg_path)
        base.convert_svg_to_png(svg_path, png_path)
        generated[key] = png_path

    for spec in PROTOTYPE_SPECS:
        svg_path = SRS_IMAGES_DIR / f"{spec['key']}.svg"
        png_path = SRS_IMAGES_DIR / f"{spec['key']}.png"
        render_prototype_figure(svg_path, spec)
        base.convert_svg_to_png(svg_path, png_path)
        generated[str(spec['key'])] = png_path
    return generated


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal_style = document.styles['Normal']
    normal_style.font.name = base.BODY_FONT
    normal_style._element.rPr.rFonts.set(base.qn('w:eastAsia'), base.BODY_FONT_FALLBACK)
    normal_style.font.size = Pt(12)
    base.configure_heading_styles(document)


def add_toc(document: Document) -> None:
    base.add_heading(document, '目录', 1)
    toc_paragraph = document.add_paragraph()
    base.configure_paragraph(toc_paragraph, first_line_indent_cm=0)
    base.insert_toc(toc_paragraph)
    document.add_section(WD_SECTION_START.NEW_PAGE)


def build_document(assets: dict[str, Path]) -> Document:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_toc(document)

    base.add_heading(document, '1. 引言', 1)
    base.add_heading(document, '1.1 编写目的', 2)
    base.add_text_paragraph(document, '本文档用于形成博士生生命周期管理系统的正式软件需求规格说明书（SRS），对系统范围、界面、模块、流程、数据、外部接口、非功能要求和验收边界进行完整定义。')
    base.add_text_paragraph(document, '本文档用于为产品、设计、开发、测试和交付提供统一基线，确保系统范围、功能边界和验收口径在各参与角色之间保持一致。')

    base.add_heading(document, '1.2 适用范围', 2)
    base.add_text_paragraph(document, '本 SRS 适用于博士生生命周期管理系统的管理端与学生门户，覆盖招生、学生与研究中心、培养、学位、统一审批、系统治理、个人中心与门户在线申请。')
    base.add_text_paragraph(document, '本文档同时覆盖统一托管部署与前后端分离部署两种运行模式下的主要接口与约束。')

    base.add_heading(document, '1.3 术语与定义', 2)
    base.add_table(document, ['术语', '说明'], TERM_ROWS)

    base.add_heading(document, '1.4 参考资料', 2)
    base.add_table(document, ['参考资料', '用途'], REFERENCE_ROWS)

    base.add_heading(document, '1.5 文档约定', 2)
    base.add_text_paragraph(document, '本文档采用“总体说明、外部接口需求、功能需求、业务流程与规则、数据需求、非功能需求、验收要求”的标准 SRS 组织方式。')
    base.add_text_paragraph(document, '所有流程图、模块功能图和界面原型图均通过 SVG 绘制，再转换为 PNG 后插入文档，以便后续持续维护与再生成。')

    base.add_heading(document, '2. 总体说明', 1)
    base.add_heading(document, '2.1 产品目标', 2)
    base.add_text_paragraph(document, '系统用于支撑博士研究生从招生报名到培养、学位、毕业归档的全生命周期管理，并通过学生门户向考生提供自助注册、认证和在线申请能力。')
    base.add_text_paragraph(document, '系统同时满足平台管理员的治理需求、导师与学位秘书的业务办理需求，以及审批角色的统一待办处理需求。')

    base.add_heading(document, '2.2 用户角色', 2)
    base.add_table(document, ['角色', '职责'], ROLE_ROWS)

    base.add_heading(document, '2.3 业务约束', 2)
    base.add_table(document, ['约束主题', '说明'], CONSTRAINT_ROWS)

    base.add_heading(document, '2.4 运行环境', 2)
    base.add_table(document, ['层级', '环境/技术', '说明'], ENVIRONMENT_ROWS)

    base.add_heading(document, '2.5 系统上下文', 2)
    add_figure(document, assets['system-context'], '图 2-1 系统上下文图')
    base.add_text_paragraph(document, '系统上下文图用于说明外部参与者、系统边界和基础依赖。管理端与门户作为两套独立交互面分别接入系统，业务与治理数据通过 PostgreSQL 持久化，会话与验证码相关场景通过 Redis 支撑，邮件服务用于通知和验证码发送。')

    base.add_heading(document, '2.6 模块功能图', 2)
    add_figure(document, assets['module-map'], '图 2-2 模块功能图')
    base.add_table(document, ['模块', '范围', '主要能力'], MODULE_ROWS)

    base.add_heading(document, '2.7 页面清单', 2)
    base.add_table(document, ['路由', '页面定位', '主要能力', '备注'], PAGE_CATALOG_ROWS)

    base.add_heading(document, '2.8 操作模式', 2)
    base.add_table(document, ['模式', '交互结构', '适用范围'], OPERATION_MODE_ROWS)
    for title, items in INTERFACE_OPERATION_SCENARIOS.items():
        base.add_heading(document, title, 3)
        base.add_operation_list(document, '建议操作步骤：', items)

    base.add_heading(document, '3. 外部接口需求', 1)
    base.add_heading(document, '3.1 用户界面原型与操作方式', 2)
    base.add_text_paragraph(document, '本节依据当前前端实现给出全量页面级原型输入，并补充关键操作方式，用于支撑 UI/UE 设计、交互走查和后续验收。')
    figure_index = 1
    prototype_map = {str(spec['key']): spec for spec in PROTOTYPE_SPECS}
    for group_title, keys in PROTOTYPE_GROUPS:
        base.add_heading(document, group_title, 3)
        for key in keys:
            spec = prototype_map[key]
            base.add_heading(document, str(spec['title']), 4)
            add_figure(document, assets[key], f'图 3-{figure_index} {spec["title"]}', width_cm=16.5)
            figure_index += 1
            base.add_text_paragraph(document, str(spec['description']))
            table_rows = [
                ('路由/界面', str(spec['route'])),
                ('使用角色', str(spec['operator'])),
                ('页面标题', str(spec['page_title'])),
                ('摘要能力', '、'.join(str(item) for item in spec['summary'])),
            ]
            base.add_table(document, ['属性', '内容'], table_rows)
            base.add_operation_list(document, '操作方式：', list(spec['panel_lines']))

    base.add_heading(document, '3.2 软件接口需求', 2)
    base.add_table(document, ['接口对象', '协议/形式', '用途'], SOFTWARE_INTERFACE_ROWS)

    base.add_heading(document, '3.3 通信接口需求', 2)
    base.add_table(document, ['主题', '方式', '说明'], COMMUNICATION_ROWS)

    base.add_heading(document, '3.4 硬件与部署接口需求', 2)
    base.add_table(document, ['对象', '建议配置/环境', '说明'], HARDWARE_ROWS)
    add_figure(document, assets['deployment-topology'], '图 3-26 部署拓扑图')
    base.add_table(document, ['部署模式', '结构说明', '适用场景'], DEPLOYMENT_ROWS)

    base.add_heading(document, '4. 功能需求', 1)
    base.add_text_paragraph(document, '本节按业务模块给出正式功能需求，以“编号 + 功能项 + 使用角色 + 触发/前置 + 系统响应”的结构表达，确保需求可实现、可验证和可追踪。')
    for section_title, rows in FUNCTIONAL_REQUIREMENTS.items():
        base.add_heading(document, section_title, 2)
        add_requirement_table(document, rows)

    base.add_heading(document, '5. 业务流程与业务规则', 1)
    base.add_heading(document, '5.1 生命周期业务主流程', 2)
    add_figure(document, assets['business-flow'], '图 5-1 生命周期业务主流程图')
    base.add_table(document, ['阶段', '触发点', '核心处理', '阶段产出'], BUSINESS_FLOW_ROWS)

    base.add_heading(document, '5.2 门户在线申请流程', 2)
    add_figure(document, assets['portal-application-flow'], '图 5-2 门户在线申请流程图')
    base.add_text_paragraph(document, '在线申请流程体现了门户认证、分章节填写、附件上传、草稿保存与最终提交的完整闭环，当前在线申请页采用章节导航与长表单结合的结构。')

    base.add_heading(document, '5.3 审批流程', 2)
    add_figure(document, assets['approval-flow'], '图 5-3 统一审批流程图')
    base.add_table(document, ['流程类型', '关键节点', '主要角色', '主要入口'], APPROVAL_FLOW_ROWS)

    base.add_heading(document, '5.4 业务规则', 2)
    base.add_table(document, ['编号', '规则描述'], BUSINESS_RULE_ROWS)

    base.add_heading(document, '6. 数据需求', 1)
    base.add_heading(document, '6.1 数据分层概览', 2)
    base.add_table(document, ['数据层', '主要对象', '说明'], DATA_SUMMARY_ROWS)

    base.add_heading(document, '6.2 核心实体关系', 2)
    add_figure(document, assets['entity-relationship'], '图 6-1 核心实体关系图')
    base.add_table(document, ['实体', '说明', '主要关联'], ENTITY_ROWS)

    base.add_heading(document, '6.3 核心实体属性模型', 2)
    base.add_text_paragraph(document, '以下属性模型覆盖系统的关键主数据、业务对象、治理对象和流程对象，用于支撑数据库设计、接口设计和页面模型设计。')
    add_entity_model_tables(document)

    base.add_heading(document, '6.4 数据约束', 2)
    base.add_text_paragraph(document, '关键对象应保持唯一标识稳定，例如学生主档、角色编码、门户账号、业务键和字典值。')
    base.add_text_paragraph(document, '需要审批托管的业务对象应通过业务键与统一流程任务进行双向关联，不得出现脱钩状态。')
    base.add_text_paragraph(document, '枚举型字段应优先从系统字典加载，以便后续扩展和治理。')

    base.add_heading(document, '7. 非功能需求', 1)
    for title, rows in NONFUNCTIONAL_ROWS.items():
        base.add_heading(document, title, 2)
        add_nonfunctional_table(document, rows)

    base.add_heading(document, '8. 验收要求', 1)
    base.add_text_paragraph(document, '本节给出需求验收的主要关注维度，作为测试设计和交付验收的基线。')
    base.add_table(document, ['验收主题', '验收重点'], ACCEPTANCE_ROWS)

    base.add_heading(document, '附录 A 文档维护说明', 1)
    base.add_text_paragraph(document, '当系统新增页面、模块、审批节点或实体模型时，应同步更新本脚本中的原型、流程图、实体属性模型和功能需求表。')
    base.add_text_paragraph(document, '所有图形资产均保存在 documents/SRS_IMAGES 目录，可重新生成 SVG/PNG 后再批量更新文档。')

    document.core_properties.title = '博士生生命周期管理系统 软件需求规格说明书（SRS版）'
    document.core_properties.subject = 'SRS 需求规格说明'
    document.core_properties.author = 'GitHub Copilot'
    return document


def main() -> None:
    ensure_asset_dirs()
    assets = generate_figure_assets()
    document = build_document(assets)
    try:
        document.save(OUTPUT_PATH)
        print(f'Generated: {OUTPUT_PATH}')
    except PermissionError:
        document.save(FALLBACK_OUTPUT_PATH)
        print(f'Generated: {FALLBACK_OUTPUT_PATH}')


if __name__ == '__main__':
    main()