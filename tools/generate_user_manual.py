from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlencode
from urllib.request import Request, urlopen
import json
import time

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.support.ui import WebDriverWait


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / 'CMMI3_Documents' / '用户手册'
IMAGE_DIR = OUTPUT_DIR / 'images'
FRONTEND_URL = 'http://127.0.0.1:5173'
API_BASE = 'http://127.0.0.1:8000/api/v1'
EDGE_PATH = Path(r'C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe')


@dataclass(frozen=True)
class CaptureItem:
    slug: str
    title: str
    route: str
    note: str


@dataclass(frozen=True)
class RoleManual:
    role_name: str
    account: str
    password: str
    scope: str
    audience: str
    permissions: list[str]
    capabilities: list[str]
    precautions: list[str]
    captures: list[CaptureItem]


COMMON_CAPTURES = [
    CaptureItem('login', '登录页', '/login', '用户输入账号密码后登录系统，成功后会按照会话状态进入目标工作页。'),
]


ROLE_MANUALS = [
    RoleManual(
        role_name='平台管理员',
        account='admin',
        password='Admin@123456',
        scope='系统治理与全业务协同',
        audience='研究生院管理人员、平台运维人员、业务总控人员',
        permissions=['*'],
        capabilities=[
            '查看并维护招生、学生、培养、学位、流程与系统治理全部模块。',
            '维护系统用户、角色、字典、接口与审计日志。',
            '通过经营总览快速查看生命周期覆盖、流程待办与预警信息。',
        ],
        precautions=[
            '平台管理员拥有全量权限，涉及主数据与流程动作时应先确认业务责任人。',
            '批量修改系统治理配置前，建议先记录原始条件与筛选口径。',
        ],
        captures=[
            CaptureItem('admin_dashboard', '管理员经营总览', '/dashboard', '经营总览集中展示学生总量、招生计划、在途审批和预警事项。'),
            CaptureItem('admin_recruitment', '管理员招生计划页', '/recruitment', '支持查看招生计划、申请池、阶段状态与处理入口。'),
            CaptureItem('admin_students', '管理员学生主档页', '/students/records', '支持查看学生状态、导师归属、团队归属与基础档案。'),
            CaptureItem('admin_training', '管理员培养方案页', '/training/plans', '支持查看培养方案周期、计划状态与明细。'),
            CaptureItem('admin_degree', '管理员论文主档页', '/degree/theses', '支持查看论文状态、答辩安排与学位进度。'),
            CaptureItem('admin_workflow', '管理员审批中心页', '/workflow/tasks', '审批中心提供任务详情、动作执行与历史轨迹。'),
            CaptureItem('admin_system_users', '管理员系统用户页', '/system/users', '系统用户页用于维护账号、角色与状态。'),
            CaptureItem('admin_dict_types', '管理员字典类型页', '/system/dict-types', '字典页用于维护下拉项、状态字典与受控数据。'),
        ],
    ),
    RoleManual(
        role_name='导师',
        account='liu.ya',
        password='LiuYa@2026',
        scope='培养与学位协同',
        audience='导师、课题组负责人',
        permissions=['dashboard:read', 'students:read', 'training:read', 'training:write', 'degree:read', 'workflow:read', 'workflow:write'],
        capabilities=[
            '查看经营总览、学生主档、培养方案、科研报告、外出研修、论文主档与流程待办。',
            '可处理科研报告、外出研修等流程任务，并推进培养相关业务。',
            '可在个人空间查看和维护个人资料。',
        ],
        precautions=[
            '导师端重点关注科研报告、外出研修和流程待办，避免逾期。',
            '导师角色不具备系统治理维护权限。',
        ],
        captures=[
            CaptureItem('advisor_dashboard', '导师经营总览', '/dashboard', '导师登录后可从经营总览快速识别待办与学生分布。'),
            CaptureItem('advisor_students', '导师学生主档页', '/students/records', '导师可查看名下学生状态、团队与导师信息。'),
            CaptureItem('advisor_reports', '导师科研报告页', '/training/reports', '导师可查看科研报告状态并结合流程待办完成审阅。'),
            CaptureItem('advisor_workflow', '导师审批中心页', '/workflow/tasks', '导师在审批中心完成任务处理与轨迹查看。'),
        ],
    ),
    RoleManual(
        role_name='学位秘书',
        account='zhou.qing',
        password='ZhouQing@2026',
        scope='学位管理与流程审批',
        audience='学位秘书、学位办工作人员',
        permissions=['dashboard:read', 'degree:read', 'degree:write', 'workflow:read', 'workflow:write'],
        capabilities=[
            '查看经营总览、论文主档、盲审意见和流程待办。',
            '可执行学位模块维护动作并处理学位相关流程任务。',
        ],
        precautions=[
            '学位秘书角色不具备招生、学生和系统治理维护权限。',
            '涉及论文与盲审数据修改时，应先核对学生与论文主档状态。',
        ],
        captures=[
            CaptureItem('secretary_dashboard', '学位秘书经营总览', '/dashboard', '学位秘书可通过看板快速识别学位收口阶段任务。'),
            CaptureItem('secretary_theses', '学位秘书论文主档页', '/degree/theses', '论文主档页用于查看论文状态与答辩安排。'),
            CaptureItem('secretary_reviews', '学位秘书盲审意见页', '/degree/reviews', '盲审意见页用于查看评审结果与处理状态。'),
            CaptureItem('secretary_workflow', '学位秘书审批中心页', '/workflow/tasks', '审批中心集中承接学位流程待办。'),
        ],
    ),
    RoleManual(
        role_name='评分人',
        account='he.lin',
        password='HeLin@2026',
        scope='招生只读评审协同',
        audience='招生评分人员',
        permissions=['dashboard:read', 'recruitment:read'],
        capabilities=[
            '查看经营总览和招生计划页。',
            '用于了解申请池、评分进度和当前阶段。',
        ],
        precautions=[
            '评分人仅具备查看权限，不执行招生写操作。',
        ],
        captures=[
            CaptureItem('reviewer_dashboard', '评分人经营总览', '/dashboard', '评分人可先通过经营总览查看招生与待办概况。'),
            CaptureItem('reviewer_recruitment', '评分人招生计划页', '/recruitment', '评分人主要在招生计划页查看申请与阶段情况。'),
        ],
    ),
    RoleManual(
        role_name='面试官',
        account='cao.bo',
        password='CaoBo@2026',
        scope='招生执行与面试协同',
        audience='招生面试官、招生执行人员',
        permissions=['dashboard:read', 'recruitment:read', 'recruitment:write'],
        capabilities=[
            '查看经营总览和招生计划页。',
            '在招生执行链路中承担面试安排、录入结果等协同动作。',
        ],
        precautions=[
            '面试官聚焦招生模块，不涉及学生、培养、学位和系统治理。',
        ],
        captures=[
            CaptureItem('interviewer_dashboard', '面试官经营总览', '/dashboard', '面试官可通过经营总览定位招生录取进度。'),
            CaptureItem('interviewer_recruitment', '面试官招生计划页', '/recruitment', '面试官在招生页查看面试待安排和面试完成申请。'),
        ],
    ),
    RoleManual(
        role_name='中心HRBP',
        account='yang.qin',
        password='YangQin@2026',
        scope='跨部门学生与培养协同',
        audience='人力资源协同人员',
        permissions=['dashboard:read', 'students:read', 'training:read'],
        capabilities=[
            '查看经营总览、学生主档与培养方案。',
            '用于掌握学生状态、团队归属和培养执行情况。',
        ],
        precautions=[
            'HRBP 角色以查询为主，不参与流程动作执行。',
        ],
        captures=[
            CaptureItem('hrbp_dashboard', 'HRBP 经营总览', '/dashboard', 'HRBP 可查看生命周期总体态势和预警信息。'),
            CaptureItem('hrbp_students', 'HRBP 学生主档页', '/students/records', '学生主档页用于查看学生状态、导师与团队归属。'),
            CaptureItem('hrbp_training', 'HRBP 培养方案页', '/training/plans', '培养方案页用于查看培养周期和计划执行概况。'),
        ],
    ),
    RoleManual(
        role_name='党群负责人',
        account='sun.wei',
        password='SunWei@2026',
        scope='学生查询与审计查看',
        audience='党群工作协同人员',
        permissions=['dashboard:read', 'students:read', 'audit:read'],
        capabilities=[
            '查看经营总览、学生主档、操作日志与同步日志。',
            '用于了解学生状态变化和治理留痕。',
        ],
        precautions=[
            '党群负责人不具备招生、培养、学位和系统写权限。',
        ],
        captures=[
            CaptureItem('party_dashboard', '党群负责人经营总览', '/dashboard', '党群负责人可先查看总体态势与预警。'),
            CaptureItem('party_students', '党群负责人学生主档页', '/students/records', '用于查询学生状态与基础档案。'),
            CaptureItem('party_logs', '党群负责人操作日志页', '/system/operation-logs', '可查看操作留痕与治理记录。'),
        ],
    ),
]


def ensure_output_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)


def create_driver() -> webdriver.Edge:
    options = EdgeOptions()
    options.binary_location = str(EDGE_PATH)
    options.add_argument('--headless=new')
    options.add_argument('--disable-gpu')
    options.add_argument('--hide-scrollbars')
    options.add_argument('--window-size=1440,1200')
    options.add_argument('--force-device-scale-factor=1')
    return webdriver.Edge(options=options)


def get_access_token(username: str, password: str) -> str:
    payload = urlencode({'username': username, 'password': password}).encode('utf-8')
    request = Request(f'{API_BASE}/auth/token', data=payload, headers={'Content-Type': 'application/x-www-form-urlencoded'})
    with urlopen(request, timeout=20) as response:
        return json.loads(response.read().decode('utf-8'))['access_token']


def wait_for_page(driver: webdriver.Edge, route: str, timeout: int = 20) -> None:
    WebDriverWait(driver, timeout).until(lambda current_driver: current_driver.execute_script('return document.readyState') == 'complete')
    WebDriverWait(driver, timeout).until(lambda current_driver: route in current_driver.current_url)
    WebDriverWait(driver, timeout).until(lambda current_driver: current_driver.find_element(By.TAG_NAME, 'body'))
    time.sleep(2.0)


def capture_page(username: str, password: str, route: str, output_file: Path) -> None:
    driver = create_driver()
    try:
        encoded_username = urlencode({'username': username})
        encoded_password = urlencode({'pwd': password})
        separator = '&' if '?' in route else '?'
        driver.get(f'{FRONTEND_URL}{route}{separator}{encoded_username}&{encoded_password}')
        wait_for_page(driver, route)
        WebDriverWait(driver, 20).until(lambda current_driver: '/login' not in current_driver.current_url)
        WebDriverWait(driver, 20).until(lambda current_driver: 'layout-shell' in current_driver.page_source)
        time.sleep(2.5)
        driver.save_screenshot(str(output_file))
    finally:
        driver.quit()


def capture_public_page(route: str, output_file: Path) -> None:
    driver = create_driver()
    try:
        driver.get(f'{FRONTEND_URL}{route}')
        wait_for_page(driver, route)
        driver.save_screenshot(str(output_file))
    finally:
        driver.quit()


def set_default_font(style) -> None:
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def format_paragraph(paragraph, bold: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(6)
    for run_item in paragraph.runs:
        run_item.font.name = 'Microsoft YaHei'
        run_item.font.size = Pt(12)
        run_item._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run_item.bold = bold if bold else run_item.bold


def add_table_border(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D9E1F2')
        borders.append(border)
    tbl_pr.append(borders)


def build_docx() -> Path:
    document = Document()
    normal_style = document.styles['Normal']
    set_default_font(normal_style)

    for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        if style_name in document.styles:
            set_default_font(document.styles[style_name])

    title = document.add_paragraph()
    title_run = title.add_run('博士生生命周期管理系统用户手册')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Microsoft YaHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    title.paragraph_format.line_spacing = 1.2
    title.alignment = 1

    subtitle = document.add_paragraph('版本：V1.0    日期：2026-04-07')
    subtitle.alignment = 1
    format_paragraph(subtitle)

    intro = document.add_paragraph(
        '本手册面向博士生生命周期管理系统的不同使用角色，说明登录方式、角色权限、功能入口与典型操作路径。'
        '当前系统覆盖经营总览、招生管理、学生管理、培养管理、学位管理、流程审批与系统治理等模块。'
    )
    format_paragraph(intro)

    document.add_heading('1. 系统概述', level=1)
    paragraphs = [
        '系统访问地址：前端 http://127.0.0.1:5173，后端接口 http://127.0.0.1:8000/api/v1。',
        '推荐浏览器：Microsoft Edge 最新版本，分辨率建议不低于 1440 × 900。',
        '登录成功后，系统会根据当前角色展示可访问菜单；无权限功能不会在侧边栏显示。',
    ]
    for text in paragraphs:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(text)
        format_paragraph(p)

    document.add_heading('2. 登录与界面说明', level=1)
    login_capture = IMAGE_DIR / 'login.png'
    if login_capture.exists():
        document.add_picture(str(login_capture), width=Cm(15.8))
        cap = document.add_paragraph('图 2-1 登录页')
        cap.alignment = 1
        format_paragraph(cap)

    login_steps = [
        '在登录页输入账号与密码。',
        '点击“登录”按钮进入系统。',
        '进入系统后，左侧为菜单导航，顶部显示当前页面标题、用户信息和退出登录按钮。',
    ]
    for index, text in enumerate(login_steps, start=1):
        p = document.add_paragraph(f'{index}. {text}')
        format_paragraph(p)

    document.add_heading('3. 角色与功能矩阵', level=1)
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    add_table_border(table)
    header = table.rows[0].cells
    header[0].text = '角色'
    header[1].text = '示例账号'
    header[2].text = '业务范围'
    header[3].text = '主要模块'
    for cell in header:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for role in ROLE_MANUALS:
        row = table.add_row().cells
        row[0].text = role.role_name
        row[1].text = role.account
        row[2].text = role.scope
        row[3].text = '；'.join(role.capabilities)
        for cell in row:
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph)

    for role in ROLE_MANUALS:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(f'4. {role.role_name}使用说明', level=1)

        meta = [
            f'适用对象：{role.audience}',
            f'示例账号：{role.account}',
            f'业务范围：{role.scope}',
            f'权限集合：{", ".join(role.permissions)}',
        ]
        for text in meta:
            p = document.add_paragraph(style='List Bullet')
            p.add_run(text)
            format_paragraph(p)

        document.add_heading('4.1 主要功能', level=2)
        for item in role.capabilities:
            p = document.add_paragraph(style='List Bullet')
            p.add_run(item)
            format_paragraph(p)

        document.add_heading('4.2 操作建议', level=2)
        for item in role.precautions:
            p = document.add_paragraph(style='List Bullet')
            p.add_run(item)
            format_paragraph(p)

        document.add_heading('4.3 典型页面与入口', level=2)
        for index, capture in enumerate(role.captures, start=1):
            image_path = IMAGE_DIR / f'{capture.slug}.png'
            heading = document.add_paragraph(f'{index}. {capture.title}')
            format_paragraph(heading)
            if image_path.exists():
                document.add_picture(str(image_path), width=Cm(15.8))
                cap = document.add_paragraph(f'图 4-{index} {capture.title}')
                cap.alignment = 1
                format_paragraph(cap)
            route_text = document.add_paragraph(f'访问路由：{capture.route}')
            format_paragraph(route_text)
            note = document.add_paragraph(capture.note)
            format_paragraph(note)

        document.add_heading('4.4 推荐使用路径', level=2)
        for step_index, capture in enumerate(role.captures, start=1):
            p = document.add_paragraph(f'{step_index}. 进入“{capture.title}”对应页面，按照页面筛选条件、列表和详情区域完成查看或处理。')
            format_paragraph(p)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading('5. 常见问题', level=1)
    faqs = [
        ('登录成功后菜单为空', '通常是权限尚未加载完成或会话已失效。可刷新页面，若仍异常请重新登录。'),
        ('驾驶舱数据加载失败', '请先确认后端服务正常、当前角色具备 dashboard:read 权限，并检查浏览器是否仍保留旧 token。'),
        ('流程任务无法处理', '请确认当前角色具备 workflow:write 权限，且任务状态仍为处理中。'),
        ('截图与实际界面不完全一致', '本手册基于当前模拟数据和默认账号生成，若生产配置、字典或权限发生变化，展示数据会随之调整。'),
    ]
    for question, answer in faqs:
        p = document.add_paragraph()
        p.add_run(f'问：{question}').bold = True
        format_paragraph(p)
        a = document.add_paragraph(f'答：{answer}')
        format_paragraph(a)

    output_file = OUTPUT_DIR / '博士生生命周期管理系统用户手册.docx'
    document.save(output_file)
    return output_file


def main() -> None:
    ensure_output_dirs()
    capture_public_page('/login', IMAGE_DIR / 'login.png')


    output_file = build_docx()
    print(output_file)


if __name__ == '__main__':
    main()